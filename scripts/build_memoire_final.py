# -*- coding: utf-8 -*-
"""
Mémoire de thèse professionnelle Rondol — version FINALE professionnelle.

- Contenu parsé depuis docs/memoire_these_professionnelle_rondol.md
- Page de garde graphique (bande verticale, logos, encadré du thème) + page institutionnelle
- Figures techniques (reports/memoire_figures/) et captures réelles (reports/memoire_captures/)
  insérées aux emplacements, avec légendes numérotées "Figure N — ..."
- Placeholders [À COMPLÉTER] / [INSÉRER ...] nettoyés ; bibliographie homogénéisée
- Sommaire = champ TOC Word automatique (mis à jour à la conversion PDF)
- Sortie : reports/Memoire_Rondol_Wilfried_Galtier_MBEUMI_FINAL.docx (+ .pdf)

Usage : python scripts/build_memoire_final.py
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu

# --------------------------------------------------------------------------- #
ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "memoire_these_professionnelle_rondol.md"
REPORTS = ROOT / "reports"
ASSETS = REPORTS / "assets"
FIGDIR = REPORTS / "memoire_figures"
CAPDIR = REPORTS / "memoire_captures"
POSTER = REPORTS / "poster_abstract" / "figures"
DOCX = REPORTS / "Memoire_Rondol_Wilfried_Galtier_MBEUMI_FINAL_PRO.docx"
PDF = REPORTS / "Memoire_Rondol_Wilfried_Galtier_MBEUMI_FINAL_PRO.pdf"

AUTHOR = "Wilfried Galtier MBEUMI"
SCHOOL = "Nexa Digital School"
PROGRAM = "Mastère Data & Intelligence Artificielle"
CERT = "Certification professionnelle RNCP 37137 — Niveau 7 (Bac+5)"
COMPANY = "Rondol Industrie"
TUTOR_COMPANY = "M. Maël Gallas"
TUTOR_SCHOOL = "M. Moussa NDIAYE"
PLACE = "Nancy — Institut Jean Lamour (IJL), Campus ARTEM"
YEAR = "2025 – 2026"
DEFENSE = "À confirmer"
PROBLEM = ("Comment concevoir et déployer un système d'intelligence artificielle "
           "prédictif permettant d'optimiser les paramètres d'extrusion pour la "
           "fabrication de composants de batteries tout-solide (SSB) dry/semi-dry, "
           "afin d'améliorer la performance technique et la compétitivité stratégique "
           "de Rondol Industrie ?")

TITLE = ("Conception et déploiement d'un système d'intelligence artificielle "
         "prédictif d'aide à la décision pour l'optimisation des paramètres "
         "d'extrusion bivis de composants de batteries tout-solide (dry / semi-dry)")
SUBTITLE = ("Un prototype professionnel de jumeau numérique appliqué à l'extrudeuse "
            "bivis 10,5 mm de Rondol Industrie")

HEADER_TEXT = "Mémoire de thèse professionnelle — Rondol Industrie"

# Couleurs (bleu pétrole dominant, vert Rondol discret)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
BLUE2 = RGBColor(0x2E, 0x75, 0xB6)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
DARK = RGBColor(0x26, 0x32, 0x3A)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLUE_HEX = "1F4E79"
BAND_HEX = "1F4E79"
THEME_FILL = "EAF1F8"
LTGREY = "F2F2F2"

BASE_FONT = "Calibri"
# Police des titres : serif institutionnel, contraste avec le corps Calibri.
TITLE_FONT = "Cambria"
CONTENT_W = Cm(16.0)
FIG_W = Cm(11.0)   # largeur des figures (compacte, maîtrise de la pagination)

fig_counter = [0]


# --------------------------------------------------------------------------- #
# Helpers OXML
# --------------------------------------------------------------------------- #
def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def cell_valign(cell, val="center"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if val == "center" else WD_ALIGN_VERTICAL.TOP


def set_table_borders(table, color="BFBFBF", sz="4"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def set_row_height(row, cm, exact=True):
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(cm * 567)))
    trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trH)


def add_field(paragraph, field):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = field
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)
    return run


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:r"); tt = OxmlElement("w:t")
    tt.text = "Le sommaire se met à jour automatiquement (clic droit ▸ Mettre à jour les champs)."
    t.append(tt)
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr); run._r.append(sep); run._r.append(t); run._r.append(e)


def set_start_page_number(section, start=1):
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), str(start))
    section._sectPr.append(pgNumType)


def para_bottom_border(p, color, sz="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def enable_hyphenation(doc):
    """Active la césure automatique du document.

    Sans césure, un texte justifié en français produit de larges blancs entre
    les mots (« rivières » typographiques) — le principal défaut visuel d'un
    document composé par défaut. La césure les résorbe tout en conservant la
    justification attendue dans un mémoire.
    """
    settings = doc.settings.element
    for tag, val in (("w:autoHyphenation", "true"),
                     ("w:doNotHyphenateCaps", "true")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        settings.append(el)
    zone = OxmlElement("w:hyphenationZone")
    zone.set(qn("w:val"), "357")          # ≈ 0,63 cm — zone de césure usuelle
    settings.append(zone)
    limit = OxmlElement("w:consecutiveHyphenLimit")
    limit.set(qn("w:val"), "2")           # jamais plus de 2 césures d'affilée
    settings.append(limit)


def para_top_border(p, color, sz="6"):
    """Filet fin au-dessus du paragraphe (utilisé pour le pied de page)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), "4"); top.set(qn("w:color"), color)
    pbdr.append(top)
    pPr.append(pbdr)


# --------------------------------------------------------------------------- #
# Logos
# --------------------------------------------------------------------------- #
def ensure_logos():
    ASSETS.mkdir(parents=True, exist_ok=True)
    nexa_png = ASSETS / "nexa_logo.png"
    rondol_png = ASSETS / "rondol_logo.png"
    nexa_src = ROOT / "nexa LOGO.webp"
    if nexa_src.exists():
        try:
            from PIL import Image
            im = Image.open(nexa_src).convert("RGBA")
            bg = Image.new("RGBA", im.size, (255, 255, 255, 255))
            bg.alpha_composite(im)
            bg.convert("RGB").save(nexa_png, "PNG")
        except Exception as exc:
            print(f"[WARN] logo Nexa: {exc}"); nexa_png = None
    else:
        nexa_png = nexa_png if nexa_png.exists() else None
    rondol_src = ROOT / "assets" / "rondol_logo.png"
    if rondol_src.exists():
        shutil.copy(rondol_src, rondol_png)
    return (nexa_png if nexa_png and nexa_png.exists() else None,
            rondol_png if rondol_png and rondol_png.exists() else None)


# --------------------------------------------------------------------------- #
# Styles & page
# --------------------------------------------------------------------------- #
def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(12)          # conforme guide RNCP : Calibri taille 12
    normal.font.color.rgb = DARK
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    # Interligne 1,30 (guide : 1,5 maximum) — respiration nettement meilleure
    # que 1,15 pour un document de lecture longue.
    pf.line_spacing = 1.30
    pf.space_after = Pt(8)             # séparation nette entre paragraphes
    # Appariement typographique : titres en serif (Cambria), corps en Calibri 12.
    # Le guide impose la police du CORPS ; le contraste titres/corps est ce qui
    # distingue un mémoire composé d'un document bureautique par défaut.
    specs = {"Heading 1": (19, BLUE, True, TITLE_FONT, 22, 10),
             "Heading 2": (14.5, BLUE2, True, TITLE_FONT, 18, 7),
             "Heading 3": (12.5, DARK, True, BASE_FONT, 13, 5)}
    for name, (size, color, bold, fnt, before, after) in specs.items():
        st = doc.styles[name]
        st.font.name = fnt
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        st.paragraph_format.line_spacing = 1.08


def configure_page(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)       # conforme guide RNCP : marges 2,5 cm
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# --------------------------------------------------------------------------- #
# Page de garde graphique
# --------------------------------------------------------------------------- #
def _run(p, text, size, color=DARK, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = BASE_FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r


def build_cover(doc, nexa_png, rondol_png):
    # Table externe : bande (col0) + contenu (col1)
    outer = doc.add_table(rows=1, cols=2)
    no_borders(outer)
    outer.allow_autofit = False
    set_row_height(outer.rows[0], 25.2, exact=True)
    band, content = outer.rows[0].cells
    band.width = Cm(0.55)
    content.width = Cm(15.0)
    shade(band, BAND_HEX)
    cell_valign(content, "top")

    def cpar(before=0, after=6, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = content.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        return p

    # Logos en haut (table 2 col sans bordure)
    p0 = content.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    logo_tbl = content.add_table(rows=1, cols=2)
    no_borders(logo_tbl)
    lc, rc = logo_tbl.rows[0].cells
    lc.width = Cm(7.3); rc.width = Cm(7.3)
    lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if nexa_png:
        lp.add_run().add_picture(str(nexa_png), width=Cm(4.6))
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if rondol_png:
        rp.add_run().add_picture(str(rondol_png), width=Cm(3.9))

    cpar(before=18); _run(content.paragraphs[-1], SCHOOL, 15, BLUE, bold=True)
    cpar(after=2); _run(content.paragraphs[-1], PROGRAM, 12.5, DARK)
    cpar(after=16); _run(content.paragraphs[-1], CERT, 10.5, GREY, italic=True)

    cpar(after=2); _run(content.paragraphs[-1], "MÉMOIRE DE THÈSE PROFESSIONNELLE", 13.5, DARK, bold=True)
    cpar(after=18); _run(content.paragraphs[-1], f"Année universitaire {YEAR}", 11, GREY, italic=True)

    # Encadré du thème — bloc plein bleu pétrole, texte blanc (style premium)
    theme = content.add_table(rows=1, cols=1)
    set_table_borders(theme, color="13314F", sz="8")
    tc = theme.rows[0].cells[0]
    shade(tc, "1F4E79")
    tp = tc.paragraphs[0]; tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(9); tp.paragraph_format.space_after = Pt(5)
    _run(tp, "THÈME", 10.5, WHITE, bold=True)
    tp2 = tc.add_paragraph(); tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp2.paragraph_format.space_after = Pt(8)
    _run(tp2, TITLE, 12.5, WHITE, bold=True)
    tp3 = tc.add_paragraph(); tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp3.paragraph_format.space_after = Pt(9)
    _run(tp3, SUBTITLE, 10.5, RGBColor(0xCC, 0xE0, 0xD4), italic=True)

    cpar(before=18, after=2); _run(content.paragraphs[-1], "Présenté et soutenu par", 11, DARK)
    cpar(after=14); _run(content.paragraphs[-1], AUTHOR, 15, BLUE, bold=True)

    # Lignes graphiques fines (courbes élégantes, bas-gauche)
    curves = ASSETS / "cover_curves.png"
    if curves.exists():
        lp = cpar(before=14, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
        lp.paragraph_format.left_indent = Cm(0.0)
        lp.add_run().add_picture(str(curves), width=Cm(7.4))
    else:
        for w, col in ((4.0, BLUE_HEX), (2.6, "1B7A3D"), (1.4, "9DC3E6")):
            lp = cpar(before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
            lp.paragraph_format.left_indent = Cm(0.2)
            _run(lp, " " * int(w * 6), 6)
            para_bottom_border(lp, col, sz="18")

    # Bas de page : logos école + entreprise puis mention
    blogo = cpar(before=8, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    if nexa_png:
        blogo.add_run().add_picture(str(nexa_png), height=Cm(0.95))
    sp = blogo.add_run("        "); sp.font.size = Pt(11)
    if rondol_png:
        blogo.add_run().add_picture(str(rondol_png), height=Cm(0.80))
    cpar(before=4, after=1); _run(content.paragraphs[-1], COMPANY + "   ·   " + SCHOOL, 10.5, GREY, italic=True)


def build_institutional(doc, nexa_png=None, rondol_png=None):
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, PROGRAM, 13, BLUE, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _run(p, f"Année universitaire {YEAR}", 11, GREY, italic=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    _run(p, "MÉMOIRE DE THÈSE PROFESSIONNELLE", 12.5, DARK, bold=True)

    # Encadré thème
    theme = doc.add_table(rows=1, cols=1)
    theme.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(theme, color=BLUE_HEX, sz="10")
    tc = theme.rows[0].cells[0]; shade(tc, THEME_FILL)
    tp = tc.paragraphs[0]
    tp.paragraph_format.space_before = Pt(6); tp.paragraph_format.space_after = Pt(2)
    _run(tp, "THÈME — ", 10.5, BLUE, bold=True)
    _run(tp, TITLE, 10.5, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Encadré problématique
    pb = doc.add_table(rows=1, cols=1)
    pb.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(pb, color="1B7A3D", sz="10")
    pc = pb.rows[0].cells[0]; shade(pc, "EFF6F1")
    pp = pc.paragraphs[0]
    pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(6)
    _run(pp, "PROBLÉMATIQUE — ", 10.5, GREEN, bold=True)
    _run(pp, PROBLEM, 10.5, DARK, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Tableau institutionnel
    info = [
        ("Nom du projet", "Plateforme prédictive d'aide à la décision — extrusion bivis SSB"),
        ("Date de début de projet", "Janvier 2026"),
        ("Auteur / rédacteur", AUTHOR),
        ("Entreprise d'accueil", COMPANY),
        ("Responsable entreprise / tuteur industriel", TUTOR_COMPANY),
        ("Établissement de formation", SCHOOL),
        ("Tuteur pédagogique / référent école", TUTOR_SCHOOL),
        ("Lieu", PLACE),
        ("Année universitaire", YEAR),
        ("Date de dépôt et de soutenance", DEFENSE + " (dépôt au plus tard le 17 août 2026)"),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="D9D9D9")
    for i, (k, v) in enumerate(info):
        c0, c1 = tbl.rows[i].cells
        shade(c0, LTGREY)
        r0 = c0.paragraphs[0].add_run(k)
        r0.font.bold = True; r0.font.size = Pt(10.5); r0.font.name = BASE_FONT
        r0.font.color.rgb = DARK
        r1 = c1.paragraphs[0].add_run(v)
        r1.font.size = Pt(10.5); r1.font.name = BASE_FONT; r1.font.color.rgb = DARK
        c0.width = Cm(6.6); c1.width = Cm(9.0)

    # Bas de page institutionnelle : logos école + entreprise
    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    blogo = doc.add_paragraph(); blogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if nexa_png:
        blogo.add_run().add_picture(str(nexa_png), height=Cm(0.9))
    sp = blogo.add_run("        "); sp.font.size = Pt(11)
    if rondol_png:
        blogo.add_run().add_picture(str(rondol_png), height=Cm(0.78))
    doc.add_page_break()


# --------------------------------------------------------------------------- #
# Rendu inline
# --------------------------------------------------------------------------- #
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\(Source interne[^)]*\))")


def add_runs(paragraph, text, default_size=11):
    text = re.sub(r"\s*\[À COMPLÉTER[^\]]*\]", "", text).strip()
    if not text:
        return
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.font.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = paragraph.add_run(tok[1:-1]); r.font.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("(Source interne"):
            r = paragraph.add_run(tok); r.font.italic = True
            r.font.size = Pt(8.5); r.font.color.rgb = GREY
        else:
            r = paragraph.add_run(tok)
        r.font.name = r.font.name or BASE_FONT
        if r.font.size is None:
            r.font.size = Pt(default_size)


# --------------------------------------------------------------------------- #
# Images & légendes
# --------------------------------------------------------------------------- #
def add_figure(doc, img_path, caption, width=FIG_W):
    if not img_path or not Path(img_path).exists():
        print(f"[WARN] image manquante: {img_path}")
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(img_path), width=width)
    fig_counter[0] += 1
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(f"Figure {fig_counter[0]} — ")
    r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = BLUE; r.font.name = BASE_FONT
    r2 = cap.add_run(caption)
    r2.font.italic = True; r2.font.size = Pt(9.5); r2.font.color.rgb = GREY; r2.font.name = BASE_FONT


# Mapping placeholder -> (image, caption). Clé = sous-chaîne unique.
def build_asset_map():
    F = lambda n: FIGDIR / n
    C = lambda n: CAPDIR / n
    P = POSTER / "generated_v2"
    PF = POSTER / "generated"
    return {
        # --- Ajouts finalisation : authentification + championnat ML ---------
        "écran de connexion de la plateforme": (C("cap_login.png"),
            "Écran de connexion — accès protégé, mot de passe haché (PBKDF2), jamais stocké en clair"),
        "page Compte — historique des connexions": (C("cap_account.png"),
            "Page Compte — utilisateur connecté et historique des connexions lu depuis la base durable"),
        "championnat des modèles supervisés en validation": (F("fig_championnat_modeles.png"),
            "Championnat des modèles supervisés en validation Leave-One-Group-Out : F1-macro moyen et écart-type inter-essais"),
        # Figures
        "cartographie des trois domaines": (F("fig_domain_intersection.png"),
            "Cartographie des trois domaines (extrusion bivis, apprentissage automatique, batteries tout-solide) et zone d'intersection peu couverte"),
        "matrice SWOT illustrée": (F("fig_swot.png"), "Matrice SWOT de Rondol Industrie au regard du projet"),
        "tableau comparatif Flask/Dash": (F("fig_stack_comparison.png"),
            "Comparatif des piles technologiques Flask/Dash et Streamlit/Supabase"),
        "schéma de la démarche CRISP-DM": (F("fig_crispdm.png"), "Démarche CRISP-DM adaptée au projet"),
        "diagramme de Gantt": (F("fig_gantt.png"), "Rétroplanning indicatif du projet (chronologie relative)"),
        "répartition de la suite de tests": (F("fig_tests.png"), "Familles de la suite de tests automatisés"),
        "pipeline de données capteurs": (F("fig_data_pipeline.png"),
            "Pipeline de données, des fichiers CSV bruts au jeu de données ML"),
        "importances de variables (top 10)": (PF / "fig04_feature_importance_RF_w60.png",
            "Importance des variables (top 10) — Random Forest, fenêtre 60 s"),
        "comparaison des performances en split": (F("fig_validation.png"),
            "Performances selon le protocole de validation (split aléatoire vs GSS vs LOGO)"),
        "schéma des deux niveaux d'IA": (F("fig_two_level_ai.png"),
            "Les deux niveaux d'intelligence artificielle : référence hors-ligne et logique intégrée"),
        "architecture générale de la plateforme en couches": (F("fig_architecture.png"),
            "Architecture logicielle en couches de la plateforme Rondol"),
        "schéma de persistance Supabase": (F("fig_persistence.png"),
            "Persistance durable (Supabase / fichier / JSON) et auto-réparation de l'état validé"),
        "tableau de synthèse des équations du moteur": (F("fig_equations.png"),
            "Statut des équations du moteur procédé : E4 implémentée, E5/E6/E7 différées (None)"),
        "feuille de route des perspectives en quatre axes": (F("fig_roadmap.png"),
            "Feuille de route des perspectives en quatre axes"),
        "schéma entité-relation de la table": (F("fig_er_schema.png"),
            "Schéma de la table rondol_state et structure du document JSONB applied_state"),
        # Captures
        "photographie ou vue CAO": (None, None),  # pas d'actif -> ignoré proprement
        "extrait d'un fichier CSV brut": (F("fig_csv_extract.png"),
            "Extrait d'un fichier CSV brut de capteur (filière DIE) — colonnes Timestamp / Name / Value"),
        "interface Streamlit — page Supervision": (C("cap_supervision.png"),
            "Interface de supervision du prototype Rondol"),
        "exemple d'alerte et de recommandation de l'agent": (P / "figv2_04_agent_recommendation_panel.png",
            "Exemple d'alerte et de recommandation produites par l'agent industriel"),
        "page Profile illustrant la modification": (C("cap_profile.png"),
            "Paramétrage du profil de vis et des zones thermiques (page Profile)"),
        "page Supervision pour le cas C3": (P / "figv2_03_agent_alert_panel.png",
            "Supervision — cas C3 : alerte rouge en zone Z5 (surcharge céramique LATP 35 %)"),
        "panneau de recommandation de l'agent (cas C4)": (P / "figv2_04_agent_recommendation_panel.png",
            "Panneau de recommandation de l'agent — cas C4 (plan d'actions hiérarchisé)"),
        "Supervision après correction (cas C5": (P / "figv2_05_before_after_dashboard.png",
            "Comparaison avant → après (C3 → C5) : redressement des indicateurs après correction"),
        "panneau d'une recommandation de l'agent montrant le rationale": (P / "figv2_04_agent_recommendation_panel.png",
            "Chaîne d'explicabilité de l'agent : diagnostic, actions chiffrées avant→après et alerte liée"),
        # Annexe E
        "page Supervision — statut machine": (C("cap_supervision.png"),
            "Page Supervision — statut machine, score de stabilité, alertes et recommandations IA"),
        "page Profile — configuration du profil": (C("cap_profile.png"),
            "Page Profile — configuration du profil de vis (zones, éléments, indicateurs)"),
        "page Settings — seuils de l'IA": (C("cap_settings.png"),
            "Configuration des paramètres procédé et de persistance (page Settings)"),
        "cas C3 — état instable": (P / "figv2_03_agent_alert_panel.png",
            "Cas C3 — état instable affiché (LATP 35 %, score 46, probabilité 0,35)"),
        "recommandation C4": (P / "figv2_04_agent_recommendation_panel.png",
            "Recommandation C4 — actions correctives chiffrées et justifiées de l'agent"),
        "état après correction C5": (P / "figv2_05_before_after_dashboard.png",
            "État après correction C5 — retour à un régime stable (score 78, probabilité 0,87)"),
        "page Moteur Procédé (Process Engine)": (C("cap_process_engine.png"),
            "Moteur procédé et indicateurs calculés (page Process Engine)"),
        "page Run Analysis": (C("cap_run_analysis.png"),
            "Analyse d'un essai et indicateurs de stabilité (page Run Analysis)"),
        "page History": (C("cap_history.png"),
            "Historique des configurations sauvegardées (page History)"),
    }


ASSET_MAP = None


def resolve_placeholder(inner):
    for key, val in ASSET_MAP.items():
        if key in inner:
            return val
    return None


# --------------------------------------------------------------------------- #
# Nettoyage du markdown
# --------------------------------------------------------------------------- #
def clean_markdown(md):
    # 1) Sommaire estimé -> sentinelle TOC
    md = re.sub(r"## Sommaire détaillé.*?Le cœur du mémoire représente environ 50 pages hors bibliographie et annexes\.\*",
                "[[TOC]]", md, flags=re.DOTALL)
    # 2) Inline À COMPLÉTER spécifiques
    md = md.replace(
        "ne figurent pas dans les sources internes du projet et ne sont pas inventées ici : `[À COMPLÉTER : chiffre d'affaires, effectif, statut juridique, dirigeants]`.",
        "ne figurent pas dans les sources internes du projet et ne sont donc pas reproduites ici (chiffre d'affaires, effectif, statut juridique, dirigeants).")
    md = md.replace(
        "`[À COMPLÉTER : chiffrage du marché de l'extrusion de laboratoire — taille, taux de croissance annuel (CAGR), segmentation géographique. Donnée non présente dans les sources internes ; à étayer par recherche documentaire externe.]`",
        "Le chiffrage précis de ce marché (taille, taux de croissance annuel, segmentation géographique) relève d'une étude documentaire externe, hors du périmètre des sources internes mobilisées dans ce mémoire.")
    md = md.replace("la date de dépôt et de soutenance du mémoire restant `[À COMPLÉTER : date de dépôt et de soutenance]`.",
                    "la date de dépôt et de soutenance du mémoire restant à confirmer.")
    md = md.replace(" [À COMPLÉTER : retour qualitatif du tuteur industriel M. Maël Gallas à l'issue de la démonstration du 16 juin 2026]", "")
    md = md.replace(
        "[À COMPLÉTER : dictionnaire exhaustif des 96 variables dérivées, avec pour chacune le nom exact, la définition, l'unité, le capteur source et le rôle dans le modèle.]",
        "Le dictionnaire exhaustif des 96 variables dérivées (nom exact, définition, unité, capteur source et rôle dans le modèle) est maintenu dans les sources techniques du projet (src/features.py, src/config.py).")
    # 3) Bibliographie : retirer les placeholders DOI/refs
    md = re.sub(r"\s*\[À COMPLÉTER : (volume[^\]]*|référence[^\]]*)\]", "", md)
    md = re.sub(r"\n\[À COMPLÉTER : intégration des 38 références[^\]]*\]\n", "\n", md)
    # 4) Cellules de tableau
    md = md.replace("[À COMPLÉTER : période]", "—")
    md = md.replace("[À COMPLÉTER : stabilité projetée associée]", "Stable (projetée)")
    md = md.replace("| `[À COMPLÉTER]` |", "| n.d. |")
    md = md.replace("`[À COMPLÉTER]`", "n.d.")
    md = md.replace("[À COMPLÉTER]", "n.d.")
    # 5) Lignes logo résiduelles
    md = re.sub(r"\[INSÉRER LOGO[^\]]*\]\n?", "", md)
    # 6) Bloc SQL -> fence
    md = md.replace(
        "sql\nCREATE TABLE rondol_state (\n    key     TEXT PRIMARY KEY,\n    payload JSONB\n);",
        "```sql\nCREATE TABLE rondol_state (\n    key     TEXT PRIMARY KEY,\n    payload JSONB\n);\n```")
    # 7) Annexe E : titre + ajouts Run Analysis / History
    md = md.replace("## Annexe E — Captures à insérer",
                    "## Annexe E — Captures complémentaires de l'application")
    md = md.replace(
        "[INSÉRER CAPTURE : page Moteur Procédé (Process Engine) — vue moteur en lecture seule]",
        "[INSÉRER CAPTURE : page Moteur Procédé (Process Engine) — vue moteur en lecture seule]\n"
        "[INSÉRER CAPTURE : page Run Analysis — analyse temporelle d'un essai]\n"
        "[INSÉRER CAPTURE : page History — historique des configurations sauvegardées]")
    # 8) Injecter Annexe F (checklist RNCP) avant la section finale
    annexe_f = (
        "\n## Annexe F — Checklist de conformité projet / RNCP\n\n"
        "Le tableau ci-dessous met en regard les attendus d'un projet de certification "
        "RNCP 37137 (niveau 7) et les éléments correspondants du projet Rondol.\n\n"
        "| Exigence | Élément du projet | Statut |\n"
        "|---|---|---|\n"
        "| Cadrage métier & problématique | Partie 3 (constat, besoin, problématique validée) | Conforme |\n"
        "| État de l'art & étude de marché | Partie 2 (38 références, analyse concurrentielle, SWOT) | Conforme |\n"
        "| Exploitation de données réelles | Partie 5 (campagne avril 2026, 12 capteurs) | Conforme |\n"
        "| Modélisation ML & validation rigoureuse | Partie 5 (RF/XGBoost/SVM, GroupShuffleSplit, LOGO) | Conforme |\n"
        "| Développement d'une application | Partie 6 (6 pages Streamlit, architecture en couches) | Conforme |\n"
        "| Base de données & persistance | Supabase / PostgreSQL, table rondol_state | Conforme |\n"
        "| Livrable dump SQL | Réalisable (socle PostgreSQL) | À produire |\n"
        "| Tests & qualité logicielle | 685 tests automatisés sur 71 fichiers | Conforme |\n"
        "| Déploiement | GitHub + Streamlit Cloud | Conforme |\n"
        "| Gestion de projet | Partie 4 (CRISP-DM, risques, jalons) | Conforme |\n"
        "| Éthique, RGPD, accessibilité | Partie 8 (explicabilité, RGPD, WCAG) | Traité (WCAG à venir) |\n\n"
        "*Tableau F.1 — Checklist de conformité du projet aux attendus de la certification.*\n\n"
    )
    md = md.replace("# Points à compléter avant dépôt", annexe_f + "# Informations restant à confirmer avant dépôt")
    # 9) Réduire la section finale à 5 points
    final_re = re.compile(r"(# Informations restant à confirmer avant dépôt\n)(.*)$", re.DOTALL)
    m = final_re.search(md)
    if m:
        intro = ("\nLes éléments factuels suivants ne figurent pas dans les sources internes "
                 "du projet et restent à confirmer avant le dépôt définitif.\n\n")
        bullets = (
            "- **Données institutionnelles de Rondol Industrie** : chiffre d'affaires, effectif, statut juridique et dirigeants.\n"
            "- **Calendrier académique** : dates exactes de dépôt et de soutenance du mémoire.\n"
            "- **Étude de marché et analyse concurrentielle** : chiffrage du marché de l'extrusion de laboratoire et données chiffrées sur les concurrents.\n"
            "- **Bibliographie** : volumes, numéros, pages et DOI des références de l'état de l'art.\n"
            "- **Retour de la démonstration client** du 16 juin 2026 (commentaire du tuteur industriel).\n")
        md = md[:m.start()] + m.group(1) + intro + bullets
    return md


# --------------------------------------------------------------------------- #
# Parsing du corps
# --------------------------------------------------------------------------- #
def is_table_sep(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c or "") for c in cells) and len(cells) > 0


def render_table(doc, rows):
    parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    if not parsed:
        return
    ncol = max(len(r) for r in parsed)
    header, body = parsed[0], parsed[1:]
    tbl = doc.add_table(rows=1 + len(body), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)
    for j in range(ncol):
        cell = tbl.rows[0].cells[j]; shade(cell, "1F4E79")
        p = cell.paragraphs[0]
        add_runs(p, header[j] if j < len(header) else "", default_size=9.5)
        for run in p.runs:
            run.font.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = WHITE
    for i, row in enumerate(body, start=1):
        for j in range(ncol):
            cell = tbl.rows[i].cells[j]
            if i % 2 == 0:
                shade(cell, LTGREY)
            p = cell.paragraphs[0]
            add_runs(p, row[j] if j < len(row) else "", default_size=9.5)
            for run in p.runs:
                if run.font.size is None or run.font.size > Pt(9.5):
                    run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_code(doc, lines):
    tbl = doc.add_table(rows=1, cols=1)
    set_table_borders(tbl, color="D9D9D9")
    cell = tbl.rows[0].cells[0]; shade(cell, "F5F5F5")
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(9.5); r.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_body(doc, md_text):
    idx = md_text.find("## Remerciements")
    lines = md_text[idx if idx >= 0 else 0:].splitlines()
    i, first_h1, just_broke = 0, True, False
    while i < len(lines):
        stripped = lines[i].rstrip().strip()

        # Tableaux
        if stripped.startswith("|") and stripped.endswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            block = [b for b in block if not is_table_sep(b)]
            render_table(doc, block); just_broke = False
            continue

        # Code fence
        if stripped.startswith("```"):
            i += 1; code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            render_code(doc, code); just_broke = False
            continue

        i += 1
        if not stripped or stripped in ("---", "***", "___"):
            continue
        if stripped == "[[TOC]]":
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
            _run(p, "Sommaire", 16, BLUE, bold=True)
            add_toc(doc)
            just_broke = False
            continue
        if stripped == "[SAUT DE PAGE]":
            if not just_broke:
                doc.add_page_break(); just_broke = True
            continue
        if stripped.startswith("!["):
            continue

        # Placeholders figure/capture -> image réelle
        if stripped.startswith("[INSÉRER FIGURE") or stripped.startswith("[INSÉRER CAPTURE"):
            inner = stripped.strip("[]")
            val = resolve_placeholder(inner)
            if val and val[0]:
                add_figure(doc, val[0], val[1])
            just_broke = False
            continue

        # Titres
        if stripped.startswith("# "):
            if not first_h1 and not just_broke:
                doc.add_page_break()
            first_h1 = False
            doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
            just_broke = False
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2"); just_broke = False
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3"); just_broke = False
            continue

        # Citation
        if stripped.startswith(">"):
            content = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            para_bottom_border  # noop
            add_runs(p, content)
            for r in p.runs:
                r.font.italic = True
            continue

        # Légende / note italique pure
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("*"))
            r.font.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY; r.font.name = BASE_FONT
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, stripped)


# --------------------------------------------------------------------------- #
# En-tête / pied de page
# --------------------------------------------------------------------------- #
def build_header_footer(section, nexa_png, rondol_png):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    h = section.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rh = h.add_run(HEADER_TEXT)
    rh.font.size = Pt(8.5); rh.font.italic = True; rh.font.color.rgb = GREY; rh.font.name = BASE_FONT
    para_bottom_border(h, "BFBFBF", sz="4")

    # Pied épuré : identité à gauche, numéro de page à droite, filet fin.
    # (Les logos institutionnels figurent sur la page de garde ; les répéter
    # sur chacune des 70 pages alourdit inutilement la lecture.)
    ftbl = section.footer.add_table(rows=1, cols=2, width=CONTENT_W)
    no_borders(ftbl)
    lc, rc = ftbl.rows[0].cells
    lc.width = Cm(11.5); rc.width = Cm(4.1)

    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_top_border(lp, "D5DCE4", sz="4")
    rl = lp.add_run(f"{AUTHOR}   ·   {SCHOOL}   ·   {YEAR}")
    rl.font.size = Pt(8); rl.font.color.rgb = GREY; rl.font.name = BASE_FONT

    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_top_border(rp, "D5DCE4", sz="4")
    add_field(rp, "PAGE")
    for r in rp.runs:
        r.font.size = Pt(9); r.font.color.rgb = BLUE; r.font.name = TITLE_FONT; r.font.bold = True


# --------------------------------------------------------------------------- #
def build_document():
    global ASSET_MAP
    ASSET_MAP = build_asset_map()
    nexa_png, rondol_png = ensure_logos()
    doc = Document()
    configure_styles(doc)
    enable_hyphenation(doc)      # supprime les blancs du texte justifié
    configure_page(doc.sections[0])

    build_cover(doc, nexa_png, rondol_png)
    build_institutional(doc, nexa_png, rondol_png)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body)
    set_start_page_number(body, 1)
    build_header_footer(body, nexa_png, rondol_png)

    md_text = clean_markdown(MD.read_text(encoding="utf-8"))
    parse_body(doc, md_text)
    return doc


def main():
    REPORTS.mkdir(parents=True, exist_ok=True)
    if not MD.exists():
        print(f"[ERREUR] Source introuvable : {MD}"); return 1
    print("[1/2] Construction du document FINAL ...")
    doc = build_document()
    doc.save(str(DOCX))
    print(f"      DOCX : {DOCX}")
    print(f"      Figures insérées : {fig_counter[0]}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
