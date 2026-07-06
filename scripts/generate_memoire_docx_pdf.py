# -*- coding: utf-8 -*-
"""
Génère le mémoire de thèse professionnelle Rondol en Word (.docx) puis en PDF.

- Contenu : parsé depuis docs/memoire_these_professionnelle_rondol.md (à partir de
  la section « Remerciements » ; la page de garde est construite ici, par code,
  pour garantir l'insertion propre des logos réels).
- Logos : nexa LOGO.webp (racine) converti en PNG, assets/rondol_logo.png.
- Sortie : reports/Memoire_Rondol_Wilfried_Galtier_MBEUMI.docx + .pdf
- PDF : tente LibreOffice/soffice en headless, sinon docx2pdf (Word COM).

Usage : python scripts/generate_memoire_docx_pdf.py
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Chemins et constantes officielles
# --------------------------------------------------------------------------- #
ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "memoire_these_professionnelle_rondol.md"
REPORTS = ROOT / "reports"
ASSETS = REPORTS / "assets"
DOCX = REPORTS / "Memoire_Rondol_Wilfried_Galtier_MBEUMI.docx"
PDF = REPORTS / "Memoire_Rondol_Wilfried_Galtier_MBEUMI.pdf"

AUTHOR = "Wilfried Galtier MBEUMI"
SCHOOL = "Nexa Digital School"
PROGRAM = "Mastère Data & Intelligence Artificielle"
CERT = "Certification professionnelle RNCP 37137 — Niveau 7 (Bac+5)"
COMPANY = "Rondol Industrie"
TUTOR_COMPANY = "M. Maël Gallas"
TUTOR_SCHOOL = "M. Moussa NDIAYE"
PLACE = "Nancy — Institut Jean Lamour (IJL), Campus ARTEM"
YEAR = "2025 – 2026"
DEFENSE = "[À COMPLÉTER : date de dépôt et de soutenance]"

TITLE = (
    "Conception et déploiement d'un système d'intelligence artificielle "
    "prédictif d'aide à la décision pour l'optimisation des paramètres "
    "d'extrusion bivis de composants de batteries tout-solide (dry / semi-dry)"
)
SUBTITLE = (
    "Un prototype professionnel de jumeau numérique appliqué à l'extrudeuse "
    "bivis 10,5 mm de Rondol Industrie"
)

HEADER_TEXT = "Mémoire de thèse professionnelle — Rondol Industrie"
FOOTER_TEXT = "Wilfried Galtier MBEUMI — Nexa Digital School — 2025/2026"

# Couleurs Rondol
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
BLUE = RGBColor(0x00, 0x5B, 0x96)
DARK = RGBColor(0x33, 0x33, 0x33)
GREY = RGBColor(0x66, 0x66, 0x66)

BASE_FONT = "Calibri"

# --------------------------------------------------------------------------- #
# Helpers OXML (bordures, fonds, champs)
# --------------------------------------------------------------------------- #
def shade(cell_or_pr, hex_fill: str) -> None:
    """Applique un fond de couleur à une cellule de tableau."""
    tcPr = cell_or_pr._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_table_borders(table, color="BFBFBF", sz="4") -> None:
    """Bordures fines uniformes sur tout le tableau."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def add_field(paragraph, field: str) -> None:
    """Insère un champ Word (ex. PAGE) dans un paragraphe."""
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1)
    run._r.append(instr)
    run._r.append(f2)


def set_start_page_number(section, start: int = 1) -> None:
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


# --------------------------------------------------------------------------- #
# Logos
# --------------------------------------------------------------------------- #
def ensure_logos():
    ASSETS.mkdir(parents=True, exist_ok=True)
    nexa_png = ASSETS / "nexa_logo.png"
    rondol_png = ASSETS / "rondol_logo.png"

    nexa_src = ROOT / "nexa LOGO.webp"
    if nexa_src.exists():
        try:
            from PIL import Image

            im = Image.open(nexa_src).convert("RGBA")
            bg = Image.new("RGBA", im.size, (255, 255, 255, 255))
            bg.alpha_composite(im)
            bg.convert("RGB").save(nexa_png, "PNG")
        except Exception as exc:  # noqa: BLE001
            print(f"[WARN] Conversion logo Nexa échouée : {exc}")
            nexa_png = None
    else:
        print("[WARN] nexa LOGO.webp introuvable à la racine.")
        nexa_png = None

    rondol_src = ROOT / "assets" / "rondol_logo.png"
    if rondol_src.exists():
        shutil.copy(rondol_src, rondol_png)
    else:
        print("[WARN] assets/rondol_logo.png introuvable.")
        rondol_png = None

    return (nexa_png if nexa_png and nexa_png.exists() else None,
            rondol_png if rondol_png and rondol_png.exists() else None)


# --------------------------------------------------------------------------- #
# Styles du document
# --------------------------------------------------------------------------- #
def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    specs = {
        "Heading 1": (16, GREEN, True),
        "Heading 2": (13, BLUE, True),
        "Heading 3": (11.5, DARK, True),
    }
    for name, (size, color, bold) in specs.items():
        st = doc.styles[name]
        st.font.name = BASE_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


def configure_page(section) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


# --------------------------------------------------------------------------- #
# Page de garde
# --------------------------------------------------------------------------- #
def build_cover(doc: Document, nexa_png, rondol_png) -> None:
    # Logos alignés en haut : tableau 2 colonnes sans bordure
    logo_tbl = doc.add_table(rows=1, cols=2)
    logo_tbl.autofit = True
    left, right = logo_tbl.rows[0].cells
    pL = left.paragraphs[0]
    pL.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if nexa_png:
        pL.add_run().add_picture(str(nexa_png), width=Cm(5.0))
    else:
        pL.add_run("[LOGO NEXA]")
    pR = right.paragraphs[0]
    pR.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if rondol_png:
        pR.add_run().add_picture(str(rondol_png), width=Cm(4.2))
    else:
        pR.add_run("[LOGO RONDOL]")

    def centered(text, size, color=DARK, bold=False, italic=False, before=0, after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.font.name = BASE_FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
        return p

    centered(SCHOOL, 14, BLUE, bold=True, before=24)
    centered(f"{PROGRAM}", 12, DARK)
    centered(CERT, 10.5, GREY, italic=True, after=18)
    centered("MÉMOIRE DE THÈSE PROFESSIONNELLE", 13, DARK, bold=True)
    centered(f"Année universitaire {YEAR}", 11, GREY, italic=True, after=24)

    centered(TITLE, 16, GREEN, bold=True, before=6, after=10)
    centered(SUBTITLE, 12, BLUE, italic=True, after=28)

    centered("Présenté et soutenu par", 11, DARK, after=2)
    centered(AUTHOR, 15, DARK, bold=True, after=24)

    # Bloc informations
    info = [
        ("Entreprise d'accueil", COMPANY),
        ("Responsable entreprise / tuteur industriel", TUTOR_COMPANY),
        ("Établissement de formation", SCHOOL),
        ("Tuteur pédagogique / référent école", TUTOR_SCHOOL),
        ("Lieu", PLACE),
        ("Année universitaire", YEAR),
        ("Date de dépôt et de soutenance", DEFENSE),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="D9D9D9")
    for i, (k, v) in enumerate(info):
        c0, c1 = tbl.rows[i].cells
        shade(c0, "F2F2F2")
        r0 = c0.paragraphs[0].add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(10.5)
        r0.font.name = BASE_FONT
        add_runs(c1.paragraphs[0], v, default_size=10.5)
    # largeurs
    for row in tbl.rows:
        row.cells[0].width = Cm(7.0)
        row.cells[1].width = Cm(9.5)


# --------------------------------------------------------------------------- #
# Rendu inline (gras, italique, code, sources internes, à compléter)
# --------------------------------------------------------------------------- #
INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[À COMPLÉTER[^\]]*\]|\(Source interne[^)]*\))"
)


def add_runs(paragraph, text: str, default_size: float = 11) -> None:
    text = text.strip()
    if not text:
        return
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.font.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.font.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
        elif tok.startswith("[À COMPLÉTER"):
            r = paragraph.add_run(tok)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x9C, 0x6500 // 256 if False else 0x65, 0x00)
            _highlight(r, "yellow")
        elif tok.startswith("(Source interne"):
            r = paragraph.add_run(tok)
            r.font.italic = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = GREY
        else:
            r = paragraph.add_run(tok)
        r.font.name = BASE_FONT
        if r.font.size is None:
            r.font.size = Pt(default_size)


def _highlight(run, color_name: str) -> None:
    rPr = run._r.get_or_add_rPr()
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), color_name)
    rPr.append(hl)


# --------------------------------------------------------------------------- #
# Encadrés (figures / à compléter)
# --------------------------------------------------------------------------- #
def add_callout(doc: Document, text: str, fill: str, label: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    set_table_borders(tbl, color="BFBFBF")
    cell = tbl.rows[0].cells[0]
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"{label} ")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.name = BASE_FONT
    r2 = p.add_run(text)
    r2.font.italic = True
    r2.font.size = Pt(10)
    r2.font.name = BASE_FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# --------------------------------------------------------------------------- #
# Parsing du Markdown (corps)
# --------------------------------------------------------------------------- #
def is_table_sep(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c or "") for c in cells) and len(cells) > 0


def render_table(doc: Document, rows: list[str]) -> None:
    parsed = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return
    ncol = max(len(r) for r in parsed)
    header = parsed[0]
    body = parsed[1:]
    tbl = doc.add_table(rows=1 + len(body), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)
    # en-tête
    for j in range(ncol):
        cell = tbl.rows[0].cells[j]
        shade(cell, "E7ECEF")
        txt = header[j] if j < len(header) else ""
        p = cell.paragraphs[0]
        add_runs(p, txt, default_size=10)
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
    # corps
    for i, row in enumerate(body, start=1):
        for j in range(ncol):
            cell = tbl.rows[i].cells[j]
            txt = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            add_runs(p, txt, default_size=10)
            for run in p.runs:
                if run.font.size is None or run.font.size > Pt(10):
                    run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_body(doc: Document, md_text: str) -> None:
    # On démarre au premier "## Remerciements"
    idx = md_text.find("## Remerciements")
    if idx == -1:
        idx = 0
    lines = md_text[idx:].splitlines()

    i = 0
    first_h1 = True
    just_broke = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # Tableaux
        if stripped.startswith("|") and stripped.endswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            # retire la ligne séparatrice
            block = [b for b in block if not is_table_sep(b)]
            render_table(doc, block)
            just_broke = False
            continue

        i += 1

        if not stripped:
            continue
        if stripped in ("---", "***", "___"):
            continue
        if stripped == "[SAUT DE PAGE]":
            if not just_broke:
                doc.add_page_break()
                just_broke = True
            continue

        # Images markdown (ignorées dans le corps : logos gérés en page de garde)
        if stripped.startswith("!["):
            continue
        # Lignes logo résiduelles -> encadré
        if stripped.startswith("[INSÉRER LOGO"):
            inner = stripped.strip("[]")
            add_callout(doc, inner.replace("INSÉRER LOGO", "").strip(" :—-"),
                        "F2F2F2", "LOGO À INSÉRER :")
            continue

        # Titres
        if stripped.startswith("# "):
            if not first_h1 and not just_broke:
                doc.add_page_break()
            first_h1 = False
            doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
            just_broke = False
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
            just_broke = False
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
            just_broke = False
            continue

        # Citation (problématique, notes)
        if stripped.startswith(">"):
            content = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_runs(p, content)
            for r in p.runs:
                r.font.italic = True
            continue

        # Encadrés figure / capture
        if stripped.startswith("[INSÉRER FIGURE"):
            inner = stripped[len("[INSÉRER FIGURE"):].strip("[] :—-")
            add_callout(doc, inner, "EAF2FB", "FIGURE À INSÉRER :")
            continue
        if stripped.startswith("[INSÉRER CAPTURE"):
            inner = stripped[len("[INSÉRER CAPTURE"):].strip("[] :—-")
            add_callout(doc, inner, "EAF2FB", "CAPTURE À INSÉRER :")
            continue

        # Ligne entièrement "[À COMPLÉTER ...]" (éventuellement entre backticks)
        bare = stripped.strip("`").strip()
        if bare.startswith("[À COMPLÉTER") and bare.endswith("]"):
            add_callout(doc, bare.strip("[]"), "FFF6CC", "À COMPLÉTER :")
            continue

        # Légende de tableau / note en italique pur : *...*
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(stripped.strip("*"))
            r.font.italic = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = GREY
            r.font.name = BASE_FONT
            continue

        # Paragraphe courant (peut contenir gras/italique/sources/à compléter)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, stripped)


# --------------------------------------------------------------------------- #
# En-tête / pied de page (section corps)
# --------------------------------------------------------------------------- #
def build_header_footer(body_section) -> None:
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False

    h = body_section.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = h.add_run(HEADER_TEXT)
    rh.font.size = Pt(9)
    rh.font.italic = True
    rh.font.color.rgb = GREY
    rh.font.name = BASE_FONT

    f = body_section.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = f.add_run(FOOTER_TEXT + "   —   Page ")
    rf.font.size = Pt(9)
    rf.font.color.rgb = GREY
    rf.font.name = BASE_FONT
    add_field(f, "PAGE")
    for r in f.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        r.font.name = BASE_FONT


# --------------------------------------------------------------------------- #
# Construction du document
# --------------------------------------------------------------------------- #
def build_document() -> Document:
    nexa_png, rondol_png = ensure_logos()
    doc = Document()
    configure_styles(doc)
    configure_page(doc.sections[0])

    # Section 1 : page de garde (sans en-tête / pied / numéro)
    build_cover(doc, nexa_png, rondol_png)

    # Saut de section -> corps
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section)
    set_start_page_number(body_section, 1)
    build_header_footer(body_section)

    md_text = MD.read_text(encoding="utf-8")
    parse_body(doc, md_text)
    return doc


# --------------------------------------------------------------------------- #
# Conversion PDF
# --------------------------------------------------------------------------- #
def convert_to_pdf() -> str:
    # 1) LibreOffice / soffice headless
    for exe in ("libreoffice", "soffice"):
        path = shutil.which(exe)
        if path:
            try:
                subprocess.run(
                    [path, "--headless", "--convert-to", "pdf",
                     "--outdir", str(REPORTS), str(DOCX)],
                    check=True, timeout=300,
                )
                if PDF.exists():
                    return f"{exe} (headless)"
            except Exception as exc:  # noqa: BLE001
                print(f"[WARN] {exe} a échoué : {exc}")
    # 2) docx2pdf (Word COM)
    try:
        from docx2pdf import convert

        convert(str(DOCX), str(PDF))
        if PDF.exists():
            return "docx2pdf (Microsoft Word)"
    except Exception as exc:  # noqa: BLE001
        print(f"[ERREUR] docx2pdf a échoué : {exc}")
    return ""


# --------------------------------------------------------------------------- #
# Contrôle qualité
# --------------------------------------------------------------------------- #
def quality_check() -> None:
    doc = Document(str(DOCX))
    full = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                full += "\n" + c.text

    checks = {
        "Auteur correct": AUTHOR in full,
        "Établissement Nexa": SCHOOL in full,
        "Référent M. Moussa NDIAYE": "Moussa NDIAYE" in full,
        "Tuteur M. Maël Gallas": "Maël Gallas" in full,
        "Rondol Industrie": "Rondol Industrie" in full,
        "Remerciements": "Remerciements" in full,
        "Résumé": "Résumé" in full or "Resume" in full,
        "Abstract": "Abstract" in full,
        "Liste des sigles": "sigles" in full.lower(),
        "Sommaire": "Sommaire" in full,
        "Introduction générale": "Introduction générale" in full,
        "Partie 1": "Partie 1" in full,
        "Partie 8": "Partie 8" in full,
        "Conclusion générale": "Conclusion générale" in full,
        "Bibliographie": "Bibliographie" in full,
        "Annexes": "Annexe" in full,
        "Pas de HTML brut": not re.search(r"<(div|br|table|tr|td)\b", full),
        "Pas de '##' résiduel": "##" not in full,
        "Pas de '**' résiduel": "**" not in full,
    }
    # images (logos)
    n_images = len(doc.inline_shapes)

    print("\n=== CONTRÔLE QUALITÉ ===")
    for k, ok in checks.items():
        print(f"  [{'OK ' if ok else 'KO '}] {k}")
    print(f"  [{'OK ' if n_images >= 2 else 'KO '}] Logos en page de garde (images détectées : {n_images})")
    print(f"  Paragraphes : {len(doc.paragraphs)} | Tableaux : {len(doc.tables)}")
    return all(checks.values()) and n_images >= 2


def estimate_pages() -> str:
    if PDF.exists():
        try:
            from pypdf import PdfReader

            return str(len(PdfReader(str(PDF)).pages))
        except Exception:
            try:
                from PyPDF2 import PdfReader

                return str(len(PdfReader(str(PDF)).pages))
            except Exception:
                pass
    return "~ estimation indisponible"


# --------------------------------------------------------------------------- #
def main() -> int:
    REPORTS.mkdir(parents=True, exist_ok=True)
    if not MD.exists():
        print(f"[ERREUR] Source introuvable : {MD}")
        return 1
    print(f"[1/4] Construction du document depuis {MD.name} ...")
    doc = build_document()
    doc.save(str(DOCX))
    print(f"[2/4] DOCX généré : {DOCX}")
    print("[3/4] Conversion PDF ...")
    engine = convert_to_pdf()
    if engine:
        print(f"      PDF généré via {engine} : {PDF}")
    else:
        print("      [AVERTISSEMENT] PDF non généré (aucun moteur disponible). DOCX disponible.")
    print("[4/4] Contrôle qualité ...")
    ok = quality_check()
    print(f"\nPages PDF : {estimate_pages()}")
    print(f"DOCX : {DOCX if DOCX.exists() else 'ABSENT'}")
    print(f"PDF  : {PDF if PDF.exists() else 'ABSENT'}")
    print(f"Qualité globale : {'CONFORME' if ok else 'À VÉRIFIER'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
