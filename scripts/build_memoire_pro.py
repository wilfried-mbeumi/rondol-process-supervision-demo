# -*- coding: utf-8 -*-
"""
Livrable final — memoire_rondol_professionnel.docx / .pdf

Réutilise le pipeline éprouvé de scripts/build_memoire_final.py (page de garde
graphique à deux logos, page institutionnelle, sommaire automatique TOC, styles
Titre 1/2/3, en-têtes/pieds avec pagination, mapping figures/captures réelles,
légendes "Figure N — …", nettoyage des placeholders) et y ajoute :

  - sortie aux noms demandés, à la racine du dépôt :
        memoire_rondol_professionnel.docx
        memoire_rondol_professionnel.pdf
  - rendu de la bibliographie en retrait suspendu (hanging indent), lisible et
    académique, à partir des 38 références consolidées de l'état de l'art V5 ;
  - copie des figures du mémoire dans figures_memoire/ (livrable annexe) ;
  - conversion PDF via Word COM (met à jour le sommaire automatique) avec repli
    docx2pdf, puis contrôle qualité.

Usage : python scripts/build_memoire_pro.py
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import build_memoire_final as bmf  # noqa: E402  (pipeline éprouvé réutilisé)
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402

OUT_DOCX = ROOT / "MBEUMI_Wilfried_THESE.docx"
OUT_PDF = ROOT / "MBEUMI_Wilfried_THESE.pdf"
FIGURES_OUT = ROOT / "figures_memoire"

BIB_ITEM_RE = re.compile(r"^\d+\.\s")

# Titre de niveau 1 qui ouvre le corps du mémoire : c'est là que la pagination
# arabe redémarre à 1.
BODY_START_TITLE = "Introduction générale"


# --------------------------------------------------------------------------- #
# Pagination conforme à l'usage académique
#   - page de garde + page institutionnelle : aucun numéro ;
#   - pages liminaires (remerciements, résumé, abstract, glossaire, sommaire) :
#     chiffres romains minuscules i, ii, iii… ;
#   - corps du mémoire, à partir de l'introduction générale : chiffres arabes
#     redémarrant à 1.
# --------------------------------------------------------------------------- #
def set_page_numbering(section, fmt: str, start: int) -> None:
    """Impose le format et le point de départ de la numérotation d'une section."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def open_body_section(doc):
    """Ouvre la section du corps : nouvelle page, en-tête/pied, arabe à partir de 1."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    bmf.configure_page(section)
    nexa_png, rondol_png = bmf.ensure_logos()
    bmf.build_header_footer(section, nexa_png, rondol_png)
    set_page_numbering(section, "decimal", 1)
    return section


# --------------------------------------------------------------------------- #
# Parser de corps « bibliography-aware » (remplace bmf.parse_body)
#   identique au pipeline éprouvé, avec un rendu dédié des références
#   bibliographiques en retrait suspendu.
# --------------------------------------------------------------------------- #
def parse_body(doc, md_text):
    idx = md_text.find("## Remerciements")
    lines = md_text[idx if idx >= 0 else 0:].splitlines()
    i, first_h1, just_broke = 0, True, False
    in_biblio = False
    body_started = [False]   # liste = cellule mutable partagée avec la boucle
    while i < len(lines):
        stripped = lines[i].rstrip().strip()

        # Tableaux
        if stripped.startswith("|") and stripped.endswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            block = [b for b in block if not bmf.is_table_sep(b)]
            bmf.render_table(doc, block); just_broke = False
            continue

        # Code fence
        if stripped.startswith("```"):
            i += 1; code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            bmf.render_code(doc, code); just_broke = False
            continue

        i += 1
        if not stripped or stripped in ("---", "***", "___"):
            continue
        if stripped == "[[TOC]]":
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
            bmf._run(p, "Sommaire", 16, bmf.BLUE, bold=True)
            bmf.add_toc(doc)
            just_broke = False
            continue
        if stripped == "[SAUT DE PAGE]":
            # Anticipation : un titre de niveau 1 produit déjà son propre saut
            # de page. Un [SAUT DE PAGE] juste avant en ajouterait un second et
            # laisserait une page entièrement blanche dans le document.
            k = i
            while k < len(lines) and (not lines[k].strip()
                                      or lines[k].strip() in ("---", "***", "___")):
                k += 1
            next_is_h1 = k < len(lines) and lines[k].strip().startswith("# ")
            if next_is_h1:
                just_broke = False        # on laisse le titre gérer le saut
            elif not just_broke:
                doc.add_page_break(); just_broke = True
            continue
        if stripped.startswith("!["):
            continue

        # Placeholders figure/capture -> image réelle
        if stripped.startswith("[INSÉRER FIGURE") or stripped.startswith("[INSÉRER CAPTURE"):
            inner = stripped.strip("[]")
            val = bmf.resolve_placeholder(inner)
            if val and val[0]:
                bmf.add_figure(doc, val[0], val[1])
            just_broke = False
            continue

        # Titres
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            # L'introduction générale ouvre le corps : on y remplace le simple
            # saut de page par une rupture de SECTION, seule façon de faire
            # repartir la numérotation à 1 en chiffres arabes après les pages
            # liminaires en chiffres romains.
            if title == BODY_START_TITLE and not body_started[0]:
                open_body_section(doc)
                body_started[0] = True
            elif not first_h1 and not just_broke:
                doc.add_page_break()
            first_h1 = False
            in_biblio = title.lower().startswith("bibliographie")
            doc.add_paragraph(title, style="Heading 1")
            just_broke = False
            continue
        if stripped.startswith("## "):
            in_biblio = False
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2"); just_broke = False
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3"); just_broke = False
            continue

        # Citation
        if stripped.startswith(">"):
            content = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            bmf.add_runs(p, content)
            for r in p.runs:
                r.font.italic = True
            continue

        # Référence bibliographique : retrait suspendu, compact, justifié
        if in_biblio and BIB_ITEM_RE.match(stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.left_indent = Cm(0.8)
            pf.first_line_indent = Cm(-0.8)
            pf.space_after = Pt(3)
            pf.line_spacing = 1.05
            bmf.add_runs(p, stripped, default_size=9.5)
            for r in p.runs:
                if r.font.size is None or r.font.size > Pt(9.5):
                    r.font.size = Pt(9.5)
            continue

        # Légende / note italique pure
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("*"))
            r.font.italic = True; r.font.size = Pt(9.5)
            r.font.color.rgb = bmf.GREY; r.font.name = bmf.BASE_FONT
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bmf.add_runs(p, stripped)


# --------------------------------------------------------------------------- #
# Conversion PDF (Word COM -> met à jour le TOC ; repli docx2pdf)
# --------------------------------------------------------------------------- #
WD_PDF = 17


def convert_pdf(docx_path: Path, pdf_path: Path) -> str:
    # 1) Word COM : met à jour les champs et le sommaire avant export
    try:
        import win32com.client as win32
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            doc = word.Documents.Open(str(docx_path))
            doc.Fields.Update()
            for k in range(1, doc.TablesOfContents.Count + 1):
                doc.TablesOfContents.Item(k).Update()
            doc.Repaginate()
            doc.Fields.Update()
            doc.Save()
            doc.ExportAsFixedFormat(OutputFileName=str(pdf_path),
                                    ExportFormat=WD_PDF, OpenAfterExport=False)
            doc.Close(False)
            return "word-com"
        finally:
            word.Quit()
    except Exception as exc:
        print(f"[WARN] Word COM indisponible: {exc}")

    # 2) Repli docx2pdf
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return "docx2pdf"
    except Exception as exc:
        print(f"[WARN] docx2pdf indisponible: {exc}")
    return "none"


# --------------------------------------------------------------------------- #
def copy_figures():
    FIGURES_OUT.mkdir(parents=True, exist_ok=True)
    n = 0
    for src_dir in (bmf.FIGDIR, bmf.CAPDIR):
        if src_dir.exists():
            for p in sorted(src_dir.glob("*.png")):
                shutil.copy(p, FIGURES_OUT / p.name); n += 1
    return n


def qc(docx_path: Path):
    from docx import Document
    doc = Document(str(docx_path))
    full = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                full += "\n" + c.text
    n_images = len(doc.inline_shapes)
    checks = {
        "Auteur Wilfried Galtier MBEUMI": "Wilfried Galtier MBEUMI" in full,
        "Nexa Digital School": "Nexa Digital School" in full,
        "Tuteur M. Moussa NDIAYE": "Moussa NDIAYE" in full,
        "Tuteur M. Maël Gallas": "Maël Gallas" in full,
        "Rondol Industrie": "Rondol Industrie" in full,
        "Remerciements": "Remerciements" in full,
        "Résumé / Abstract": "Résumé" in full and "Abstract" in full,
        "Sommaire (TOC)": "Sommaire" in full,
        "Introduction générale": "Introduction générale" in full,
        "Conclusion générale": "Conclusion générale" in full,
        "Bibliographie 38 réfs (38.)": "38." in full,
        "Coquille C4 corrigée (7.5 Cas C4)": "Cas C4 — recommandation" in full,
        "Annexes": "Annexe" in full,
        "Légendes Figure N": "Figure 1" in full,
        "PAS de 'À COMPLÉTER'": "À COMPLÉTER" not in full,
        "PAS de 'INSÉRER'": "INSÉRER" not in full,
        "PAS de '##' résiduel": "##" not in full,
        "PAS de '**' résiduel": "**" not in full,
        "Images >= 25": n_images >= 25,
    }
    print("\n=== CONTRÔLE QUALITÉ ===")
    ok_all = True
    for k, ok in checks.items():
        print(f"  [{'OK ' if ok else 'KO '}] {k}")
        ok_all = ok_all and ok
    print(f"  Images inline : {n_images} | Paragraphes : {len(doc.paragraphs)} | Tableaux : {len(doc.tables)}")
    return ok_all


def main():
    # Réorientation du pipeline éprouvé vers notre rendu et nos sorties
    bmf.parse_body = parse_body
    bmf.DOCX = OUT_DOCX
    bmf.PDF = OUT_PDF

    if not bmf.MD.exists():
        print(f"[ERREUR] Source introuvable : {bmf.MD}"); return 1

    print("[1/4] Construction du document professionnel ...")
    doc = bmf.build_document()

    # Pagination : les pages liminaires (section 2, ouverte par bmf après la
    # page de garde et la page institutionnelle) passent en chiffres romains
    # minuscules. La section du corps, ouverte à l'introduction générale par
    # open_body_section(), est déjà en arabe redémarrant à 1.
    # Les sections 0 et 1 — page de garde et page institutionnelle — ne sont
    # pas touchées : elles restent sans en-tête, sans pied et sans numéro.
    sections = doc.sections
    if len(sections) >= 2:
        set_page_numbering(sections[1], "lowerRoman", 1)
    n_body = sum(1 for s in sections[2:])
    print(f"      Pagination : liminaires en romains, "
          f"corps en arabes ({n_body} section(s) de corps)")

    doc.save(str(OUT_DOCX))
    print(f"      DOCX : {OUT_DOCX}")
    print(f"      Figures insérées : {bmf.fig_counter[0]}")

    print("[2/4] Copie des figures (figures_memoire/) ...")
    nfig = copy_figures()
    print(f"      {nfig} fichiers copiés dans {FIGURES_OUT}")

    print("[3/4] Conversion PDF (mise à jour du sommaire) ...")
    engine = convert_pdf(OUT_DOCX, OUT_PDF)
    ok_pdf = OUT_PDF.exists()
    print(f"      PDF  : {'OK (' + engine + ') ' + str(OUT_PDF) if ok_pdf else 'ÉCHEC'}")

    print("[4/4] Contrôle qualité ...")
    ok_all = qc(OUT_DOCX)
    try:
        from pypdf import PdfReader
        npages = len(PdfReader(str(OUT_PDF)).pages) if ok_pdf else None
    except Exception:
        npages = None
    print(f"\nDOCX : {OUT_DOCX if OUT_DOCX.exists() else 'ABSENT'}")
    print(f"PDF  : {OUT_PDF if ok_pdf else 'ABSENT'} | Pages : {npages}")
    print(f"Qualité globale : {'CONFORME' if ok_all else 'À VÉRIFIER'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
