"""
Build one-page industrial scientific abstract v2.

Output:
    reports/poster_abstract/Mbeumi_2026_Abstract_v2.docx
    reports/poster_abstract/Mbeumi_2026_Abstract_v2.pdf

Constraints:
    - Single A4 page
    - 2 figures only (Streamlit cockpit + before/after C3 to C5)
    - Discreet Rondol logo in header
    - Industrial / scientific tone — not marketing
    - Explicit coverage of manager's feature requests
    - Cases C1..C5 narrated as a scientific demonstration loop
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets"
FIG = ROOT / "reports" / "poster_abstract" / "figures" / "generated_v2"
OUT_DIR = ROOT / "reports" / "poster_abstract"
OUT_DOCX = OUT_DIR / "Mbeumi_2026_Abstract_v2.docx"
OUT_PDF = OUT_DIR / "Mbeumi_2026_Abstract_v2.pdf"

LOGO = ASSETS / "rondol_logo.png"
FIG1 = FIG / "figv2_03_agent_alert_panel.png"          # cockpit Streamlit
FIG2 = FIG / "figv2_05_before_after_dashboard.png"     # before / after C3 -> C5

DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x0B, 0x3D, 0x91)
RED = RGBColor(0xB0, 0x1C, 0x1C)
GREEN = RGBColor(0x1C, 0x6E, 0x2E)

BASE_FONT = "Calibri"


# ----------------------------------------------------------------------------
# helpers
# ----------------------------------------------------------------------------


def set_run(run, *, bold=False, italic=False, size=9, color=DARK, font=BASE_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, space_before=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.08
    return p


def add_run(p, text, **kwargs):
    r = p.add_run(text)
    set_run(r, **kwargs)
    return r


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def remove_table_borders(table):
    tbl_pr = table._element.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)


def thin_bottom_border(paragraph, color="0B3D91"):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


# ----------------------------------------------------------------------------
# build
# ----------------------------------------------------------------------------


def build():
    doc = Document()

    # Page setup — A4, tight margins to fit on one page
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    style = doc.styles["Normal"]
    style.font.name = BASE_FONT
    style.font.size = Pt(9)

    # -------------------------------------------------------------------------
    # Header band: small logo (right) + conference identifiers (left)
    # -------------------------------------------------------------------------
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl.autofit = False
    header_tbl.columns[0].width = Cm(13.5)
    header_tbl.columns[1].width = Cm(4.1)
    remove_table_borders(header_tbl)

    left = header_tbl.cell(0, 0)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left.width = Cm(13.5)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Industrial abstract  ·  AI-assisted twin-screw extrusion",
            size=8, color=GREY, italic=True)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, "Submission target: scientific poster / abstract — 2026",
            size=8, color=GREY)

    right = header_tbl.cell(0, 1)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.width = Cm(4.1)
    p_logo = right.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_logo = p_logo.add_run()
    r_logo.add_picture(str(LOGO), width=Cm(2.6))

    rule = add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    thin_bottom_border(rule)

    # -------------------------------------------------------------------------
    # Title + author block
    # -------------------------------------------------------------------------
    p_title = add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    add_run(p_title,
            "AI-assisted twin-screw extrusion: real-time process supervision "
            "and parameter optimisation for Li-bearing dry / semi-dry battery formulations",
            bold=True, size=12, color=DARK)

    p_auth = add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    add_run(p_auth, "Wilfried Galtier Mbeumi", size=9, color=DARK)
    add_run(p_auth, "  ·  Rondol Industrie, Nancy (FR)  ·  Institut Jean Lamour (IJL), Campus ARTEM, Nancy (FR)",
            size=9, color=GREY)
    p_sup = add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_run(p_sup, "Industrial supervisor: Maël Gallas (Rondol Industrie).  ",
            size=8, italic=True, color=GREY)
    add_run(p_sup,
            "Context: Hot-Melt Extrusion (HME) transfer to Li-ion / solid-state battery components — an emerging, under-published intersection of extrusion engineering, machine learning and electrode formulation.",
            size=8, italic=True, color=GREY)

    # -------------------------------------------------------------------------
    # Section helper
    # -------------------------------------------------------------------------
    def section_title(text):
        p = add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=3, space_after=1)
        add_run(p, text.upper(), bold=True, size=9, color=ACCENT)

    # -------------------------------------------------------------------------
    # Background
    # -------------------------------------------------------------------------
    section_title("Background & objective")
    p = add_paragraph(doc)
    add_run(p,
            "HME is a mature, solvent-free continuous process whose extension to lithium-ion and "
            "solid-state battery electrodes remains scarcely published. We report a software agent "
            "that ingests the full set of extrusion control variables — ", size=9)
    add_run(p,
            "positionable screw elements, temperature profile from zone 1 to die, screw speed, "
            "specific mechanical energy (SME), fill factor, residence time, free volume and material "
            "throughput",
            size=9, italic=True)
    add_run(p,
            " — together with a multi-feeder description (1 to 5 feeders; position, speed, mass flow "
            "rate, density and thermal expansion; feed media including granules, powders, liquid, "
            "semi-liquid, gas and supercritical fluids). The agent returns hierarchised process "
            "recommendations and risk alerts in real time, with torque and pressure foreseen as "
            "first-class signals in version v2.",
            size=9)

    # -------------------------------------------------------------------------
    # Materials & Methods
    # -------------------------------------------------------------------------
    section_title("Materials & methods")
    p = add_paragraph(doc)
    add_run(p, "Reference recipe (lithium-bearing, semi-dry): ", bold=True, size=9)
    add_run(p,
            "LFP 65 / PVDF 8 / Super-P 5 / LATP 17 / LiTFSI 5 wt% on a Rondol 10.5 mm twin-screw "
            "extruder (L/D 40, 8 thermal zones). ",
            size=9)
    add_run(p, "Dataset: ", bold=True, size=9)
    add_run(p,
            "11 industrial runs (7–13 April 2026), 8 retained after duration filter, segmented into "
            "627 sliding windows of 60 s (step 30 s); GroupShuffleSplit by run_id guarantees zero "
            "inter-run leakage. ",
            size=9)
    add_run(p, "AI stack: ", bold=True, size=9)
    add_run(p,
            "Streamlit frontend (4 pages); physics-based KPIs (fill factor, residence time, SME, "
            "free volume per zone); Random Forest / XGBoost / SVM classifiers; rule-based "
            "formulation score (5 weighted criteria) fused with the data-driven stability "
            "probability into a single hierarchised recommendation panel — proposing zone "
            "temperatures, screw element substitutions and feeder set-points.",
            size=9)

    # -------------------------------------------------------------------------
    # Results — narrated as C1..C5
    # -------------------------------------------------------------------------
    section_title("Results — demonstration loop on the Li recipe (C1 → C5)")
    p = add_paragraph(doc)
    add_run(p, "Model performance ", bold=True, size=9)
    add_run(p,
            "(Random Forest, 60 s window, n = 627): test accuracy ",
            size=9)
    add_run(p, "0.950", bold=True, size=9, color=ACCENT)
    add_run(p, ", F1-macro ", size=9)
    add_run(p, "0.917", bold=True, size=9, color=ACCENT)
    add_run(p, ", CV ROC-AUC ", size=9)
    add_run(p, "0.976 ± 0.021", bold=True, size=9, color=ACCENT)
    add_run(p,
            ". Top discriminating features (Gini) are dominated by downstream cast-film and die "
            "thermal variability, confirming that melt-quality stability — not raw zone set-points — "
            "drives classification.",
            size=9)

    # Compact case table (C1..C5)
    case_rows = [
        ("C1", "Baseline LFP",                     "score 65 / p_stab 0.84 / FF Z5 0.71",  "stable",   "green",  "proceed"),
        ("C2", "Process-optimised variant",        "score 72 / p_stab 0.88 / FF Z5 0.68",  "stable",   "green",  "favourable window"),
        ("C3", "Ceramic overload LATP 35 wt%",     "score 46 / p_stab 0.35 / FF Z5 0.97",  "unstable", "red",    "Z5 overflow, torque 84 %"),
        ("C4", "AI recommendation issued",         "LATP→17 % · Z4 kneading 30° · Z5 +5 °C", "advisory", "amber", "ranked actions"),
        ("C5", "Post-recommendation",              "score 78 / p_stab 0.87 / FF Z5 0.72",  "stable",   "green",  "alert cleared, Δ +32 pts"),
    ]
    tbl = doc.add_table(rows=len(case_rows) + 1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [Cm(0.9), Cm(4.2), Cm(6.7), Cm(2.2), Cm(3.8)]
    for col, w in zip(tbl.columns, widths):
        col.width = w
    remove_table_borders(tbl)

    # Header row
    hdr = ["", "Case", "Observable / agent output", "Class.", "Action"]
    for i, txt in enumerate(hdr):
        cell = tbl.rows[0].cells[i]
        cell.width = widths[i]
        shade_cell(cell, "0B3D91")
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_before = Pt(0)
        p_h.paragraph_format.space_after = Pt(0)
        add_run(p_h, txt, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))

    severity_color = {
        "green": GREEN,
        "red": RED,
        "amber": RGBColor(0xC2, 0x7A, 0x00),
    }
    severity_fill = {
        "green": "EAF4EC",
        "red":   "F8E6E6",
        "amber": "FBF1DD",
    }
    for r, (cid, label, obs, klass, sev, action) in enumerate(case_rows, start=1):
        row = tbl.rows[r]
        for i, w in enumerate(widths):
            row.cells[i].width = w
        shade_cell(row.cells[0], severity_fill[sev])
        p_id = row.cells[0].paragraphs[0]
        p_id.paragraph_format.space_before = Pt(0)
        p_id.paragraph_format.space_after = Pt(0)
        p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_id, cid, bold=True, size=8, color=severity_color[sev])

        for col_idx, txt in enumerate([label, obs, klass, action], start=1):
            cell = row.cells[col_idx]
            p_c = cell.paragraphs[0]
            p_c.paragraph_format.space_before = Pt(0)
            p_c.paragraph_format.space_after = Pt(0)
            bold = col_idx == 1
            italic = col_idx == 3
            add_run(p_c, txt, bold=bold, italic=italic, size=8, color=DARK)

    # Narrative point about predicted-vs-real drift — motivates v2 torque + pressure
    p = add_paragraph(doc, space_before=2)
    add_run(p, "Predicted-vs-real drift. ", bold=True, italic=True, size=9, color=DARK)
    add_run(p,
            "On one run the initial prediction was nominally stable (p_stable ≈ 0.78) yet downstream "
            "cast-film sensors revealed a slow drift the agent could not anticipate from "
            "temperatures and screw speed alone — motivating the v2 integration of live torque and "
            "pressure as first-class inputs.",
            italic=True, size=9, color=DARK)

    # -------------------------------------------------------------------------
    # Figures — side by side
    # -------------------------------------------------------------------------
    fig_tbl = doc.add_table(rows=2, cols=2)
    fig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fig_tbl.autofit = False
    fig_tbl.columns[0].width = Cm(8.8)
    fig_tbl.columns[1].width = Cm(8.8)
    remove_table_borders(fig_tbl)

    # row 0: images
    for idx, (cell, path) in enumerate(
        [(fig_tbl.cell(0, 0), FIG1), (fig_tbl.cell(0, 1), FIG2)]
    ):
        cell.width = Cm(8.8)
        p_fig = cell.paragraphs[0]
        p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fig.paragraph_format.space_before = Pt(2)
        p_fig.paragraph_format.space_after = Pt(0)
        run = p_fig.add_run()
        run.add_picture(str(path), width=Cm(8.4))

    # row 1: captions
    captions = [
        ("Figure 1.",
         " Streamlit supervision view — feeder set-points, extrusion parameters, "
         "AI compatibility score, stability probability, fill-factor / torque alerts, "
         "ranked recommendations."),
        ("Figure 2.",
         " Before → after (C3 → C5): applying the AI recommendation restores compatibility "
         "(46 → 78), stability probability (0.35 → 0.87), fill factor Z5 (0.97 → 0.72) "
         "and torque margin (84 → 62 %)."),
    ]
    for i, (head, body) in enumerate(captions):
        cell = fig_tbl.cell(1, i)
        cell.width = Cm(8.8)
        p_cap = cell.paragraphs[0]
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(0)
        add_run(p_cap, head, bold=True, size=7, color=ACCENT)
        add_run(p_cap, body, size=7, color=GREY)

    # -------------------------------------------------------------------------
    # Discussion & conclusion (compact, single paragraph)
    # -------------------------------------------------------------------------
    section_title("Discussion & outlook")
    p = add_paragraph(doc, space_after=0)
    add_run(p,
            "Results are encouraging on a limited 8-run industrial dataset and remain advisory: "
            "final decisions stay with the process engineer. The rule-based formulation score will "
            "be replaced by a regression model trained on a literature-derived recipe corpus; the "
            "agent will be extended to NMC and sulfide solid electrolytes, with in-line torque and "
            "pressure closing the loop in v2. The combination ",
            size=9)
    add_run(p, "extrusion + AI + battery formulation ", bold=True, size=9, color=ACCENT)
    add_run(p,
            "addresses a documented scientific gap and supports a credible TRL 4–5 industrial path "
            "for Rondol Industrie.",
            size=9)

    # Footer micro-line
    p_foot = add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=2)
    add_run(p_foot,
            "Acknowledgements: Rondol Industrie · Institut Jean Lamour · Campus ARTEM · Mastère Data & IA (RNCP 37137).",
            size=7, italic=True, color=GREY)

    # -------------------------------------------------------------------------
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"DOCX written: {OUT_DOCX}")

    # PDF
    from docx2pdf import convert
    convert(str(OUT_DOCX), str(OUT_PDF))
    print(f"PDF written:  {OUT_PDF}")


if __name__ == "__main__":
    build()
