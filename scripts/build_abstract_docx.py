"""Genere l'abstract scientifique final (DOCX + PDF) pour le 15 mai 2026.

Sortie :
  reports/poster_abstract/Mbeumi_2026_AbstractEN.docx
  reports/poster_abstract/Mbeumi_2026_AbstractEN.pdf

Convention :
  A4 portrait, marges 1.6 cm, Calibri 10 pt, 1 colonne, 2 figures integrees,
  format "extended abstract" 1 page (texte) + figures inline.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "reports" / "poster_abstract" / "figures" / "generated"
OUT_DIR = ROOT / "reports" / "poster_abstract"
DOCX_PATH = OUT_DIR / "Mbeumi_2026_AbstractEN.docx"
PDF_PATH = OUT_DIR / "Mbeumi_2026_AbstractEN.pdf"

GREEN = RGBColor(0x1B, 0x7A, 0x3D)
BLUE = RGBColor(0x00, 0x5B, 0x96)
GREY = RGBColor(0x33, 0x33, 0x33)
MID_GREY = RGBColor(0x66, 0x66, 0x66)


def _set_run_font(run, *, size=10, bold=False, italic=False, color=GREY, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), name)


def add_paragraph(doc, text, *, size=9.5, bold=False, italic=False, color=GREY,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, *, color=GREEN, size=10.5, space_before=2, keep_next=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.keep_with_next = keep_next
    run = p.add_run(text.upper())
    _set_run_font(run, size=size, bold=True, color=color)
    return p


def add_runs(doc, segments, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1):
    """segments : list of (text, dict) where dict has font kwargs."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for text, kwargs in segments:
        run = p.add_run(text)
        kw = {"size": 9.5, "color": GREY}
        kw.update(kwargs)
        _set_run_font(run, **kw)
    return p


def add_figure_centered(doc, image_path: Path, *, width_cm: float, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.space_before = Pt(0)
    run = cap.add_run(caption)
    _set_run_font(run, size=8.5, italic=True, color=MID_GREY)


def setup_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)
        section.footer_distance = Cm(0.5)


def install_page_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Rondol Industrie  ·  Institut Jean Lamour  ·  Mastère Data & IA "
        "(RNCP 37137)  ·  Corresponding author: wilfried.mbeumi@rondol.com"
    )
    _set_run_font(run, size=7.5, italic=True, color=MID_GREY)


def build_document():
    doc = Document()
    setup_margins(doc)

    # ----- Title block -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    title.paragraph_format.space_before = Pt(0)
    run = title.add_run(
        "AI-Assisted Twin-Screw Extrusion for Lithium-Bearing Battery Electrode "
        "Manufacturing — A Decision-Support Demonstrator"
    )
    _set_run_font(run, size=13, bold=True, color=GREEN)

    add_runs(doc, [
        ("Wilfried Galtier Mbeumi", {"size": 9.5, "bold": True, "color": GREY}),
        ("¹·²", {"size": 7.5, "color": GREY}),
        ("  ·  Maël Gallas", {"size": 9.5, "color": GREY}),
        ("¹", {"size": 7.5, "color": GREY}),
        ("    ", {"size": 9.5, "color": GREY}),
        ("¹ Rondol Industrie, Nancy, France  ·  ² Institut Jean Lamour (IJL), "
         "Campus ARTEM, Université de Lorraine, Nancy, France",
         {"size": 8, "italic": True, "color": MID_GREY}),
    ], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # ----- 1. Introduction -----
    add_heading(doc, "1. Introduction")
    add_runs(doc, [
        ("Hot Melt Extrusion (HME) is a mature, solvent-free continuous process "
         "originally developed for pharmaceutical applications [1]. Its transfer to "
         "lithium-ion (LIB) and emerging solid-state battery (SSB) electrode "
         "manufacturing — cathodes, anodes, composite electrolytes — has attracted "
         "growing interest as a route to dry / semi-dry processing aligned with the "
         "European proposed restriction of PVDF and other PFAS [2]. However, "
         "peer-reviewed work integrating ", {}),
        ("extrusion engineering, machine learning, and battery electrode "
         "formulation", {"bold": True}),
        (" in a single framework remains scarce [3–5]. We introduce an "
         "AI-assisted decision-support demonstrator linking ", {}),
        ("formulation → process → screw profile → risk → recommendation",
         {"italic": True, "color": BLUE}),
        (" on a lithium-bearing dry / semi-dry recipe.", {}),
    ])

    # ----- 2. Materials & Methods -----
    add_heading(doc, "2. Materials & Methods")
    add_paragraph(doc,
        "Eleven industrial runs were collected on a Rondol Ø 10.5 mm twin-screw "
        "extruder (L/D 40:1, horizontal) between 7 and 13 April 2026. After a "
        "duration filter (≥ 15 min per run), eight runs were retained and segmented "
        "into 627 sliding windows of 60 s (step 30 s; 87 features per window). "
        "Binary stability labels were assigned by expert rules on the local thermal "
        "variability of barrel zones Z1–Z8, the die head (DIE), and three downstream "
        "cast-film sensors (P1, P2, Body). Three classifiers — Random Forest, "
        "XGBoost, SVM (RBF) — were trained with GroupShuffleSplit over run_id "
        "(5 train / 3 test) to prevent inter-run leakage. The AI agent combines this "
        "data-driven stability classification with a rule-based compatibility score "
        "(5 weighted criteria: ceramic load, viscosity, thermal compatibility, "
        "abrasion risk, material availability) evaluated on the lithium reference "
        "recipe LFP 65 / PVDF 8 / Super P 5 / LATP 17 / LiTFSI 5 wt%, and emits a "
        "hierarchised textual recommendation.",
    )

    # ----- 3. Results -----
    add_heading(doc, "3. Results")
    add_runs(doc, [
        ("On the held-out test set, Random Forest reaches ", {}),
        ("accuracy 0.950, F1-macro 0.917, ROC-AUC 0.976 ± 0.021 (5-fold CV)", {"bold": True}),
        (" and is retained as the production model on the basis of balanced "
         "performance and inference cost (Fig. 1). The top-10 features by Gini "
         "importance are dominated by ", {}),
        ("downstream cast-film and die sensors", {"bold": True}),
        (" (CastFilmP2_iqr 0.072 ; DIE_std 0.067 ; CastFilmBody_std 0.066), with "
         "barrel-zone temperatures appearing only beyond rank 13 — confirming that "
         "melt-quality variability, rather than nominal zone setpoints, is the most "
         "discriminating signal. The 60 s window outperforms 30 s (F1 0.916) and "
         "120 s (F1 0.843). Five test cases built on the lithium reference recipe "
         "exercise the full agent loop: ", {}),
        ("C1", {"bold": True}),
        (" — baseline, score 65/100, stable ; ", {}),
        ("C3", {"bold": True}),
        (" — LATP overload at 35 wt%, score 46/100, red alert on zone Z5 ; ", {}),
        ("C4", {"bold": True}),
        (" — AI recommendation (reduce LATP to 17–20 %, soften Z4 kneading to 30°, "
         "+5 °C on Z5) ; ", {}),
        ("C5", {"bold": True}),
        (" — recommendation applied, score 78/100, p_stable = 0.87, alert cleared "
         "(Fig. 2).", {}),
    ])

    # Two figures side by side: ML perf + before/after
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    cell_tl, cell_tr = table.cell(0, 0), table.cell(0, 1)
    cell_bl, cell_br = table.cell(1, 0), table.cell(1, 1)
    for cell in (cell_tl, cell_tr, cell_bl, cell_br):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
    # figure 1 (ML perf)
    p = cell_tl.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / "fig03_ml_performance_w60.png"), width=Cm(8.2))
    cap = cell_bl.paragraphs[0]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Fig. 1  —  Test confusion matrices for RF / XGBoost / SVM "
                    "(60 s window, n=340).")
    _set_run_font(r, size=7.5, italic=True, color=MID_GREY)
    # figure 2 (before/after)
    p = cell_tr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / "fig05_before_after_recommendation.png"), width=Cm(8.2))
    cap = cell_br.paragraphs[0]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Fig. 2  —  Effect of the AI recommendation: KPI before (C3) "
                    "vs. after (C5).")
    _set_run_font(r, size=7.5, italic=True, color=MID_GREY)

    # ----- 4. Discussion -----
    add_heading(doc, "4. Discussion")
    add_paragraph(doc,
        "Results are encouraging but rest on a limited industrial dataset (8 runs, "
        "627 windows) and a rule-based compatibility score that will be progressively "
        "replaced by a regression model trained on a literature-derived dataset of "
        "~50 lithium-bearing recipes. The dominance of downstream sensors in the "
        "feature ranking is consistent with a process well-regulated upstream of the "
        "screw, yet underlines that the most informative observability for stability "
        "decisions sits at the die and the cast film, not at the barrel setpoints. "
        "The 32-point compatibility-score gain and 0.52 absolute increase in "
        "stability probability between C3 and C5 demonstrate that the recommendation "
        "loop is non-trivial. Known limitations: no electrochemical validation yet "
        "(thermal proxy only), simplified abrasion model, and no closed-loop "
        "integration of in-line sensors [7].",
    )

    # ----- 5. Conclusion -----
    add_heading(doc, "5. Conclusion")
    add_paragraph(doc,
        "This work demonstrates the technical and industrial feasibility (TRL 4–5) of "
        "an AI-augmented decision-support tool for the HME of lithium-bearing "
        "battery components. Planned extensions include a literature-derived dataset "
        "of ~50 recipes, SHAP-based local interpretability, in-line sensor closure "
        "of the recommendation loop, and scale-up validation on the Rondol 21 mm "
        "platform. The integration extrusion + AI + batteries addresses a documented "
        "scientific gap and a strategic opportunity in the SSB value chain.",
    )

    # ----- References (compact inline) -----
    add_heading(doc, "References", color=GREY, size=9.5, space_before=2, keep_next=False)
    refs_inline = (
        "[1] Repka et al., Int. J. Pharm. 535, 2018.  "
        "[2] ECHA, Annex XV restriction report on PFAS, 2023.  "
        "[3] Drakopoulos et al., Cell Rep. Phys. Sci. 2, 100689, 2021.  "
        "[4] Haarmann et al., Energy Technol. 9, 2021.  "
        "[5] Seeba et al., Batteries 10, 2024.  "
        "[6] Kim et al., Nat. Commun. 14, 2023.  "
        "[7] Maia et al., AMI Plastics & Completion AI, 2025.  "
        "[8] Wang et al., Nano-Micro Lett. 17, 2025.  "
        "[9] Daoudi et al., J. Power Sources 590, 2024.  "
        "[10] Fraunhofer IWS, DRYtraec® continuous manufacturing, 2024."
    )
    add_paragraph(doc, refs_inline, size=7.5, color=MID_GREY, space_after=0)

    install_page_footer(doc)
    doc.save(DOCX_PATH)
    print(f"[OK] DOCX  -> {DOCX_PATH.relative_to(ROOT)}")


def convert_to_pdf():
    try:
        from docx2pdf import convert
    except ImportError:
        print("[WARN] docx2pdf manquant — pip install docx2pdf")
        return
    convert(str(DOCX_PATH), str(PDF_PATH))
    print(f"[OK] PDF   -> {PDF_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    build_document()
    convert_to_pdf()
