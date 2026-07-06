"""
Build the FINAL one-page industrial scientific abstract.

Output:
    reports/poster_abstract/Mbeumi_2026_Abstract_FINAL.docx
    reports/poster_abstract/Mbeumi_2026_Abstract_FINAL.pdf

Light finishing pass on top of v3 — no new section, no new figure, same density:
    - Side margins   1.7 → 1.5 cm  (gives back ~4 mm of usable width)
    - Figure cells   8.8 → 9.0 cm  ·  pictures 8.4 → 8.85 cm  (+5 % wider, better legibility)
    - Header / cases / figures tables widened to span the new usable area
    - Governance sentence merged with field-validation: "...manually reviewed by
      process engineers and subsequently tested on extrusion trials under industrial
      supervision (human-in-the-loop governance)."
    - Figures reused from generated_v3_abstract/ (already publication-ready)
"""

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_abstract_v3 as v3  # noqa: E402  (constants + helpers + figure generators)

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "reports" / "poster_abstract"
OUT_DOCX = OUT_DIR / "Mbeumi_2026_Abstract_FINAL.docx"
OUT_PDF = OUT_DIR / "Mbeumi_2026_Abstract_FINAL.pdf"


# Pull all the helpers from v3 so styling stays 1:1
_par = v3._par
_r = v3._r
_no_borders = v3._no_borders
_shade = v3._shade
_rule_under = v3._rule_under
DARK, GREY, ACCENT, RED, GREEN_D = v3.DARK, v3.GREY, v3.ACCENT, v3.RED, v3.GREEN_D
BASE_FONT = v3.BASE_FONT
LOGO = v3.LOGO
FIG1 = v3.FIG1
FIG2 = v3.FIG2


# Layout constants (the only real change vs v3)
SIDE_MARGIN = Cm(1.5)         # was 1.7
USABLE_W = Cm(18.0)           # 21 - 2*1.5
HEADER_LEFT_W = Cm(13.7)
HEADER_RIGHT_W = Cm(4.3)
FIG_CELL_W = Cm(9.0)          # was 8.8
FIG_PICT_W = Cm(8.85)         # was 8.4 — ~+5% wider
CASE_WIDTHS = [Cm(0.95), Cm(4.30), Cm(6.85), Cm(2.25), Cm(3.85)]   # ~+0.2 cm total


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = SIDE_MARGIN
    section.right_margin = SIDE_MARGIN

    style = doc.styles["Normal"]
    style.font.name = BASE_FONT
    style.font.size = Pt(9)

    # ---- Header band ----
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl.autofit = False
    header_tbl.columns[0].width = HEADER_LEFT_W
    header_tbl.columns[1].width = HEADER_RIGHT_W
    _no_borders(header_tbl)

    left = header_tbl.cell(0, 0)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left.width = HEADER_LEFT_W
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _r(p, "Industrial abstract  ·  AI-assisted twin-screw extrusion",
       size=8, color=GREY, italic=True)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _r(p2, "Submission target: scientific poster / abstract — 2026",
       size=8, color=GREY)

    right = header_tbl.cell(0, 1)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.width = HEADER_RIGHT_W
    p_logo = right.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_logo = p_logo.add_run()
    r_logo.add_picture(str(LOGO), width=Cm(2.6))

    rule = _par(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    _rule_under(rule)

    # ---- Title block ----
    p_title = _par(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    _r(p_title,
       "AI-assisted twin-screw extrusion: real-time process supervision "
       "and parameter optimisation for Li-bearing dry / semi-dry battery formulations",
       bold=True, size=12, color=DARK)

    p_auth = _par(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    _r(p_auth, "Wilfried Galtier Mbeumi", size=9, color=DARK)
    _r(p_auth,
       "  ·  Rondol Industrie, Nancy (FR)  ·  Institut Jean Lamour (IJL), Campus ARTEM, Nancy (FR)",
       size=9, color=GREY)
    p_sup = _par(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    _r(p_sup, "Industrial supervisor: Maël Gallas (Rondol Industrie).  ",
       size=8, italic=True, color=GREY)
    _r(p_sup,
       "Context: Hot-Melt Extrusion (HME) transfer to Li-ion / solid-state battery components — "
       "an emerging, under-published intersection of extrusion engineering, machine learning and "
       "electrode formulation.",
       size=8, italic=True, color=GREY)

    def sec(text_):
        p = _par(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=3, space_after=1)
        _r(p, text_.upper(), bold=True, size=9, color=ACCENT)

    # ---- Background ----
    sec("Background & objective")
    p = _par(doc)
    _r(p,
       "HME is a mature, solvent-free continuous process whose extension to lithium-ion and "
       "solid-state battery electrodes remains scarcely published. We report a software agent "
       "that ingests the full set of extrusion control variables — ", size=9)
    _r(p,
       "positionable screw elements, temperature profile from zone 1 to die, screw speed, "
       "specific mechanical energy (SME), fill factor, residence time, free volume and material "
       "throughput",
       size=9, italic=True)
    _r(p,
       " — together with a multi-feeder description (1 to 5 feeders; position, speed, mass flow "
       "rate, density and thermal expansion; feed media including granules, powders, liquid, "
       "semi-liquid, gas and supercritical fluids). The agent returns hierarchised process "
       "recommendations and risk alerts in real time, with torque and pressure foreseen as "
       "first-class signals in version v2.",
       size=9)

    # ---- Materials & methods (governance + field-test merged in one sentence) ----
    sec("Materials & methods")
    p = _par(doc)
    _r(p, "Reference recipe (lithium-bearing, semi-dry): ", bold=True, size=9)
    _r(p,
       "LFP 65 / PVDF 8 / Super-P 5 / LATP 17 / LiTFSI 5 wt% on a Rondol 10.5 mm twin-screw "
       "extruder (L/D 40, 8 thermal zones). ",
       size=9)
    _r(p, "Dataset: ", bold=True, size=9)
    _r(p,
       "11 industrial runs (7–13 April 2026), 8 retained after duration filter, segmented into "
       "627 sliding windows of 60 s (step 30 s); GroupShuffleSplit by run_id guarantees zero "
       "inter-run leakage. ",
       size=9)
    _r(p, "AI stack: ", bold=True, size=9)
    _r(p,
       "Streamlit frontend (4 pages); physics-based KPIs (fill factor, residence time, SME, "
       "free volume per zone); Random Forest / XGBoost / SVM classifiers; a rule-based "
       "formulation score (5 weighted criteria) is fused with the data-driven stability "
       "probability into a single hierarchised recommendation panel — proposing zone "
       "temperatures, screw element substitutions and feeder set-points. ",
       size=9)
    _r(p,
       "All agent recommendations were manually reviewed by process engineers and "
       "subsequently tested on extrusion trials under industrial supervision "
       "(human-in-the-loop governance).",
       size=9, bold=True)

    # ---- Results ----
    sec("Results — demonstration loop on the Li recipe (C1 → C5)")
    p = _par(doc)
    _r(p, "Model performance ", bold=True, size=9)
    _r(p,
       "(Random Forest, 60 s window, n = 627): test accuracy ", size=9)
    _r(p, "0.950", bold=True, size=9, color=ACCENT)
    _r(p, ", F1-macro ", size=9)
    _r(p, "0.917", bold=True, size=9, color=ACCENT)
    _r(p, ", CV ROC-AUC ", size=9)
    _r(p, "0.976 ± 0.021", bold=True, size=9, color=ACCENT)
    _r(p,
       ". Top discriminating features (Gini) are dominated by downstream cast-film and die "
       "thermal variability, confirming that melt-quality stability — not raw zone set-points — "
       "drives classification.",
       size=9)

    case_rows = [
        ("C1", "Baseline LFP",                 "score 65 / p_stab 0.84 / FF Z5 0.71",   "stable",   "green",  "proceed"),
        ("C2", "Process-optimised variant",    "score 72 / p_stab 0.88 / FF Z5 0.68",   "stable",   "green",  "favourable window"),
        ("C3", "Ceramic overload LATP 35 wt%", "score 46 / p_stab 0.35 / FF Z5 0.97",   "unstable", "red",    "Z5 overflow, torque 84 %"),
        ("C4", "AI recommendation issued",     "LATP→17 % · Z4 kneading 30° · Z5 +5 °C", "advisory","amber",  "ranked actions"),
        ("C5", "Post-recommendation",          "score 78 / p_stab 0.87 / FF Z5 0.72",   "stable",   "green",  "alert cleared, Δ +32 pts"),
    ]
    tbl = doc.add_table(rows=len(case_rows) + 1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for col, w in zip(tbl.columns, CASE_WIDTHS):
        col.width = w
    _no_borders(tbl)

    hdr = ["", "Case", "Observable / agent output", "Class.", "Action"]
    for i, t in enumerate(hdr):
        cell = tbl.rows[0].cells[i]
        cell.width = CASE_WIDTHS[i]
        _shade(cell, "0B3D91")
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_before = Pt(0)
        p_h.paragraph_format.space_after = Pt(0)
        _r(p_h, t, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))

    sev_text = {"green": GREEN_D, "red": RED,
                "amber": RGBColor(0xC2, 0x7A, 0x00)}
    sev_fill = {"green": "EAF4EC", "red": "F8E6E6", "amber": "FBF1DD"}
    for r, (cid, label, obs, klass, sev, action) in enumerate(case_rows, start=1):
        row = tbl.rows[r]
        for i, w in enumerate(CASE_WIDTHS):
            row.cells[i].width = w
        _shade(row.cells[0], sev_fill[sev])
        p_id = row.cells[0].paragraphs[0]
        p_id.paragraph_format.space_before = Pt(0)
        p_id.paragraph_format.space_after = Pt(0)
        p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _r(p_id, cid, bold=True, size=8, color=sev_text[sev])

        for col_idx, t in enumerate([label, obs, klass, action], start=1):
            cell = row.cells[col_idx]
            p_c = cell.paragraphs[0]
            p_c.paragraph_format.space_before = Pt(0)
            p_c.paragraph_format.space_after = Pt(0)
            _r(p_c, t,
               bold=(col_idx == 1), italic=(col_idx == 3),
               size=8, color=DARK)

    p = _par(doc, space_before=2)
    _r(p, "Predicted-vs-real drift. ",
       bold=True, italic=True, size=9, color=DARK)
    _r(p,
       "On one run the initial prediction was nominally stable (p_stable ≈ 0.78) yet downstream "
       "cast-film sensors revealed a slow drift the agent could not anticipate from temperatures "
       "and screw speed alone — motivating the v2 integration of live torque and pressure as "
       "first-class inputs.",
       italic=True, size=9, color=DARK)

    # ---- Figures (slightly wider for legibility) ----
    fig_tbl = doc.add_table(rows=2, cols=2)
    fig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fig_tbl.autofit = False
    fig_tbl.columns[0].width = FIG_CELL_W
    fig_tbl.columns[1].width = FIG_CELL_W
    _no_borders(fig_tbl)

    for cell, path in [(fig_tbl.cell(0, 0), FIG1), (fig_tbl.cell(0, 1), FIG2)]:
        cell.width = FIG_CELL_W
        p_fig = cell.paragraphs[0]
        p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fig.paragraph_format.space_before = Pt(2)
        p_fig.paragraph_format.space_after = Pt(0)
        run = p_fig.add_run()
        run.add_picture(str(path), width=FIG_PICT_W)

    captions = [
        ("Figure 1.",
         " Streamlit supervision view (case C3, raw screenshot) — feeders 1–5 with media, "
         "mass flow rate, density, thermal expansion and feed position; extrusion parameters "
         "(screw speed, throughput, SME, residence time, fill factor, free volume); temperature "
         "profile zone 1 → die; AI compatibility score, stability probability and ranked "
         "recommendations."),
        ("Figure 2.",
         " Run-analysis view (raw screenshot) — applying the AI recommendation moves "
         "compatibility 46 → 78, p_stable 0.35 → 0.87, fill factor Z5 0.97 → 0.72 and estimated "
         "torque 84 → 62 %."),
    ]
    for i, (head, body) in enumerate(captions):
        cell = fig_tbl.cell(1, i)
        cell.width = FIG_CELL_W
        p_cap = cell.paragraphs[0]
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(0)
        _r(p_cap, head, bold=True, size=7, color=ACCENT)
        _r(p_cap, body, size=7, color=GREY)

    # ---- Discussion ----
    sec("Discussion & outlook")
    p = _par(doc, space_after=0)
    _r(p,
       "Results are encouraging on a limited 8-run industrial dataset; the agent's outputs remain "
       "advisory and final process decisions stay with the engineer. The rule-based formulation "
       "score will be replaced by a regression model trained on a literature-derived recipe "
       "corpus; the agent will be extended to NMC and sulfide solid electrolytes, with in-line "
       "torque and pressure closing the loop in v2. The combination ",
       size=9)
    _r(p, "extrusion + AI + battery formulation ",
       bold=True, size=9, color=ACCENT)
    _r(p,
       "addresses a documented scientific gap and supports a credible TRL 4–5 industrial path "
       "for Rondol Industrie.",
       size=9)

    p_foot = _par(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=2)
    _r(p_foot,
       "Acknowledgements: Rondol Industrie · Institut Jean Lamour · Campus ARTEM · "
       "Mastère Data & IA (RNCP 37137).",
       size=7, italic=True, color=GREY)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"DOCX written: {OUT_DOCX}")

    from docx2pdf import convert
    convert(str(OUT_DOCX), str(OUT_PDF))
    print(f"PDF written:  {OUT_PDF}")


if __name__ == "__main__":
    # Figures are already generated in reports/poster_abstract/figures/generated_v3_abstract/
    # — they are publication-ready, reuse them as-is. No regeneration needed.
    assert FIG1.exists(), f"missing {FIG1} — run build_abstract_v3.py first"
    assert FIG2.exists(), f"missing {FIG2} — run build_abstract_v3.py first"
    build_docx()
