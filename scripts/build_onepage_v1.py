"""
Poster une-page v1 — réaligné sur cas industriels C1->C6.

Sorties (NOUVELLES, n'écrasent rien) :
  reports/poster_abstract/Mbeumi_2026_ManagerAligned_OnePagePoster_v1.docx
  reports/poster_abstract/Mbeumi_2026_ManagerAligned_OnePagePoster_v1.pdf

Le script génère aussi à la volée :
  reports/poster_abstract/figures/generated_v3/v1_six_cases_strip.png
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets"
FIG_V2 = ROOT / "reports" / "poster_abstract" / "figures" / "generated_v2"
FIG_V3 = ROOT / "reports" / "poster_abstract" / "figures" / "generated_v3"
FIG_OUT = ROOT / "reports" / "poster_abstract" / "figures" / "generated_v3"
FIG_OUT.mkdir(parents=True, exist_ok=True)
OUT_DIR = ROOT / "reports" / "poster_abstract"
OUT_DOCX = OUT_DIR / "Mbeumi_2026_ManagerAligned_OnePagePoster_v1.docx"
OUT_PDF = OUT_DIR / "Mbeumi_2026_ManagerAligned_OnePagePoster_v1.pdf"

C_GREEN = RGBColor(0x0F, 0x6A, 0x3A)
C_BLUE = RGBColor(0x14, 0x49, 0x8B)
C_RED = RGBColor(0xB7, 0x2A, 0x2A)
C_ORANGE = RGBColor(0xC0, 0x6A, 0x10)
C_DARK = RGBColor(0x1A, 0x1A, 0x1A)
C_GREY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# 1) Figure six-case strip C1..C6 (matplotlib)
# ---------------------------------------------------------------------------
def build_six_cases_strip(out_path: Path) -> None:
    cases = [
        {
            "id": "C1", "label": "BASELINE", "color": "#0F6A3A",
            "head_text": "LFP 65 · LATP 17", "verdict": "stable",
            "lines": [("score", "65", "#0F6A3A"), ("p (RF)", "0.84", "#333333"), ("alert", "none", "#0F6A3A")],
        },
        {
            "id": "C2", "label": "OPTIMISED", "color": "#0F6A3A",
            "head_text": "Z4 +1 kneading · Z5 45° → 30°", "verdict": "stable",
            "lines": [("score", "82", "#0F6A3A"), ("p (RF)", "0.91", "#333333"), ("alert", "none", "#0F6A3A")],
        },
        {
            "id": "C3", "label": "AT RISK", "color": "#B72A2A",
            "head_text": "LATP 17 → 35 wt% · 180 rpm", "verdict": "unstable",
            "lines": [("score", "46", "#B72A2A"), ("p (RF)", "0.35", "#B72A2A"), ("alert", "Z5 overflow", "#B72A2A")],
        },
        {
            "id": "C4", "label": "AI RECO", "color": "#C06A10",
            "head_text": "4 ranked actions on C3", "verdict": "diagnostic",
            "lines": [("formulation", "LATP −18 pts", "#333333"),
                      ("screw", "Z4 45° → 30°", "#333333"),
                      ("process", "SME −15 % · Z5 +5 °C", "#333333")],
        },
        {
            "id": "C5", "label": "APPLIED", "color": "#0F6A3A",
            "head_text": "Reco appliquée → essai", "verdict": "stable",
            "lines": [("score", "78", "#0F6A3A"), ("p (RF)", "0.87", "#333333"), ("alert", "cleared", "#0F6A3A")],
        },
        {
            "id": "C6", "label": "LIMIT CASE", "color": "#7A5BAF",
            "head_text": "LATP 25 wt% (borderline)", "verdict": "borderline",
            "lines": [("score", "58", "#7A5BAF"),
                      ("p (RF)", "0.55", "#7A5BAF"),
                      ("note", "operator hand-off", "#7A5BAF")],
        },
    ]

    n = len(cases)
    fig_w = 11.2
    fig_h = 2.55
    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=200)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    title = ax.text(
        50, 96,
        "SIX-CASE INDUSTRIAL DEMONSTRATION  ·  baseline → optimised → risk → AI reco → applied → limit",
        ha="center", va="top", fontsize=10, weight="bold", color="#1A1A1A",
    )

    card_w = 100 / n - 1.2
    gap = 1.2

    for i, c in enumerate(cases):
        x0 = i * (card_w + gap) + 0.6
        y_top = 88
        head_h = 14
        body_h = 70

        # Head
        head = mpatches.FancyBboxPatch(
            (x0, y_top - head_h), card_w, head_h,
            boxstyle="round,pad=0.02,rounding_size=1.4",
            linewidth=0, facecolor=c["color"],
        )
        ax.add_patch(head)
        ax.text(x0 + 2.0, y_top - head_h / 2, c["id"],
                ha="left", va="center", fontsize=12, weight="bold", color="white")
        ax.text(x0 + card_w - 2.0, y_top - head_h / 2, c["label"],
                ha="right", va="center", fontsize=8, weight="bold", color="white")

        # Body
        body = mpatches.FancyBboxPatch(
            (x0, y_top - head_h - body_h), card_w, body_h,
            boxstyle="round,pad=0.02,rounding_size=1.4",
            linewidth=1.0, edgecolor=c["color"], facecolor="white",
        )
        ax.add_patch(body)

        # Head-text (formulation/changement)
        ax.text(x0 + card_w / 2, y_top - head_h - 4,
                c["head_text"], ha="center", va="top",
                fontsize=7.0, style="italic", color="#333333")

        # 3 KPI rows
        row_y_start = y_top - head_h - 14
        row_h = 14
        for k, (lbl, val, col) in enumerate(c["lines"]):
            yy = row_y_start - k * row_h
            ax.text(x0 + 3, yy - row_h / 2, lbl,
                    ha="left", va="center", fontsize=6.8, color="#555555")
            ax.text(x0 + card_w - 3, yy - row_h / 2, val,
                    ha="right", va="center", fontsize=8.0, weight="bold", color=col)

        # Verdict bar bas
        ax.text(x0 + card_w / 2, y_top - head_h - body_h + 4,
                c["verdict"].upper(),
                ha="center", va="center", fontsize=7.0, weight="bold",
                color=c["color"])

        # Arrow to next
        if i < n - 1:
            arr_x = x0 + card_w + gap / 2
            ax.annotate("", xy=(arr_x + 0.5, 55), xytext=(arr_x - 0.5, 55),
                        arrowprops=dict(arrowstyle="->", color="#777777", lw=1.2))

    fig.tight_layout(pad=0.2)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ---------------------------------------------------------------------------
# 2) DOCX helpers
# ---------------------------------------------------------------------------
def shade(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_borders(cell, color_hex: str = "BFBFBF", size: int = 4) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def no_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_margins(cell, top=30, bottom=30, left=60, right=60) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_run(par, text, *, bold=False, italic=False, size=7.8, color=C_DARK, font="Calibri"):
    r = par.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def set_par_spacing(par, before=0, after=0, line=1.05):
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def section_title(par, num, label, color, size=10.0):
    set_par_spacing(par, 0, 1, 1.0)
    add_run(par, f"{num}  ", bold=True, size=size, color=color)
    add_run(par, label.upper(), bold=True, size=size, color=color)


def bullet(cell, runs, after=0):
    p = cell.add_paragraph()
    set_par_spacing(p, 0, after, 1.05)
    add_run(p, "• ", bold=True, size=7.6, color=C_GREEN)
    for txt, kw in runs:
        kw.setdefault("size", 7.6)
        add_run(p, txt, **kw)
    return p


# ---------------------------------------------------------------------------
# 3) Build
# ---------------------------------------------------------------------------
def build() -> None:
    six_strip = FIG_OUT / "v1_six_cases_strip.png"
    build_six_cases_strip(six_strip)

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0.55)
    section.bottom_margin = Cm(0.55)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(8.0)

    # ===== HEADER (logo + title) =====
    hdr = doc.add_table(rows=1, cols=2)
    hdr.autofit = False
    set_cell_width(hdr.rows[0].cells[0], 3.2)
    set_cell_width(hdr.rows[0].cells[1], 16.4)
    logo_cell = hdr.rows[0].cells[0]
    title_cell = hdr.rows[0].cells[1]
    for c in (logo_cell, title_cell):
        no_borders(c); cell_margins(c, 0, 0, 10, 10)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_l = logo_cell.paragraphs[0]
    set_par_spacing(p_l, 0, 0, 1.0)
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if (ASSETS / "rondol_logo.png").exists():
        p_l.add_run().add_picture(str(ASSETS / "rondol_logo.png"), width=Cm(2.8))

    p_t = title_cell.paragraphs[0]
    set_par_spacing(p_t, 0, 0, 1.0)
    add_run(p_t,
            "AI-Assisted Twin-Screw Hot-Melt Extrusion for Lithium / Solid-State Battery Components",
            bold=True, size=12.0, color=C_DARK)
    p_t2 = title_cell.add_paragraph()
    set_par_spacing(p_t2, 0, 0, 1.0)
    add_run(p_t2, "An industrial AI decision-support tool — demonstrated on six real extrusion runs (C1→C6)",
            italic=True, size=8.0, color=C_GREY)
    p_t3 = title_cell.add_paragraph()
    set_par_spacing(p_t3, 0, 0, 1.0)
    add_run(p_t3, "Wilfried Galtier Mbeumi", bold=True, size=7.5, color=C_DARK)
    add_run(p_t3, "  ·  Rondol Industrie & Institut Jean Lamour (IJL), Campus ARTEM, Nancy — ",
            size=7.5, color=C_GREY)
    add_run(p_t3, "Supervisor: ", size=7.5, color=C_GREY)
    add_run(p_t3, "Maël Gallas", bold=True, size=7.5, color=C_DARK)
    add_run(p_t3, "  ·  Symposium IA + Extrusion + Batteries Li/SSB — May 2026  ·  Mastère Data & IA (RNCP 37137)",
            size=7.5, color=C_GREY)

    # ===== § 1 Contexte / Objectif (compact, 4 puces sur 1 ligne) =====
    p_ctx = doc.add_paragraph()
    set_par_spacing(p_ctx, 2, 0, 1.0)
    add_run(p_ctx, "1  CONTEXT & OBJECTIVE", bold=True, size=9.5, color=C_BLUE)

    ctx = doc.add_paragraph()
    set_par_spacing(ctx, 0, 1, 1.10)
    add_run(ctx, "Hot Melt Extrusion (HME) ", bold=True, size=7.8, color=C_DARK)
    add_run(ctx,
            "— solvent-free, continuous, mature in pharma — is being transferred to lithium-ion and "
            "solid-state battery (SSB) electrode / electrolyte manufacturing. The intersection ",
            size=7.8, color=C_DARK)
    add_run(ctx, "extrusion + AI + batteries", italic=True, bold=True, size=7.8, color=C_RED)
    add_run(ctx,
            " is a documented scientific gap (Drakopoulos 2021, Haarmann 2021, Kim 2023, Seeba 2024, Maia 2025). ",
            size=7.8, color=C_DARK)
    add_run(ctx, "Goal: ", bold=True, size=7.8, color=C_GREEN)
    add_run(ctx,
            "show that an AI decision-support tool (Streamlit + ML + rule-based score) can detect risk, "
            "diagnose, recommend and improve a lithium-bearing HME run on a real twin-screw line.",
            size=7.8, color=C_DARK)

    # ===== § 2 SIX-CASE INDUSTRIAL STORY (cœur du poster) =====
    p_s2 = doc.add_paragraph()
    set_par_spacing(p_s2, 3, 0, 1.0)
    add_run(p_s2, "2  INDUSTRIAL CASE STUDIES  —  C1 → C6", bold=True, size=10.0, color=C_BLUE)
    add_run(p_s2, "    ‧  formulation → process → screw → AI verdict → action",
            italic=True, size=7.5, color=C_GREY)

    p_strip = doc.add_paragraph()
    p_strip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(p_strip, 0, 0, 1.0)
    p_strip.add_run().add_picture(str(six_strip), width=Cm(19.4))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(cap, 0, 1, 1.0)
    add_run(cap, "Fig. 1  ", bold=True, size=7.0, color=C_BLUE)
    add_run(cap,
            "Six industrial runs covering the full agent loop. "
            "C1 baseline · C2 process optimisation · C3 detected risk (LATP 35 wt% → Z5 overflow alert) · "
            "C4 AI diagnostic + 4 ranked actions · C5 re-run after applying the recommendation · "
            "C6 borderline case (LATP 25 wt%) — operator hand-off, illustrates the model's uncertainty zone.",
            italic=True, size=7.0, color=C_GREY)

    # ===== § 3 ZOOM C3 (ALERTE) + C4 (RECO) + AVANT/APRÈS (3 colonnes) =====
    panels = doc.add_table(rows=1, cols=3)
    panels.autofit = False
    panels.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_width(panels.rows[0].cells[0], 6.4)
    set_cell_width(panels.rows[0].cells[1], 6.4)
    set_cell_width(panels.rows[0].cells[2], 6.7)
    cA, cB, cC = panels.rows[0].cells
    for c in (cA, cB, cC):
        no_borders(c); cell_margins(c, 20, 10, 30, 30)

    # ---- A : C3 supervision (alert)
    pA = cA.paragraphs[0]
    set_par_spacing(pA, 0, 1, 1.0)
    add_run(pA, "C3  ", bold=True, size=8.5, color=C_RED)
    add_run(pA, "AGENT SUPERVISION — alert raised", bold=True, size=7.6, color=C_DARK)
    pA_img = cA.add_paragraph()
    pA_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(pA_img, 0, 1, 1.0)
    img_a = FIG_V2 / "figv2_03_agent_alert_panel.png"
    if img_a.exists():
        pA_img.add_run().add_picture(str(img_a), width=Cm(6.0))
    pA_cap = cA.add_paragraph()
    set_par_spacing(pA_cap, 0, 0, 1.05)
    add_run(pA_cap, "Live KPI tiles : ", bold=True, size=7.0, color=C_DARK)
    add_run(pA_cap, "score 46/100 · p_stable=0.35 · fill Z5=0.97 · torque 84 %. "
                    "Per-zone risk strip flags Z5 overflow ; agent log traces ML + physics rules.",
            size=7.0, color=C_GREY)

    # ---- B : C4 reco panel
    pB = cB.paragraphs[0]
    set_par_spacing(pB, 0, 1, 1.0)
    add_run(pB, "C4  ", bold=True, size=8.5, color=C_ORANGE)
    add_run(pB, "AI RECOMMENDATION — 4 ranked actions", bold=True, size=7.6, color=C_DARK)
    pB_img = cB.add_paragraph()
    pB_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(pB_img, 0, 1, 1.0)
    img_b = FIG_V2 / "figv2_04_agent_recommendation_panel.png"
    if img_b.exists():
        pB_img.add_run().add_picture(str(img_b), width=Cm(6.0))
    pB_cap = cB.add_paragraph()
    set_par_spacing(pB_cap, 0, 0, 1.05)
    add_run(pB_cap, "Diagnostic + ", bold=True, size=7.0, color=C_DARK)
    add_run(pB_cap,
            "actions hiérarchisées : (1) LATP 35 → 17 wt% · (2) Z4 kneading 45° → 30° · "
            "(3) SME −15 % · (4) Z5 +5 °C. Projected gain +30 pts.",
            size=7.0, color=C_GREY)

    # ---- C : Before / after C3 → C5
    pC = cC.paragraphs[0]
    set_par_spacing(pC, 0, 1, 1.0)
    add_run(pC, "C3 → C5  ", bold=True, size=8.5, color=C_GREEN)
    add_run(pC, "BEFORE / AFTER applying the AI reco", bold=True, size=7.6, color=C_DARK)
    pC_img = cC.add_paragraph()
    pC_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(pC_img, 0, 1, 1.0)
    img_c = ROOT / "reports" / "poster_abstract" / "figures" / "generated" / "fig05_before_after_recommendation.png"
    if img_c.exists():
        pC_img.add_run().add_picture(str(img_c), width=Cm(6.3))
    # mini table Δ
    deltas = [
        ("Score /100",        "46",   "78",   "+32",  C_GREEN),
        ("p_stable (RF)",     "0.35", "0.87", "+0.52", C_GREEN),
        ("Fill factor Z5",    "0.97", "0.72", "−0.25", C_GREEN),
        ("Torque %",          "84",   "62",   "−22",   C_GREEN),
    ]
    dt = cC.add_table(rows=len(deltas), cols=4)
    for i, (k, b, a, d, col) in enumerate(deltas):
        cells = dt.rows[i].cells
        for cc in cells:
            set_borders(cc, "EAEAEA", 4); cell_margins(cc, 4, 4, 18, 18)
        cells[0].paragraphs[0].add_run(k).font.size = Pt(6.8)
        for cc, val, color in [
            (cells[1], b, C_RED),
            (cells[2], a, C_GREEN),
            (cells[3], d, col),
        ]:
            pp = cc.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_par_spacing(pp, 0, 0, 1.0)
            add_run(pp, val, bold=True, size=6.8, color=color)

    # ===== § 3b Honesty box : C6 limit case =====
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_width(box.rows[0].cells[0], 19.6)
    cb = box.rows[0].cells[0]
    cell_margins(cb, 20, 20, 50, 50)
    set_borders(cb, "7A5BAF", 6)
    shade(cb, "F4F0FA")
    p1 = cb.paragraphs[0]
    set_par_spacing(p1, 0, 0, 1.0)
    add_run(p1, "C6  HONEST LIMIT  —  ", bold=True, size=8.5, color=RGBColor(0x55, 0x3A, 0x90))
    add_run(p1,
            "borderline recipe LFP 65 / LATP 25 wt% (between C1 and C3) : score 58, p_stable 0.55. "
            "The rule-based formulation score and the RF classifier disagree → ",
            size=7.6, color=C_DARK)
    add_run(p1, "operator hand-off triggered, no automated green light. ",
            bold=True, size=7.6, color=RGBColor(0x55, 0x3A, 0x90))
    add_run(p1,
            "Reflects the ~5 % residual error on test (17/340 misclassified) and the ceiling of the "
            "5-criteria heuristic — documents the uncertainty zone instead of hiding it.",
            size=7.6, color=C_DARK)

    # ===== § 4 ML RESULTS (compact, 2 cells) =====
    p_s4 = doc.add_paragraph()
    set_par_spacing(p_s4, 3, 1, 1.0)
    add_run(p_s4, "3  ML RESULTS — stability classifier (window 60 s, n=627, GroupShuffleSplit by run_id)",
            bold=True, size=9.5, color=C_BLUE)

    ml = doc.add_table(rows=1, cols=2)
    ml.autofit = False
    ml.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_width(ml.rows[0].cells[0], 8.6)
    set_cell_width(ml.rows[0].cells[1], 10.9)
    mlL = ml.rows[0].cells[0]
    mlR = ml.rows[0].cells[1]
    for c in (mlL, mlR):
        no_borders(c); cell_margins(c, 10, 10, 20, 20)

    # mini-table KPI ML
    kpi = mlL.add_table(rows=4, cols=4)
    kpi.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Model", "Accuracy", "F1-macro", "ROC-AUC CV"]
    values = [
        ("Random Forest", "0.950", "0.917", "0.976 ± 0.021"),
        ("XGBoost",        "0.882", "0.827", "0.983 ± 0.020"),
        ("SVM (RBF)",      "0.953", "0.916", "0.957 ± 0.037"),
    ]
    for i, h in enumerate(headers):
        c = kpi.rows[0].cells[i]
        shade(c, "0F6A3A"); set_borders(c, "FFFFFF", 4); cell_margins(c, 6, 6, 30, 30)
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_par_spacing(pp, 0, 0, 1.0)
        add_run(pp, h, bold=True, size=7.2, color=RGBColor(0xFF, 0xFF, 0xFF))
    for r, vals in enumerate(values):
        row = kpi.rows[r + 1]
        for i, v in enumerate(vals):
            c = row.cells[i]; set_borders(c, "DADADA", 4); cell_margins(c, 6, 6, 30, 30)
            pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_par_spacing(pp, 0, 0, 1.0)
            bold = (r == 0)
            col = C_GREEN if r == 0 else C_DARK
            add_run(pp, v, bold=bold, size=7.2, color=col)
    cap_ml = mlL.add_paragraph()
    set_par_spacing(cap_ml, 1, 0, 1.0)
    add_run(cap_ml, "Random Forest", bold=True, size=7.0, color=C_GREEN)
    add_run(cap_ml,
            "  retenu en production (équilibre perf / interprétabilité). 8 runs (≥ 15 min) sur 11. "
            "RF test CM = [[54, 8], [9, 269]] sur 340 fenêtres.",
            size=7.0, color=C_GREY)

    # right : top features + insight
    p_tf = mlR.paragraphs[0]
    set_par_spacing(p_tf, 0, 1, 1.05)
    add_run(p_tf, "Top features (Gini, RF w=60 s) : ", bold=True, size=7.4, color=C_DARK)
    add_run(p_tf,
            "CastFilmP2_iqr 0.072 · DIE_std 0.067 · CastFilmBody_std 0.066 · CastFilmP2_std 0.052 · "
            "DIE_iqr 0.052 · CastFilmBody_iqr 0.051 · CastFilmP1_std 0.050.",
            size=7.4, color=C_DARK)
    p_insight = mlR.add_paragraph()
    set_par_spacing(p_insight, 0, 0, 1.05)
    add_run(p_insight, "Insight clé : ", bold=True, size=7.4, color=C_RED)
    add_run(p_insight,
            "la stabilité dérive ", size=7.4, color=C_DARK)
    add_run(p_insight, "aval (cast-film + DIE)", bold=True, size=7.4, color=C_DARK)
    add_run(p_insight,
            ", pas des consignes de zone. La qualité du fondu en sortie pilote la classification — "
            "cohérent avec la physique de l’extrusion réactive et confirme la valeur de la chaîne capteurs aval.",
            size=7.4, color=C_DARK)

    # ===== § 5 Discussion + Conclusion (2 col) =====
    foot = doc.add_table(rows=1, cols=2)
    foot.autofit = False
    foot.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_width(foot.rows[0].cells[0], 9.7)
    set_cell_width(foot.rows[0].cells[1], 9.9)
    fL = foot.rows[0].cells[0]; fR = foot.rows[0].cells[1]
    for c in (fL, fR):
        no_borders(c); cell_margins(c, 30, 10, 20, 20)

    section_title(fL.paragraphs[0], "4", "Discussion / Limites", C_BLUE, size=9.5)
    bullet(fL, [
        ("Dataset limité : ", dict(bold=True, color=C_DARK)),
        ("8 runs / 627 fenêtres ; généralisation à NMC et sulfures non encore testée.", dict(color=C_DARK)),
    ])
    bullet(fL, [
        ("Score formulation rule-based : ", dict(bold=True, color=C_DARK)),
        ("5 critères pondérés, à remplacer par régression entraînée sur ~50 recettes littérature.",
         dict(color=C_DARK)),
    ])
    bullet(fL, [
        ("Pas de validation électrochimique : ", dict(bold=True, color=C_DARK)),
        ("proxy thermique, abrasion simplifiée, pas de bouclage in-line des cellules battery-grade.",
         dict(color=C_DARK)),
    ])
    bullet(fL, [
        ("Zone d’incertitude assumée : ", dict(bold=True, color=C_DARK)),
        ("cas C6 montre la nécessité d’un opérateur en boucle ; l’IA est un copilot, pas un autopilote.",
         dict(color=C_DARK)),
    ])

    section_title(fR.paragraphs[0], "5", "Conclusion", C_GREEN, size=9.5)
    bullet(fR, [
        ("Démonstrateur industriel crédible : ", dict(bold=True, color=C_DARK)),
        ("agent IA testé sur 6 cas réels d’extrusion lithiée — boucle complète détection → diagnostic → reco → essai amélioré.",
         dict(color=C_DARK)),
    ])
    bullet(fR, [
        ("Aide décisionnelle opérateur : ", dict(bold=True, color=C_DARK)),
        ("Streamlit Supervision / Profile / Settings / Run Analysis, alertes & reco hiérarchisées en temps réel.",
         dict(color=C_DARK)),
    ])
    bullet(fR, [
        ("Réduction des essais itératifs : ", dict(bold=True, color=C_DARK)),
        ("le diagnostic C3 → C5 économise un cycle d’essai full-scale.",
         dict(color=C_DARK)),
    ])
    bullet(fR, [
        ("Potentiel scale-up Rondol 21 mm + SSB : ", dict(bold=True, color=C_GREEN)),
        ("contribution à un gap publié, voie stratégique extrusion + IA + batteries.",
         dict(color=C_DARK)),
    ])

    # ===== Footer : refs + ack =====
    refs = doc.add_paragraph()
    set_par_spacing(refs, 3, 0, 1.05)
    refs.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(refs, "References (sel.) : ", bold=True, size=7.0, color=C_GREY)
    add_run(refs,
            "Drakopoulos 2021 · Haarmann 2021 · Kim 2023 · Seeba 2024 · Daoudi 2024 · Maia 2025 · Wang 2025 · "
            "Fraunhofer IWS — DRYtraec®. ",
            italic=True, size=7.0, color=C_GREY)
    add_run(refs, "Acknowledgements : ", bold=True, size=7.0, color=C_GREY)
    add_run(refs,
            "Rondol Industrie ; Institut Jean Lamour (IJL) ; Campus ARTEM ; Mastère Data & IA (RNCP 37137). "
            "PFAS / PVDF — ECHA Feb. 2023.",
            italic=True, size=7.0, color=C_GREY)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"[OK] DOCX  : {OUT_DOCX}")

    try:
        from docx2pdf import convert
        convert(str(OUT_DOCX), str(OUT_PDF))
        print(f"[OK] PDF   : {OUT_PDF}")
    except Exception as exc:
        print(f"[WARN] PDF non généré ({exc!r})")


if __name__ == "__main__":
    build()
