# -*- coding: utf-8 -*-
"""Met à jour le sommaire (TOC) + champs via Word COM, exporte en PDF, puis QC."""
from __future__ import annotations
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
DOCX = REPORTS / "Memoire_Rondol_Wilfried_Galtier_MBEUMI_FINAL_PRO.docx"
PDF = REPORTS / "Memoire_Rondol_Wilfried_Galtier_MBEUMI_FINAL_PRO.pdf"

WD_PDF = 17
WD_FIELD_TOC = None


def convert():
    import win32com.client as win32
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(DOCX))
        # Met à jour tous les champs (PAGE, etc.)
        doc.Fields.Update()
        # Met à jour explicitement chaque table des matières
        for i in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents.Item(i).Update()
        doc.Repaginate()
        doc.Fields.Update()
        doc.Save()
        doc.ExportAsFixedFormat(OutputFileName=str(PDF), ExportFormat=WD_PDF,
                                OpenAfterExport=False)
        doc.Close(False)
        return True
    except Exception as exc:
        print(f"[ERREUR] Word COM: {exc}")
        return False
    finally:
        word.Quit()


def pages():
    try:
        from pypdf import PdfReader
        return len(PdfReader(str(PDF)).pages)
    except Exception as exc:
        print(f"[WARN] comptage pages: {exc}")
        return None


def qc():
    from docx import Document
    doc = Document(str(DOCX))
    full = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                full += "\n" + c.text
    n_images = len(doc.inline_shapes)
    checks = {
        "Auteur Wilfried Galtier MBEUMI": "Wilfried Galtier MBEUMI" in full,
        "Nexa Digital School": "Nexa Digital School" in full,
        "M. Moussa NDIAYE": "Moussa NDIAYE" in full,
        "M. Maël Gallas": "Maël Gallas" in full,
        "Rondol Industrie": "Rondol Industrie" in full,
        "Remerciements": "Remerciements" in full,
        "Résumé": "Résumé" in full,
        "Abstract": "Abstract" in full,
        "Sigles": "sigles" in full.lower(),
        "Sommaire": "Sommaire" in full,
        "Introduction générale": "Introduction générale" in full,
        "Partie 1": "Partie 1" in full,
        "Partie 8": "Partie 8" in full,
        "Conclusion générale": "Conclusion générale" in full,
        "Bibliographie": "Bibliographie" in full,
        "Annexes": "Annexe" in full,
        "Checklist RNCP (Annexe F)": "Annexe F" in full,
        "Légendes Figure N": "Figure 1" in full,
        "PAS de 'À COMPLÉTER'": "À COMPLÉTER" not in full,
        "PAS de 'INSÉRER'": "INSÉRER" not in full,
        "PAS de 'À INSÉRER'": "À INSÉRER" not in full,
        "PAS de 'pagination estimée'": "pagination estimée" not in full.lower(),
        "PAS de '##' résiduel": "##" not in full,
        "PAS de '**' résiduel": "**" not in full,
        "PAS de HTML brut": not re.search(r"<(div|br|table|tr|td|span)\b", full),
        "Images >= 25": n_images >= 25,
    }
    print("\n=== CONTRÔLE QUALITÉ ===")
    ok_all = True
    for k, ok in checks.items():
        print(f"  [{'OK ' if ok else 'KO '}] {k}")
        ok_all = ok_all and ok
    print(f"  Images inline détectées : {n_images}")
    print(f"  Paragraphes : {len(doc.paragraphs)} | Tableaux : {len(doc.tables)}")
    return ok_all, n_images


def main():
    if not DOCX.exists():
        print(f"[ERREUR] DOCX introuvable: {DOCX}"); return 1
    print("[1/3] Conversion PDF (Word COM, mise à jour du sommaire) ...")
    pdf_ok = convert()
    print(f"      PDF : {'OK ' + str(PDF) if pdf_ok and PDF.exists() else 'ÉCHEC'}")
    print("[2/3] Contrôle qualité ...")
    ok_all, n_images = qc()
    print("[3/3] Bilan")
    np = pages()
    print(f"\nDOCX : {DOCX if DOCX.exists() else 'ABSENT'}")
    print(f"PDF  : {PDF if PDF.exists() else 'ABSENT'}")
    print(f"Pages PDF : {np}")
    print(f"Qualité globale : {'CONFORME' if ok_all else 'À VÉRIFIER'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
