# -*- coding: utf-8 -*-
"""
Note de synthèse (1 page) pour présentation à l'encadrant : état des livrables
et indicateurs clés du projet. Document simple, séparé du mémoire.

Usage : python scripts/build_synthese_encadrant.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "NOTE_SYNTHESE_ENCADRANT.docx"
OUT_PDF = ROOT / "NOTE_SYNTHESE_ENCADRANT.pdf"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x26, 0x32, 0x3A)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
BASE_FONT = "Calibri"


def shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def borders(table, color="D9D9D9", sz="4"):
    tbl_pr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)
        b.append(el)
    tbl_pr.append(b)


def run(p, text, size=11, color=DARK, bold=False, italic=False, font=BASE_FONT):
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = font
    r.font.bold = bold; r.font.italic = italic
    return r


def h(doc, text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run(p, text, size=size, color=BLUE, bold=True)
    return p


def para(doc, text, size=10.5, color=DARK, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run(p, text, size=size, color=color, italic=italic)
    return p


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.6); s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(title, "Note de synthèse — État du projet", size=17, color=BLUE, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(14)
    run(sub, "Plateforme prédictive d'aide à la décision — extrusion bivis SSB (Rondol Industrie)",
        size=11, color=GREY, italic=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run(p, "Auteur : ", bold=True); run(p, "Wilfried Galtier MBEUMI    ")
    run(p, "Date : ", bold=True); run(p, "24 juillet 2026")

    # --- Statut des livrables ---
    h(doc, "1. Statut des livrables")
    rows = [
        ("Livrable", "Statut", "Détail"),
        ("Mémoire (MBEUMI_Wilfried_THESE.pdf)", "Prêt", "76 pages, conforme au guide RNCP (méthode, rétroplanning, budget, veille, risques, tableau de bord)"),
        ("Notebook d'analyse", "Prêt", "26 cellules exécutées : données, ML, moteur de règles expert (indépendant du modèle ML)"),
        ("Application (Streamlit Cloud)", "Prêt", "7 pages, authentification active, 705 tests automatisés passants"),
        ("Code source (ZIP)", "Prêt", "Périmètre strict app (220 fichiers, 6,3 Mo), dump SQL inclus"),
        ("Support de soutenance (PPTX/PDF)", "Prêt", "14 diapositives, méthode Kanban explicitée"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(tbl)
    tbl.columns[0].width = Cm(5.2); tbl.columns[1].width = Cm(2.2); tbl.columns[2].width = Cm(8.6)
    for j, val in enumerate(rows[0]):
        c = tbl.rows[0].cells[j]; shade(c, "1F4E79")
        rr = c.paragraphs[0].add_run(val); rr.font.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = WHITE; rr.font.name = BASE_FONT
    for i, row in enumerate(rows[1:], start=1):
        for j, val in enumerate(row):
            c = tbl.rows[i].cells[j]
            if j == 1:
                shade(c, "E4F3E9")
            rr = c.paragraphs[0].add_run(val); rr.font.size = Pt(9.5); rr.font.color.rgb = DARK; rr.font.name = BASE_FONT
            if j == 1:
                rr.font.bold = True; rr.font.color.rgb = GREEN

    # --- Indicateurs clés ---
    h(doc, "2. Indicateurs clés")
    kpis = [
        ("Tests automatisés", "705 / 705 passants (indépendants de l'ordre d'exécution)"),
        ("Modèle retenu", "RandomForest (fenêtre 60 s), F1-macro 0,918 ± 0,054 sur essai réel"),
        ("Validation externe", "AUC 0,753 sur jeu consolidé (3 479 fenêtres, hors entraînement)"),
        ("Jeu d'apprentissage", "798 fenêtres, 8 essais réels (campagne du 7 au 13 avril 2026)"),
        ("Pages applicatives", "7 (Supervision, Profile, Settings, Run Analysis, History, Process Engine, Compte)"),
    ]
    tbl2 = doc.add_table(rows=len(kpis), cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(tbl2, color="D9D9D9")
    tbl2.columns[0].width = Cm(4.5); tbl2.columns[1].width = Cm(11.5)
    for i, (k, v) in enumerate(kpis):
        c0, c1 = tbl2.rows[i].cells
        shade(c0, "F2F2F2")
        r0 = c0.paragraphs[0].add_run(k); r0.font.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = DARK; r0.font.name = BASE_FONT
        r1 = c1.paragraphs[0].add_run(v); r1.font.size = Pt(9.5); r1.font.color.rgb = DARK; r1.font.name = BASE_FONT

    # --- Nouveautés récentes ---
    h(doc, "3. Dernières mises à jour (session du 23–24 juillet 2026)")
    for item in [
        "Conformité guide RNCP vérifiée point par point (méthode Kanban explicitée, rétroplanning complet, cartographie des risques enrichie, RGPD/éthique mis à jour après l'ajout de l'authentification).",
        "Tableau de bord de suivi (§4.8) : remplacement d'une simple phrase par un vrai tableau d'indicateurs objectifs (Tableau 4.5).",
        "Page de garde du mémoire refaite avec un objet natif Word (page de garde intégrée), signature visuelle personnelle.",
        "Notebook d'analyse étoffé : démonstration exécutée que les recommandations proviennent d'un moteur de règles expert, distinct du modèle ML.",
        "Dépôt nettoyé et ZIP recentré sur le strict nécessaire à l'exécution de l'application (220 fichiers, 6,3 Mo).",
    ]:
        pb = doc.add_paragraph(style="List Bullet")
        pb.paragraph_format.space_after = Pt(4)
        run(pb, item, size=10)

    # --- Conclusion ---
    h(doc, "4. Conclusion")
    para(doc, "Les trois livrables (mémoire, notebook, application) sont synchronisés et cohérents "
              "(mêmes chiffres partout, vérifié par contrôle automatisé). Le projet est en état d'être "
              "présenté pour retour de l'encadrant avant la soutenance finale.", size=10.5)

    doc.save(str(OUT_DOCX))
    print("[OK] DOCX:", OUT_DOCX)

    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(str(OUT_DOCX))
        d.SaveAs(str(OUT_PDF), FileFormat=17)
        d.Close()
        word.Quit()
        print("[OK] PDF:", OUT_PDF)
    except Exception as exc:
        print("[WARN] Conversion PDF via Word COM impossible :", exc)


if __name__ == "__main__":
    raise SystemExit(main())
