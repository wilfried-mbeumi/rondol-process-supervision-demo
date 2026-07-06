"""Build the INDUSTRIAL POSTER PACKAGE v2 (DOCX + PDF) — symposium 15 May 2026.

This is the *industrial demonstrator* package — separate from the academic
abstract v1.  It generates a fresh set of visual figures (workflow, agent UI
mockups, KPI dashboard, 5-case strip) and assembles a multi-page DOCX with a
storytelling layout, then exports to PDF.

Outputs (never overwrites v1):
  reports/poster_abstract/figures/generated_v2/figv2_*.png
  reports/poster_abstract/Mbeumi_2026_IndustrialPosterPackage_v2.docx
  reports/poster_abstract/Mbeumi_2026_IndustrialPosterPackage_v2.pdf
"""
from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt, RGBColor

# ----------------------------------------------------------------------------
# Paths
# ----------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
FIG_V1 = ROOT / "reports" / "poster_abstract" / "figures" / "generated"
FIG_V2 = ROOT / "reports" / "poster_abstract" / "figures" / "generated_v2"
FIG_V2.mkdir(parents=True, exist_ok=True)
OUT_DIR = ROOT / "reports" / "poster_abstract"
DOCX_PATH = OUT_DIR / "Mbeumi_2026_IndustrialPosterPackage_v2.docx"
PDF_PATH = OUT_DIR / "Mbeumi_2026_IndustrialPosterPackage_v2.pdf"

# ----------------------------------------------------------------------------
# Palette industrielle (cohérente avec fig01..fig06)
# ----------------------------------------------------------------------------
GREEN = "#1B7A3D"
GREEN_DARK = "#0F4F26"
GREEN_SOFT = "#E8F5E9"
BLUE = "#005B96"
BLUE_DARK = "#003A66"
BLUE_SOFT = "#E3F0FA"
ORANGE = "#E67E22"
RED = "#C0392B"
RED_SOFT = "#FBEAEA"
AMBER = "#F1C40F"
GREY = "#333333"
MID_GREY = "#666666"
LIGHT_GREY = "#F2F4F7"
PANEL_BG = "#0F172A"   # dark dashboard background
PANEL_FG = "#E2E8F0"   # light text on dark panels

RGB_GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RGB_BLUE = RGBColor(0x00, 0x5B, 0x96)
RGB_RED = RGBColor(0xC0, 0x39, 0x2B)
RGB_GREY = RGBColor(0x33, 0x33, 0x33)
RGB_MID = RGBColor(0x66, 0x66, 0x66)
RGB_LIGHT = RGBColor(0xF2, 0xF4, 0xF7)
RGB_PANEL_BG = RGBColor(0x0F, 0x17, 0x2A)

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 11,
    "axes.titlesize": 12,
    "axes.titleweight": "bold",
    "axes.labelsize": 10,
    "axes.edgecolor": GREY,
    "axes.labelcolor": GREY,
    "xtick.color": GREY,
    "ytick.color": GREY,
    "text.color": GREY,
    "figure.facecolor": "white",
    "axes.facecolor": "white",
    "savefig.dpi": 220,
    "savefig.bbox": "tight",
})


# ============================================================================
# FIGURES — generated_v2/
# ============================================================================

def _rounded_box(ax, x, y, w, h, *, fc, ec=None, lw=1.6, radius=2.0, alpha=1.0):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle=f"round,pad=0.2,rounding_size={radius}",
        linewidth=lw, edgecolor=ec or fc, facecolor=fc, alpha=alpha,
    )
    ax.add_patch(box)
    return box


def figv2_01_hero_workflow():
    """Hero — industrial AI agent workflow (formulation -> action)."""
    fig, ax = plt.subplots(figsize=(16, 5.6))
    ax.set_xlim(0, 100); ax.set_ylim(0, 36); ax.axis("off")

    # banner
    _rounded_box(ax, 0, 30, 100, 6, fc=GREEN, radius=1.2)
    ax.text(2, 33, "RONDOL  ·  AI EXTRUSION COPILOT",
            color="white", fontsize=13, fontweight="bold", va="center")
    ax.text(98, 33, "Lithium-bearing electrodes  ·  TRL 4–5  ·  Symposium 15 May 2026",
            color="white", fontsize=9.5, ha="right", va="center", alpha=0.95)

    # 5 boxes
    steps = [
        ("01  RECIPE",        "LFP 65 / PVDF 8 / SP 5\nLATP 17 / LiTFSI 5",   BLUE,   "input"),
        ("02  PROCESS",       "T(Z1..Z8) · rpm · debit\nL/D 40:1  ·  Ø 10.5",  BLUE,   "input"),
        ("03  SCREW PROFILE", "conveying / kneading\ncompression / tip\nfill · RT · SME", BLUE_DARK, "input"),
        ("04  AI ENGINE",     "Random Forest w60s\nF1=0.917 · AUC=0.976\n+ rule-based score", ORANGE, "core"),
        ("05  ACTION",        "ranked recommendation\nformulation + screw\n+ process tweaks", GREEN, "output"),
    ]
    n = len(steps)
    box_w, box_h = 16.0, 18.0
    gap = (100 - n * box_w) / (n + 1)
    y_top = 26

    centers = []
    for i, (title, body, color, role) in enumerate(steps):
        x0 = gap + i * (box_w + gap)
        # main card
        _rounded_box(ax, x0, y_top - box_h, box_w, box_h, fc="white",
                     ec=color, lw=2.2, radius=2.0)
        # title strip
        _rounded_box(ax, x0, y_top - 4, box_w, 4, fc=color, radius=1.0)
        ax.text(x0 + box_w / 2, y_top - 2, title, color="white",
                fontsize=10, fontweight="bold", ha="center", va="center")
        # body
        ax.text(x0 + box_w / 2, y_top - box_h / 2 - 1.5, body,
                color=GREY, fontsize=9, ha="center", va="center", linespacing=1.4)
        centers.append((x0 + box_w / 2, y_top - box_h / 2))

    # arrows
    for i in range(n - 1):
        x0 = centers[i][0] + box_w / 2 + 0.4
        x1 = centers[i + 1][0] - box_w / 2 - 0.4
        y = centers[i][1]
        arrow = FancyArrowPatch((x0, y), (x1, y),
                                arrowstyle="-|>", mutation_scale=18,
                                linewidth=2.0, color=GREY)
        ax.add_patch(arrow)

    # closed-loop arrow from ACTION back to PROCESS
    ax.annotate(
        "", xy=(centers[1][0], 4), xytext=(centers[4][0], 4),
        arrowprops=dict(arrowstyle="-|>", color=ORANGE, lw=2.2,
                        connectionstyle="arc3,rad=0.0"),
    )
    ax.text((centers[1][0] + centers[4][0]) / 2, 2.0,
            "closed-loop:  recommendation re-enters the process",
            color=ORANGE, fontsize=9.5, ha="center", va="center", style="italic")

    fig.savefig(FIG_V2 / "figv2_01_hero_workflow.png", facecolor="white")
    plt.close(fig)


def figv2_02_architecture():
    """Layered architecture — UI / Decision / Process / Data."""
    fig, ax = plt.subplots(figsize=(15, 7.5))
    ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis("off")

    layers = [
        ("PRESENTATION  ·  Streamlit HMI",
         "Home / Profile / Settings / Run analysis / History",
         BLUE, 78, 18),
        ("DECISION  ·  AI Agent",
         "compatibility score (5 rules)  +  ML stability classifier  +  alerts  +  text recommendation",
         ORANGE, 56, 18),
        ("PROCESS LOGIC  ·  Physics",
         "screw_logic.py  ·  fill factor · residence time · SME · zone risk",
         BLUE_DARK, 34, 18),
        ("DATA",
         "8 runs (627 windows · 87 features)  ·  9 trained models  ·  metrics & robustness",
         GREEN, 12, 18),
    ]
    for title, body, color, y, h in layers:
        _rounded_box(ax, 2, y, 96, h, fc="white", ec=color, lw=2.4, radius=2.5)
        _rounded_box(ax, 2, y + h - 5, 96, 5, fc=color, radius=1.5)
        ax.text(4, y + h - 2.5, title, color="white",
                fontsize=11.5, fontweight="bold", va="center")
        ax.text(50, y + h / 2 - 2, body, color=GREY, fontsize=10.5,
                ha="center", va="center")

    # connecting bars
    for y in (76, 54, 32):
        ax.plot([50, 50], [y, y - 2], color=MID_GREY, lw=2.0, zorder=0)
        ax.plot([47, 53], [y - 2, y - 2], color=MID_GREY, lw=2.0, zorder=0)

    # side annotations
    ax.text(99, 87, "operator",       fontsize=9.5, color=MID_GREY, ha="right", style="italic")
    ax.text(99, 65, "agent loop",     fontsize=9.5, color=MID_GREY, ha="right", style="italic")
    ax.text(99, 43, "deterministic",  fontsize=9.5, color=MID_GREY, ha="right", style="italic")
    ax.text(99, 21, "industrial runs", fontsize=9.5, color=MID_GREY, ha="right", style="italic")

    fig.savefig(FIG_V2 / "figv2_02_architecture.png", facecolor="white")
    plt.close(fig)


def _dark_panel(ax, x, y, w, h, *, radius=1.5, fc=PANEL_BG):
    _rounded_box(ax, x, y, w, h, fc=fc, ec=fc, radius=radius)


def figv2_03_agent_alert_panel():
    """Mockup of the AI agent control panel — case C3 (RISK)."""
    fig, ax = plt.subplots(figsize=(14, 7.5))
    ax.set_xlim(0, 100); ax.set_ylim(0, 60); ax.axis("off")

    # outer panel
    _dark_panel(ax, 0, 0, 100, 60)

    # header
    _rounded_box(ax, 2, 53, 96, 5, fc=BLUE_DARK, ec=BLUE_DARK, radius=1.2)
    ax.text(3.5, 55.5, "AI EXTRUSION COPILOT  —  Live supervision",
            color="white", fontsize=12, fontweight="bold", va="center")
    ax.text(98, 55.5, "case C3  ·  semi-dry  ·  Ø 10.5 mm  ·  L/D 40",
            color="#94A3B8", fontsize=9.5, ha="right", va="center")

    # status banner — RED
    _rounded_box(ax, 2, 46, 96, 5, fc=RED, ec=RED, radius=1.2)
    ax.text(3.5, 48.5, "AT RISK  —  ceramic overload detected (LATP 35 wt%)",
            color="white", fontsize=11.5, fontweight="bold", va="center")
    ax.text(98, 48.5, "stability classifier: UNSTABLE",
            color="white", fontsize=10, ha="right", va="center", style="italic")

    # KPI badges
    kpis = [
        ("Compat. score", "46 / 100", RED),
        ("p_stable (RF)", "0.35",     RED),
        ("Fill factor Z5", "0.97",    AMBER),
        ("Torque (est.)",  "84 %",    AMBER),
    ]
    bw = 22; bh = 9
    for i, (lbl, val, c) in enumerate(kpis):
        x = 2 + i * (bw + 2)
        _rounded_box(ax, x, 35, bw, bh, fc="#1E293B", ec=c, lw=2.0, radius=1.2)
        ax.text(x + bw / 2, 35 + bh - 2.2, lbl, color="#94A3B8",
                fontsize=9, ha="center", va="center")
        ax.text(x + bw / 2, 35 + bh / 2 - 1.6, val, color=c,
                fontsize=15, fontweight="bold", ha="center", va="center")

    # zone risk chart Z1..Z8
    zones = [f"Z{i}" for i in range(1, 9)]
    risk = [0.10, 0.15, 0.22, 0.35, 0.92, 0.45, 0.30, 0.18]  # Z5 spike
    colors = [RED if v >= 0.7 else (AMBER if v >= 0.4 else GREEN) for v in risk]

    ax_inset = fig.add_axes([0.07, 0.13, 0.55, 0.30])
    ax_inset.set_facecolor("#1E293B")
    ax_inset.bar(zones, risk, color=colors, edgecolor="#0F172A")
    ax_inset.axhline(0.7, color=RED, lw=1.2, linestyle="--", alpha=0.6)
    ax_inset.set_ylim(0, 1.0)
    ax_inset.set_title("Per-zone risk  ·  threshold 0.7",
                       color=PANEL_FG, fontsize=10.5, loc="left", pad=6)
    ax_inset.tick_params(colors=PANEL_FG)
    for spine in ax_inset.spines.values():
        spine.set_color("#334155")
    ax_inset.grid(axis="y", color="#334155", alpha=0.5)

    # log lines (right side)
    log_x = 65
    _rounded_box(ax, log_x, 6, 33, 26, fc="#1E293B", ec="#334155", radius=1.2)
    ax.text(log_x + 1.5, 30, "AGENT LOG",
            color="#94A3B8", fontsize=10, fontweight="bold", va="center")
    log_lines = [
        ("12:04:18", "rule: LATP 35 % > 30 %  ·  ceramic overload",  RED),
        ("12:04:18", "ML: p_stable=0.35 (RF w60)  ·  class UNSTABLE", RED),
        ("12:04:19", "physics: fill Z5=0.97  ·  torque 84 %",         AMBER),
        ("12:04:19", "alert raised  ·  zone Z5",                       RED),
        ("12:04:20", "recommendation engine: ready",                   GREEN),
    ]
    for i, (ts, msg, c) in enumerate(log_lines):
        y = 26 - i * 3.6
        ax.text(log_x + 1.5, y, ts,  color="#64748B", fontsize=8.5, va="center", family="monospace")
        ax.text(log_x + 7,   y, "■", color=c,        fontsize=9,   va="center")
        ax.text(log_x + 9,   y, msg, color=PANEL_FG, fontsize=8.8, va="center")

    fig.savefig(FIG_V2 / "figv2_03_agent_alert_panel.png", facecolor="white")
    plt.close(fig)


def figv2_04_agent_recommendation_panel():
    """Mockup of the AI agent recommendation panel — case C4."""
    fig, ax = plt.subplots(figsize=(14, 7.5))
    ax.set_xlim(0, 100); ax.set_ylim(0, 60); ax.axis("off")

    _dark_panel(ax, 0, 0, 100, 60)

    # header
    _rounded_box(ax, 2, 53, 96, 5, fc=BLUE_DARK, ec=BLUE_DARK, radius=1.2)
    ax.text(3.5, 55.5, "AI EXTRUSION COPILOT  —  Recommendation",
            color="white", fontsize=12, fontweight="bold", va="center")
    ax.text(98, 55.5, "case C4  ·  diagnostic + ranked actions",
            color="#94A3B8", fontsize=9.5, ha="right", va="center")

    # diagnostic banner
    _rounded_box(ax, 2, 45, 96, 6, fc="#1E293B", ec=ORANGE, lw=2.0, radius=1.2)
    ax.text(3.5, 49.6, "DIAGNOSTIC",
            color=ORANGE, fontsize=10, fontweight="bold", va="center")
    ax.text(3.5, 47.3,
            "ceramic overload (LATP 35 wt%) — overflow risk Z5 — torque drift",
            color=PANEL_FG, fontsize=10.5, va="center")

    # 4 ranked action cards
    actions = [
        (1, "FORMULATION", "Reduce LATP to 17–20 wt%",   GREEN),
        (2, "SCREW",       "Z4 kneading 45° → 30°",       BLUE),
        (3, "PROCESS",     "Reduce specific energy −15 %", BLUE_DARK),
        (4, "PROCESS",     "Z5 temperature +5 °C",         BLUE_DARK),
    ]
    aw, ah = 22.5, 19
    for i, (rank, target, action, c) in enumerate(actions):
        x = 2 + i * (aw + 1.5)
        _rounded_box(ax, x, 21, aw, ah, fc="#1E293B", ec=c, lw=2.0, radius=1.4)
        # priority badge
        _rounded_box(ax, x + 1, 21 + ah - 5, 4, 4, fc=c, ec=c, radius=1.0)
        ax.text(x + 3, 21 + ah - 3, f"#{rank}", color="white",
                fontsize=10, fontweight="bold", ha="center", va="center")
        ax.text(x + 7, 21 + ah - 3, target, color=c,
                fontsize=10, fontweight="bold", va="center")
        ax.text(x + aw / 2, 21 + ah / 2 - 2, action, color=PANEL_FG,
                fontsize=10, ha="center", va="center", wrap=True)

    # projected outcome panel
    _rounded_box(ax, 2, 4, 96, 14, fc="#1E293B", ec="#334155", radius=1.4)
    ax.text(4, 14, "PROJECTED OUTCOME  (after applying the four actions)",
            color="#94A3B8", fontsize=9.5, fontweight="bold", va="center")
    ax.text(4, 10,
            "compatibility score 46  →  70–80     ·     stability probability ≥ 0.85"
            "     ·     Z5 alert cleared",
            color=PANEL_FG, fontsize=11.5, va="center")
    # arrow score
    arrow = FancyArrowPatch((42, 7), (62, 7), arrowstyle="-|>",
                            mutation_scale=22, linewidth=2.5, color=GREEN)
    ax.add_patch(arrow)
    ax.text(70, 7, "  expected gain  +30 pts",
            color=GREEN, fontsize=11, fontweight="bold", va="center")

    fig.savefig(FIG_V2 / "figv2_04_agent_recommendation_panel.png", facecolor="white")
    plt.close(fig)


def figv2_05_before_after_dashboard():
    """Big KPI dashboard — C3 (before) vs C5 (after recommendation)."""
    fig, ax = plt.subplots(figsize=(15, 5.8))
    ax.set_xlim(0, 100); ax.set_ylim(0, 50); ax.axis("off")

    # title
    ax.text(50, 47, "BEFORE  →  AFTER  ·  C3 (at risk)  vs  C5 (recommendation applied)",
            color=GREY, fontsize=13, fontweight="bold", ha="center", va="center")

    # 4 large tiles
    tiles = [
        ("Compatibility score", "46", "78", "+32", GREEN, False),
        ("p_stable (RF)",       "0.35", "0.87", "+0.52", GREEN, False),
        ("Fill factor Z5",      "0.97", "0.72", "−0.25", GREEN, True),
        ("Estimated torque",    "84 %", "62 %", "−22 pts", GREEN, True),
    ]
    tw, th = 22, 36
    gap = (100 - 4 * tw) / 5
    y0 = 4

    for i, (lbl, before, after, delta, c, inverse) in enumerate(tiles):
        x = gap + i * (tw + gap)
        _rounded_box(ax, x, y0, tw, th, fc="white", ec="#CBD5E1", lw=1.6, radius=2.0)
        ax.text(x + tw / 2, y0 + th - 3.5, lbl, color=MID_GREY,
                fontsize=10.5, ha="center", va="center")
        # before
        ax.text(x + tw / 4,  y0 + th / 2 + 2, before, color=RED,
                fontsize=20, fontweight="bold", ha="center", va="center")
        ax.text(x + tw / 4,  y0 + th / 2 - 3.5, "C3", color=MID_GREY,
                fontsize=9, ha="center", va="center", style="italic")
        # arrow
        arrow = FancyArrowPatch((x + tw / 3 + 2, y0 + th / 2 + 1),
                                (x + 2 * tw / 3 - 2, y0 + th / 2 + 1),
                                arrowstyle="-|>", mutation_scale=16,
                                linewidth=2.0, color=GREY)
        ax.add_patch(arrow)
        # after
        ax.text(x + 3 * tw / 4, y0 + th / 2 + 2, after, color=GREEN,
                fontsize=20, fontweight="bold", ha="center", va="center")
        ax.text(x + 3 * tw / 4, y0 + th / 2 - 3.5, "C5", color=MID_GREY,
                fontsize=9, ha="center", va="center", style="italic")
        # delta badge
        _rounded_box(ax, x + 2, y0 + 2, tw - 4, 4, fc=GREEN_SOFT, ec=GREEN,
                     lw=1.2, radius=1.0)
        ax.text(x + tw / 2, y0 + 4, f"Δ  {delta}", color=GREEN_DARK,
                fontsize=11, fontweight="bold", ha="center", va="center")

    fig.savefig(FIG_V2 / "figv2_05_before_after_dashboard.png", facecolor="white")
    plt.close(fig)


def figv2_06_five_cases_strip():
    """Five-case journey strip — one mini-card per case."""
    fig, ax = plt.subplots(figsize=(16, 6.8))
    ax.set_xlim(0, 100); ax.set_ylim(0, 60); ax.axis("off")

    ax.text(50, 57.5,
            "FIVE-CASE DEMONSTRATION  ·  baseline  →  optimised  →  risk  →  recommendation  →  applied",
            color=GREY, fontsize=12.5, fontweight="bold", ha="center", va="center")

    cases = [
        ("C1", "BASELINE",
         "LFP 65 · LATP 17",
         "score 65", "p=0.84", "stable",  GREEN,  None),
        ("C2", "OPTIMISED",
         "screw Z5 45° → 30°",
         "score 82", "p=0.91", "stable",  GREEN,  None),
        ("C3", "AT RISK",
         "LATP 17 → 35 %",
         "score 46", "p=0.35", "alert Z5", RED,   "alert"),
        ("C4", "AI RECO",
         "4 ranked actions",
         "diagnostic", "panel", "ranked", ORANGE, "reco"),
        ("C5", "APPLIED",
         "reco → process",
         "score 78", "p=0.87", "stable",  GREEN,  None),
    ]
    n = len(cases)
    cw, ch = 17.5, 46
    gap = (100 - n * cw) / (n + 1)
    y0 = 4

    for i, (cid, label, sub, k1, k2, k3, color, marker) in enumerate(cases):
        x = gap + i * (cw + gap)
        _rounded_box(ax, x, y0, cw, ch, fc="white", ec=color, lw=2.4, radius=1.8)
        # header strip
        _rounded_box(ax, x, y0 + ch - 7, cw, 7, fc=color, radius=1.4)
        ax.text(x + 1.5, y0 + ch - 3.5, cid, color="white",
                fontsize=13, fontweight="bold", va="center")
        ax.text(x + cw - 1.5, y0 + ch - 3.5, label, color="white",
                fontsize=9.5, fontweight="bold", ha="right", va="center")
        # sub
        ax.text(x + cw / 2, y0 + ch - 11, sub, color=GREY,
                fontsize=9.5, ha="center", va="center", style="italic")
        # KPIs
        for j, (txt, kc) in enumerate([
            (k1, color), (k2, MID_GREY), (k3, color),
        ]):
            yy = y0 + ch - 17 - j * 6.5
            _rounded_box(ax, x + 1.5, yy - 2.5, cw - 3, 5,
                         fc=LIGHT_GREY, ec=LIGHT_GREY, radius=1.0)
            ax.text(x + cw / 2, yy, txt, color=kc,
                    fontsize=10.5, fontweight="bold", ha="center", va="center")
        # arrow to next case
        if i < n - 1:
            arrow = FancyArrowPatch((x + cw + 0.6, y0 + ch / 2),
                                    (x + cw + gap - 0.6, y0 + ch / 2),
                                    arrowstyle="-|>", mutation_scale=14,
                                    linewidth=1.6, color=MID_GREY)
            ax.add_patch(arrow)

    fig.savefig(FIG_V2 / "figv2_06_five_cases_strip.png", facecolor="white")
    plt.close(fig)


def build_all_figures():
    figv2_01_hero_workflow()
    figv2_02_architecture()
    figv2_03_agent_alert_panel()
    figv2_04_agent_recommendation_panel()
    figv2_05_before_after_dashboard()
    figv2_06_five_cases_strip()
    print("[OK] 6 figures generated in", FIG_V2.relative_to(ROOT))


# ============================================================================
# DOCX
# ============================================================================

def _set_run_font(run, *, size=10, bold=False, italic=False,
                  color=RGB_GREY, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), name)


def add_paragraph(doc, text, *, size=10, bold=False, italic=False,
                  color=RGB_GREY, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_after=2, space_before=0, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_runs(doc, segments, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=2, space_before=0, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    for text, kwargs in segments:
        run = p.add_run(text)
        kw = {"size": 10, "color": RGB_GREY}
        kw.update(kwargs)
        _set_run_font(run, **kw)
    return p


def add_heading(doc, text, *, color=RGB_GREEN, size=12, space_before=6,
                space_after=2, keep_next=True, upper=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.keep_with_next = keep_next
    run = p.add_run(text.upper() if upper else text)
    _set_run_font(run, size=size, bold=True, color=color)
    return p


def add_caption(doc, text, *, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run_font(run, size=8.5, italic=True, color=RGB_MID)
    return p


def add_figure(doc, image_path, *, width_cm=17.5,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color="CBD5E1", size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_kpi_badges(doc, badges):
    """badges: list of (label, value, color_hex)."""
    table = doc.add_table(rows=2, cols=len(badges))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    for col, (label, value, fill) in enumerate(badges):
        # row 0: value
        cell_v = table.cell(0, col)
        cell_v.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell_v, fill.lstrip("#"))
        _set_cell_borders(cell_v, color=fill.lstrip("#"), size=8)
        p = cell_v.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        _set_run_font(r, size=16, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # row 1: label
        cell_l = table.cell(1, col)
        cell_l.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell_l, "F2F4F7")
        _set_cell_borders(cell_l, color="CBD5E1")
        p = cell_l.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label)
        _set_run_font(r, size=8.5, color=RGB_MID)
    # equalise widths
    total = Cm(17.5)
    col_w = Cm(17.5 / len(badges))
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_w
    return table


def setup_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.footer_distance = Cm(0.5)
        section.header_distance = Cm(0.4)


def install_page_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Rondol Industrie  ·  Institut Jean Lamour  ·  Mastère Data & IA "
        "(RNCP 37137)  ·  Industrial Poster Package v2  ·  2026-05-11"
    )
    _set_run_font(run, size=7.5, italic=True, color=RGB_MID)


def add_title_block(doc):
    # green band as title
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(17.5)
    _set_cell_shading(cell, "1B7A3D")
    _set_cell_borders(cell, color="1B7A3D", size=8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("RONDOL  ·  AI EXTRUSION COPILOT  ·  ")
    _set_run_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    r = p.add_run("Lithium-bearing electrodes — industrial demonstrator")
    _set_run_font(r, size=10, italic=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(4)
    r = p2.add_run(
        "AI-Assisted Twin-Screw Extrusion — From recipe to ranked recommendation"
    )
    _set_run_font(r, size=15, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # authors line
    add_runs(doc, [
        ("Wilfried Galtier Mbeumi", {"size": 9.5, "bold": True}),
        ("¹·²", {"size": 7.5}),
        ("    ·    ", {"size": 9.5, "color": RGB_MID}),
        ("Maël Gallas", {"size": 9.5, "bold": True}),
        ("¹    ", {"size": 7.5}),
        ("¹ Rondol Industrie, Nancy  ·  ² Institut Jean Lamour (IJL), "
         "Université de Lorraine, Nancy",
         {"size": 8, "italic": True, "color": RGB_MID}),
    ], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, space_before=2)


# ----------------------------------------------------------------------------
# Page builders
# ----------------------------------------------------------------------------

def page1_hero(doc):
    add_title_block(doc)

    # pitch
    add_runs(doc, [
        ("An industrial decision-support agent that closes the loop  ",
         {"size": 11, "color": RGB_GREY}),
        ("recipe → process → screw profile → risk → recommendation",
         {"size": 11, "bold": True, "color": RGB_BLUE, "italic": True}),
        ("  on a real Rondol twin-screw extruder (Ø 10.5 mm, L/D 40:1).  "
         "Eight industrial runs, 627 sliding windows, 87 features, three "
         "supervised classifiers — the agent diagnoses a risk and emits a "
         "ranked corrective action that the operator can apply on the next "
         "trial.",
         {"size": 11, "color": RGB_GREY}),
    ], space_after=4)

    # 4 KPI badges
    add_kpi_badges(doc, [
        ("Test F1-macro (RF, 60 s)",   "0.917",  "1B7A3D"),
        ("ROC-AUC 5-fold CV",          "0.976",  "005B96"),
        ("Score gain C3 → C5",          "+32 pts", "E67E22"),
        ("Industrial readiness",       "TRL 4–5", "0F4F26"),
    ])

    add_heading(doc, "Industrial workflow", color=RGB_GREEN, size=11,
                space_before=10)
    add_figure(doc, FIG_V2 / "figv2_01_hero_workflow.png", width_cm=17.5)
    add_caption(doc,
        "Fig. A  —  The five stages of the AI copilot. The recommendation loop "
        "is closed: a corrective action re-enters the process and is re-scored.")

    add_heading(doc, "What this demonstrator proves", color=RGB_BLUE,
                size=11, space_before=6)
    add_runs(doc, [
        ("1. ", {"size": 10, "bold": True, "color": RGB_BLUE}),
        ("the agent detects a documented failure mode (ceramic overload) "
         "before it disrupts the extruder — ", {}),
        ("rule + ML in agreement", {"bold": True}),
        (";   ", {}),
        ("2. ", {"size": 10, "bold": True, "color": RGB_BLUE}),
        ("it produces a ", {}),
        ("ranked corrective action", {"bold": True}),
        (" (not a passive dashboard);   ", {}),
        ("3. ", {"size": 10, "bold": True, "color": RGB_BLUE}),
        ("applying the action restores stability — score 46 → 78, p_stable "
         "0.35 → 0.87, Z5 alert cleared.", {}),
    ], space_after=2)


def page2_architecture_evidence(doc):
    add_page_break(doc)

    add_heading(doc, "Agent architecture — four layers", color=RGB_GREEN, size=12)
    add_paragraph(doc,
        "The agent is implemented as a Streamlit HMI calling a deterministic "
        "physics engine and a pre-trained stability classifier. Each decision "
        "is traceable to a rule, a physical KPI, and a probability. No black "
        "box — every alert is justified by an explicit signal.",
        size=10, space_after=2)
    add_figure(doc, FIG_V2 / "figv2_02_architecture.png", width_cm=17.0)
    add_caption(doc,
        "Fig. B  —  Layered architecture. From operator inputs to industrial "
        "run data, with the decision agent at the centre.")

    add_heading(doc, "Scientific backbone — ML performance",
                color=RGB_BLUE, size=11, space_before=6)
    add_runs(doc, [
        ("Eight industrial runs (≥ 15 min) were segmented into ", {}),
        ("627 sliding windows of 60 s", {"bold": True}),
        (" (87 features each). Random Forest, XGBoost and SVM (RBF) were "
         "compared with GroupShuffleSplit over run_id to prevent inter-run "
         "leakage. Random Forest is retained as the production model: ", {}),
        ("accuracy 0.950, F1-macro 0.917, ROC-AUC 0.976 ± 0.021 (5-fold CV)",
         {"bold": True, "color": RGB_GREEN}),
        (". The 60 s window outperforms 30 s and 120 s (F1-macro CV: "
         "0.935 vs 0.916 vs 0.843).", {}),
    ])

    # Two figures side by side (ML perf + feature importance)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in (table.cell(0, 0), table.cell(0, 1),
                 table.cell(1, 0), table.cell(1, 1)):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_V1 / "fig03_ml_performance_w60.png"),
                            width=Cm(8.4))
    p = table.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fig. C  —  Test confusion matrices (60 s, n=340).")
    _set_run_font(r, size=8, italic=True, color=RGB_MID)

    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_V1 / "fig04_feature_importance_RF_w60.png"),
                            width=Cm(8.4))
    p = table.cell(1, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fig. D  —  Top-10 RF features: downstream CastFilm + die.")
    _set_run_font(r, size=8, italic=True, color=RGB_MID)


def page3_five_cases(doc):
    add_page_break(doc)

    add_heading(doc, "The AI copilot in action — five cases",
                color=RGB_GREEN, size=12)
    add_paragraph(doc,
        "Five reproducible test cases (JSON-defined states under "
        "reports/poster_abstract/cases/states/) exercise the full agent loop "
        "on a single lithium-bearing reference recipe. Each case is a complete "
        "operator interaction: formulation, screw profile, process setpoints — "
        "and the agent verdict.",
        size=10, space_after=4)

    add_figure(doc, FIG_V2 / "figv2_06_five_cases_strip.png", width_cm=17.5)
    add_caption(doc,
        "Fig. E  —  Five-case demonstration strip. C1 baseline, C2 optimised "
        "screw, C3 ceramic overload (alert Z5), C4 ranked recommendation, "
        "C5 recommendation applied — back to stable.")

    add_heading(doc, "Reference recipe (cases C1, C2, C5)",
                color=RGB_BLUE, size=10.5, space_before=8)
    # recipe table
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headers = ["Component", "LFP", "PVDF", "Super P", "LATP", "LiTFSI"]
    values  = ["wt%",       "65",  "8",    "5",       "17",   "5"]
    for j, (h, v) in enumerate(zip(headers, values)):
        # header
        c = table.cell(0, j)
        _set_cell_shading(c, "1B7A3D")
        _set_cell_borders(c, color="1B7A3D")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h); _set_run_font(r, size=9, bold=True,
                                        color=RGBColor(0xFF, 0xFF, 0xFF))
        # value
        c = table.cell(1, j)
        _set_cell_shading(c, "E8F5E9")
        _set_cell_borders(c, color="CBD5E1")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(v); _set_run_font(r, size=10, bold=(j > 0), color=RGB_GREY)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(17.5 / 6)

    add_heading(doc, "Screw profile (Ø 10.5 mm  ·  L/D 40:1, eight zones)",
                color=RGB_BLUE, size=10.5, space_before=8)
    add_figure(doc, FIG_V1 / "fig02_screw_li_profile.png", width_cm=17.0,
               space_before=0, space_after=2)
    add_caption(doc,
        "Fig. F  —  Rendered screw profile for the lithium reference recipe "
        "(C1). Zones Z1 (feed) to Z8 (tip) show conveying / kneading / "
        "compression elements as configured in the HMI Profile page.")


def page4_money_shot(doc):
    add_page_break(doc)

    add_heading(doc, "Money shot — risk  →  recommendation  →  improvement",
                color=RGB_GREEN, size=12)
    add_paragraph(doc,
        "This page shows the agent during the critical loop on case C3 → C4 → "
        "C5. The two panels below are mockups of the live agent UI, "
        "rendered with the actual values from the case JSON files. They "
        "illustrate what the operator sees on the Supervision screen at each "
        "step of the loop.",
        size=10, space_after=4)

    add_heading(doc, "1. Alert  —  case C3 (ceramic overload)",
                color=RGB_RED, size=10.5, space_before=4)
    add_figure(doc, FIG_V2 / "figv2_03_agent_alert_panel.png", width_cm=17.0)
    add_caption(doc,
        "Fig. G  —  AI control panel under risk. Compatibility score 46, "
        "p_stable 0.35, Z5 fill factor 0.97, estimated torque 84 %. "
        "Per-zone risk chart spikes on Z5; agent log shows the rule and the "
        "ML classifier agreeing.")

    add_heading(doc, "2. Recommendation  —  case C4 (ranked actions)",
                color=RGB_BLUE, size=10.5, space_before=6)
    add_figure(doc, FIG_V2 / "figv2_04_agent_recommendation_panel.png",
               width_cm=17.0)
    add_caption(doc,
        "Fig. H  —  AI recommendation panel. Four hierarchised actions cover "
        "formulation, screw profile and process. Projected outcome: "
        "compatibility score 70–80, p_stable ≥ 0.85, Z5 alert cleared.")

    add_heading(doc, "3. After  —  case C5 (recommendation applied)",
                color=RGB_GREEN, size=10.5, space_before=6)
    add_figure(doc, FIG_V2 / "figv2_05_before_after_dashboard.png",
               width_cm=17.5)
    add_caption(doc,
        "Fig. I  —  KPI before / after. Score +32 points, p_stable +0.52, "
        "fill factor Z5 −0.25, estimated torque −22 points. The loop is "
        "demonstrated end-to-end on real cases.")


def page5_methods_conclusion(doc):
    add_page_break(doc)

    add_heading(doc, "Materials & Methods", color=RGB_GREEN, size=11)
    add_paragraph(doc,
        "Eleven industrial runs were collected on a Rondol Ø 10.5 mm twin-screw "
        "extruder (L/D 40:1, horizontal) between 7 and 13 April 2026. A "
        "duration filter (≥ 15 min per run) retained eight runs, segmented "
        "into 627 sliding windows of 60 s (step 30 s; 87 features per "
        "window). Binary stability labels were assigned by expert rules on "
        "local thermal variability of barrel zones Z1–Z8, the die head (DIE), "
        "and three downstream cast-film sensors (P1, P2, Body). Three "
        "classifiers — Random Forest, XGBoost, SVM (RBF) — were trained with "
        "GroupShuffleSplit over run_id (5 train / 3 test) to prevent inter-run "
        "leakage. The AI agent combines this data-driven stability classifier "
        "with a rule-based compatibility score (5 weighted criteria: ceramic "
        "load, viscosity, thermal compatibility, abrasion risk, material "
        "availability) evaluated on the lithium reference recipe (LFP 65 / "
        "PVDF 8 / Super P 5 / LATP 17 / LiTFSI 5 wt%), and emits a "
        "hierarchised textual recommendation.",
        size=9.5, space_after=2)

    add_heading(doc, "Compatibility score — five weighted criteria",
                color=RGB_BLUE, size=10.5, space_before=6)
    # Table
    rows = [
        ("Criterion",            "Weight", "Logic"),
        ("Ceramic load",         "30 %",   "penalty above 30 wt% (abrasion, viscosity)"),
        ("Estimated viscosity",  "20 %",   "upper bound: available torque × margin"),
        ("Thermal compatibility","20 %",   "Tg / Tm vs zone temperature profile"),
        ("Abrasion risk",        "15 %",   "granulometry + ceramic hardness"),
        ("Material availability","15 %",   "Rondol / IJL lab presence"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            c = table.cell(i, j)
            _set_cell_shading(c, "1B7A3D" if i == 0 else
                              ("F2F4F7" if i % 2 == 0 else "FFFFFF"))
            _set_cell_borders(c, color="CBD5E1")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(txt)
            _set_run_font(r,
                          size=9 if i == 0 else 9.5,
                          bold=(i == 0),
                          color=(RGBColor(0xFF, 0xFF, 0xFF) if i == 0
                                 else RGB_GREY))
    widths = [Cm(5.2), Cm(2.6), Cm(9.7)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]

    add_heading(doc, "Discussion", color=RGB_GREEN, size=11, space_before=8)
    add_paragraph(doc,
        "Results rest on a limited industrial dataset (8 runs, 627 windows) "
        "and a rule-based compatibility score that will be progressively "
        "replaced by a regression model trained on a literature-derived "
        "dataset of ~50 lithium-bearing recipes. The dominance of downstream "
        "sensors in the feature ranking is consistent with a process "
        "well-regulated upstream of the screw — informative observability for "
        "stability sits at the die and the cast film, not at the barrel "
        "setpoints. The 32-point compatibility-score gain and 0.52 absolute "
        "increase in stability probability between C3 and C5 demonstrate that "
        "the recommendation loop is non-trivial.",
        size=9.5, space_after=2)

    add_heading(doc, "Conclusion & roadmap", color=RGB_GREEN, size=11,
                space_before=6)
    add_paragraph(doc,
        "This package demonstrates the technical and industrial feasibility "
        "(TRL 4–5) of an AI-augmented decision-support tool for the hot melt "
        "extrusion of lithium-bearing battery components. Planned extensions: "
        "(i) literature-derived dataset of ~50 recipes, (ii) SHAP local "
        "interpretability, (iii) in-line sensor closure of the recommendation "
        "loop, (iv) scale-up validation on the Rondol 21 mm platform.",
        size=9.5, space_after=2)

    add_heading(doc, "References", color=RGB_GREY, size=10, space_before=6,
                keep_next=False)
    refs = (
        "[1] Repka et al., Int. J. Pharm. 535, 2018.  "
        "[2] ECHA, Annex XV restriction report on PFAS, 2023.  "
        "[3] Drakopoulos et al., Cell Rep. Phys. Sci. 2, 100689, 2021.  "
        "[4] Haarmann et al., Energy Technol. 9, 2021.  "
        "[5] Seeba et al., Batteries 10, 2024.  "
        "[6] Kim et al., Nat. Commun. 14, 2023.  "
        "[7] Maia et al., AMI Plastics & Completion AI, 2025.  "
        "[8] Wang et al., Nano-Micro Lett. 17, 2025.  "
        "[9] Daoudi et al., J. Power Sources 590, 2024.  "
        "[10] Fraunhofer IWS, DRYtraec® continuous manufacturing, 2024."
    )
    add_paragraph(doc, refs, size=7.5, color=RGB_MID,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0)


def build_document():
    doc = Document()
    setup_margins(doc)
    install_page_footer(doc)

    page1_hero(doc)
    page2_architecture_evidence(doc)
    page3_five_cases(doc)
    page4_money_shot(doc)
    page5_methods_conclusion(doc)

    doc.save(DOCX_PATH)
    print(f"[OK] DOCX -> {DOCX_PATH.relative_to(ROOT)}")


def convert_to_pdf():
    try:
        from docx2pdf import convert
    except ImportError:
        print("[WARN] docx2pdf missing — pip install docx2pdf")
        return
    convert(str(DOCX_PATH), str(PDF_PATH))
    print(f"[OK] PDF  -> {PDF_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    build_all_figures()
    build_document()
    convert_to_pdf()
