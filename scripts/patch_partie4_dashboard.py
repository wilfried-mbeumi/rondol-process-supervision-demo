# -*- coding: utf-8 -*-
"""
Édition chirurgicale de MBEUMI_Wilfried_THESE.docx : ajoute le vrai tableau de
bord (Tableau 4.5, §4.8) et la phrase de renvoi vers la charte éthique (§4.4),
SANS régénérer le document (préserve la page de garde native Word insérée
manuellement par l'auteur — les tout premiers éléments du corps du document).

Usage : python scripts/patch_partie4_dashboard.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "MBEUMI_Wilfried_THESE.docx"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x26, 0x32, 0x3A)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LTGREY = "F2F2F2"
BASE_FONT = "Calibri"

DASHBOARD_ROWS = [
    ["Indicateur", "Cible / seuil", "Valeur mesurée", "Statut"],
    ["Tests automatisés passants", "100 % de la suite", "705 / 705", "Atteint"],
    ["Couverture fonctionnelle (pages)", "7 pages opérationnelles", "7 / 7", "Atteint"],
    ["Volumétrie du jeu d'apprentissage", "≥ 5 essais exploitables", "798 fenêtres, 8 essais", "Atteint"],
    ["Performance du modèle retenu", "F1-macro ≥ 0,85", "0,918 ± 0,054 (essai réel)", "Atteint"],
    ["Jalon — campagne d'essais", "7 – 13 avril 2026", "Réalisée dans les délais", "Tenu"],
    ["Jalon — démonstration client", "16 juin 2026", "Réalisée le 16 juin 2026", "Tenu"],
    ["Incidents résolus, figés en non-régression", "Suivi continu", "Cf. Tableau 6.1, Annexe F", "Suivi actif"],
]

NEW_INTRO = (
    "Le suivi du projet s'est appuyé sur un tableau de bord d'indicateurs "
    "objectifs, réévalué au fil des itérations et présenté au tuteur "
    "industriel (Direction) lors des points d'avancement."
)

CAPTION_45 = (
    "Tableau 4.5 — Tableau de bord d'indicateurs de suivi, réévalué à chaque "
    "itération et présenté à la Direction (tuteur industriel)."
)

ETHICS_XREF = (
    "L'anticipation des déviances morales possibles (interprétation excessive "
    "de valeurs non calibrées comme industrielles, usage de la donnée "
    "d'authentification au-delà de sa finalité déclarée) est formalisée dans "
    "une charte éthique dédiée en sept principes, présentée en section 8.3 "
    "pour rester adossée à la discussion des limites du système qu'elle encadre."
)


def _set_table_borders(table, color="BFBFBF", sz="4"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _shade(cell, hex_fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def find_paragraph(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise RuntimeError(f"Paragraphe introuvable : {needle!r}")


def build_dashboard_table(doc):
    ncol = len(DASHBOARD_ROWS[0])
    tbl = doc.add_table(rows=len(DASHBOARD_ROWS), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(tbl)
    header, body = DASHBOARD_ROWS[0], DASHBOARD_ROWS[1:]
    for j, val in enumerate(header):
        cell = tbl.rows[0].cells[j]
        _shade(cell, "1F4E79")
        r = cell.paragraphs[0].add_run(val)
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE; r.font.name = BASE_FONT
    for i, row in enumerate(body, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            if i % 2 == 0:
                _shade(cell, LTGREY)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9.5); r.font.color.rgb = DARK; r.font.name = BASE_FONT
    return tbl


def main():
    doc = Document(str(DOCX_PATH))

    # 1) §4.8 — remplace la phrase-prose par une intro courte + insère le vrai tableau
    p_dash = find_paragraph(doc, "Le suivi du projet s'est appuyé sur un tableau de bord")
    for r in list(p_dash.runs):
        r.text = ""
    p_dash.runs[0].text = NEW_INTRO if p_dash.runs else None
    if not p_dash.runs:
        p_dash.add_run(NEW_INTRO)
    else:
        p_dash.runs[0].text = NEW_INTRO
        p_dash.runs[0].font.size = Pt(11)
        p_dash.runs[0].font.color.rgb = DARK
        p_dash.runs[0].font.name = BASE_FONT

    tbl = build_dashboard_table(doc)
    p_dash._p.addnext(tbl._tbl)

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = caption_p.add_run(CAPTION_45)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = GREY; cr.font.name = BASE_FONT
    tbl._tbl.addnext(caption_p._p)

    # 2) §4.4 — renvoi vers la charte éthique, juste après la légende Tableau 4.2
    p_42 = find_paragraph(doc, "Tableau 4.2 — Cartographie des risques")
    xref_p = doc.add_paragraph()
    xr = xref_p.add_run(ETHICS_XREF)
    xr.font.size = Pt(11); xr.font.color.rgb = DARK; xr.font.name = BASE_FONT
    p_42._p.addnext(xref_p._p)

    doc.save(str(DOCX_PATH))
    print("[OK] Patch appliqué :", DOCX_PATH)


if __name__ == "__main__":
    raise SystemExit(main())
