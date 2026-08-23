"""build_dossier_soutenance.py — Assemble le dossier de soutenance en un DOCX imprimable.

Source  : reports/soutenance/DOSSIER_FINAL/_source/*.md (un fichier par bloc).
Sorties : reports/soutenance/DOSSIER_FINAL/
            - DOSSIER_SOUTENANCE_MBEUMI.docx  (dossier complet, avec couverture)
            - « 0 - Lire d'abord.docx » … « 6 - Checklist jour J.docx » (un par bloc)

Le Markdown reste la source versionnée — Git sait montrer les différences d'un
.md, pas d'un .docx binaire, et c'est ce qui permet de vérifier qu'aucun chiffre
périmé ne survit d'une révision à l'autre. Les documents bureautiques sont les
livrables : ils vivent à la racine du dossier, les sources dans _source/.

Ce script met en page, il ne réécrit aucun contenu et ne recalcule aucun chiffre.

Usage : python scripts/build_dossier_soutenance.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
DEST = ROOT / "reports" / "soutenance" / "DOSSIER_FINAL"
# Les Markdown sont la source versionnée (Git sait en montrer les différences,
# pas un .docx binaire). Ils vivent dans _source/ pour que le dossier ouvert par
# le candidat ne contienne que des documents bureautiques.
SRC = DEST / "_source"
OUT = DEST / "DOSSIER_SOUTENANCE_MBEUMI.docx"

TEAL = RGBColor(0x0D, 0x94, 0x88)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x55, 0x5F, 0x6D)

ORDER = [
    "00_LIRE_DABORD.md",
    "01_PRESENTATION_30MIN.md",
    "02_QUESTIONS_REPONSES_15MIN.md",
    "03_ENTRETIEN_PRO_15MIN.md",
    "04_JEU_DE_ROLE_30MIN.md",
    "05_ANTISECHE_A4.md",
    "06_CHECKLIST_JOUR_J.md",
]

# --- inline markdown (gras / italique / code) -------------------------------
_INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|~~.+?~~|\*[^*]+?\*)")


def add_inline(paragraph, text: str, *, size=10.5, color=None, base_bold=False):
    """Écrit `text` dans `paragraph` en interprétant **gras**, *ital*, `code`, ~~barré~~."""
    for chunk in _INLINE.split(text):
        if not chunk:
            continue
        run = paragraph.add_run()
        run.font.size = Pt(size)
        run.font.color.rgb = color or DARK
        run.bold = base_bold
        if chunk.startswith("**") and chunk.endswith("**"):
            run.text = chunk[2:-2]
            run.bold = True
        elif chunk.startswith("~~") and chunk.endswith("~~"):
            run.text = chunk[2:-2]
            run.font.strike = True
            run.font.color.rgb = GREY
        elif chunk.startswith("`") and chunk.endswith("`"):
            run.text = chunk[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(size - 1)
        elif chunk.startswith("*") and chunk.endswith("*"):
            run.text = chunk[1:-1]
            run.italic = True
        else:
            run.text = chunk


def shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_color)
    tcPr.append(el)


def parse_table(lines: list[str], i: int):
    """Retourne (rows, next_index) si un tableau Markdown commence à i."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        # Un séparateur contient au moins un tiret : sans ce test, une ligne
        # d'en-tête vide « | | | » serait prise pour un séparateur.
        is_sep = (any("-" in c for c in cells)
                  and all(set(c) <= set("-: ") for c in cells))
        if not is_sep:
            rows.append(cells)
        i += 1
    return rows, i


def build() -> Path:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(5)

    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)

    # ---------- page de garde ----------
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOSSIER DE SOUTENANCE")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = TEAL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Wilfried Galtier MBEUMI")
    r.font.size = Pt(15)
    r.font.color.rgb = DARK

    for line, sz in [
        ("Mastère 2 Data & Intelligence Artificielle — RNCP 37137 (niveau 7)", 11),
        ("Nexa Digital School · Rondol Industrie", 11),
        ("", 8),
        ("Pré-soutenance : 25 août 2026 · Soutenance : 9 septembre 2026", 11),
        ("Épreuve de 90 minutes en quatre blocs", 11),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(sz)
        r.font.color.rgb = GREY

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------- corps ----------
    for fname in ORDER:
        path = SRC / fname
        if not path.exists():
            print(f"  [!] absent : {fname}")
            continue
        lines = path.read_text(encoding="utf-8").splitlines()

        i = 0
        in_code = False
        while i < len(lines):
            line = lines[i]
            s = line.strip()

            if s.startswith("```"):
                in_code = not in_code
                i += 1
                continue
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                i += 1
                continue

            if not s:
                i += 1
                continue

            # tableau
            if s.startswith("|"):
                rows, i = parse_table(lines, i)
                if not rows:
                    continue
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        txt = row[ci] if ci < len(row) else ""
                        add_inline(para, txt, size=9, base_bold=(ri == 0))
                        if ri == 0:
                            shade(cell, "E8F1F0")
                doc.add_paragraph()
                continue

            # séparateur horizontal
            if s in ("---", "***", "___"):
                i += 1
                continue

            # titres
            m = re.match(r"^(#{1,4})\s+(.*)", s)
            if m:
                level, text = len(m.group(1)), m.group(2)
                if level == 1:
                    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14 if level <= 2 else 9)
                p.paragraph_format.space_after = Pt(5)
                p.paragraph_format.keep_with_next = True
                add_inline(
                    p, text,
                    size={1: 19, 2: 14.5, 3: 12, 4: 11}[level],
                    color=TEAL if level <= 2 else DARK,
                    base_bold=True,
                )
                i += 1
                continue

            # citation
            if s.startswith(">"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.7)
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, s.lstrip("> ").strip(), size=10.5, color=DARK)
                i += 1
                continue

            # cases à cocher
            if re.match(r"^-\s*\[\s*\]", s):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                add_inline(p, "☐  " + re.sub(r"^-\s*\[\s*\]\s*", "", s), size=10.5)
                i += 1
                continue

            # puces
            if s.startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, s[2:], size=10.5)
                i += 1
                continue

            # listes numérotées
            if re.match(r"^\d+\.\s", s):
                p = doc.add_paragraph(style="List Number")
                add_inline(p, re.sub(r"^\d+\.\s", "", s), size=10.5)
                i += 1
                continue

            # paragraphe courant
            p = doc.add_paragraph()
            add_inline(p, s, size=10.5)
            i += 1

    doc.save(OUT)
    return OUT


def build_one(md_name: str, docx_name: str, *, base_size: float = 10.5,
              margin_cm: float = 1.7, space_after: float = 4) -> Path:
    """Génère un DOCX autonome pour un seul fichier source (un bloc d'épreuve)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(base_size)
    style.paragraph_format.space_after = Pt(space_after)

    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm + 0.2)
        section.right_margin = Cm(margin_cm + 0.2)

    lines = (SRC / md_name).read_text(encoding="utf-8").splitlines()
    i, in_code = 0, False
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if s.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            r = doc.add_paragraph().add_run(line)
            r.font.name, r.font.size = "Consolas", Pt(base_size - 1.5)
            i += 1
            continue
        if not s or s in ("---", "***", "___"):
            i += 1
            continue

        if s.startswith("|"):
            rows, i = parse_table(lines, i)
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = table.rows[ri].cells[ci]
                    cell.text = ""
                    add_inline(cell.paragraphs[0], row[ci] if ci < len(row) else "",
                               size=base_size - 1.5, base_bold=(ri == 0))
                    if ri == 0:
                        shade(cell, "E8F1F0")
            doc.add_paragraph()
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", s)
        if m:
            level, text = len(m.group(1)), m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(13 if level <= 2 else 8)
            p.paragraph_format.keep_with_next = True
            add_inline(p, text,
                       size={1: 18, 2: 14, 3: 11.5, 4: 10.5}[level],
                       color=TEAL if level <= 2 else DARK, base_bold=True)
            i += 1
            continue

        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_inline(p, s.lstrip("> ").strip(), size=base_size)
            i += 1
            continue

        if re.match(r"^-\s*\[\s*\]", s):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            add_inline(p, "☐  " + re.sub(r"^-\s*\[\s*\]\s*", "", s), size=base_size)
            i += 1
            continue

        if s.startswith(("- ", "* ")):
            add_inline(doc.add_paragraph(style="List Bullet"), s[2:], size=base_size)
            i += 1
            continue

        if re.match(r"^\d+\.\s", s):
            add_inline(doc.add_paragraph(style="List Number"),
                       re.sub(r"^\d+\.\s", "", s), size=base_size)
            i += 1
            continue

        add_inline(doc.add_paragraph(), s, size=base_size)
        i += 1

    out = DEST / docx_name
    doc.save(out)
    return out


# Un Word autonome par bloc d'épreuve — c'est ce que le candidat ouvre et annote.
# L'antisèche est resserrée pour tenir sur un recto-verso : c'est la seule feuille
# emportée le jour J, elle perd sa fonction si elle déborde sur une 3e page.
SPLIT = {
    "00_LIRE_DABORD.md": ("0 - Lire d'abord.docx", {}),
    "01_PRESENTATION_30MIN.md": ("1 - Presentation 30 min.docx", {}),
    "02_QUESTIONS_REPONSES_15MIN.md": ("2 - Questions du jury 15 min.docx", {}),
    "03_ENTRETIEN_PRO_15MIN.md": ("3 - Entretien professionnel 15 min.docx", {}),
    "04_JEU_DE_ROLE_30MIN.md": ("4 - Jeu de role 30 min.docx", {}),
    "05_ANTISECHE_A4.md": ("5 - Antiseche a imprimer.docx",
                           {"base_size": 8.0, "margin_cm": 1.0, "space_after": 1.5}),
    "06_CHECKLIST_JOUR_J.md": ("6 - Checklist jour J.docx",
                               {"base_size": 9.5, "margin_cm": 1.4, "space_after": 2.5}),
}


if __name__ == "__main__":
    out = build()
    print(f"[OK] {out.name} — {out.stat().st_size / 1024:.0f} Ko (dossier complet)")
    for md, (docx, opts) in SPLIT.items():
        if (SRC / md).exists():
            f = build_one(md, docx, **opts)
            print(f"[OK] {f.name} — {f.stat().st_size / 1024:.0f} Ko")
