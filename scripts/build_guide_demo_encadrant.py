# -*- coding: utf-8 -*-
"""
Guide de présentation en direct pour l'encadrant : checklist + timing + Q&A sur le modèle + commandes dataset.
À imprimer ou consulter en parallèle pendant la démo.

Usage : python scripts/build_guide_demo_encadrant.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "GUIDE_DEMO_LIVE_ENCADRANT.docx"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x26, 0x32, 0x3A)
GREY = RGBColor(0x59, 0x59, 0x59)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0xC0)


def shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def run(p, text, sz=11, color=DARK, bold=False, italic=False, font="Calibri"):
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.color.rgb = color; r.font.name = font
    r.font.bold = bold; r.font.italic = italic
    return r


def h(doc, text, sz=14, color=BLUE):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    run(p, text, sz=sz, color=color, bold=True)
    return p


def checkbox(doc, label, detail=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run(p, "☐ " + label, bold=True, color=DARK, sz=10)
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.0); p2.paragraph_format.space_after = Pt(2)
        run(p2, detail, sz=9.5, color=GREY, italic=True)


def code_box(doc, code_lines):
    tbl = doc.add_table(rows=1, cols=1)
    c = tbl.rows[0].cells[0]
    shade(c, "F5F5F5")
    for i, ln in enumerate(code_lines):
        if i > 0:
            c.add_paragraph()
        p = c.paragraphs[i] if i < len(c.paragraphs) else c.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Courier New"; r.font.size = Pt(8.5); r.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.2); s.bottom_margin = Cm(1.2)
        s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(title, "GUIDE DÉMO LIVE — Présentation Encadrant", sz=16, color=BLUE, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(10)
    run(sub, "24 juillet 2026  •  Plateforme Rondol SSB", sz=10, color=GREY, italic=True)

    # --- PHASE 1: Documents (5 min) ---
    h(doc, "PHASE 1 : Documents (5 min)")
    checkbox(doc, "Ouvrir NOTE_SYNTHESE_ENCADRANT.pdf", "Présentation synthétique 2 pages (livrables, indicateurs)")
    checkbox(doc, "Ouvrir MBEUMI_Wilfried_THESE.pdf", "Page 1 : couverture native Word • Page 7 : sommaire • Pages 25–32 : Partie 4 (Gestion de projet, Tableau 4.5 nouveau) • Pages 43–50+ : Partie 5 (ML, modèles)")
    checkbox(doc, "Points clés à couvrir :", "✓ Conforme guide RNCP (méthode Kanban, rétroplanning, budget, tableau de bord, risques) ✓ 76 pages, figures + captures réelles intégrées ✓ Charte éthique 7 principes (§8.3)")

    # --- PHASE 2: Notebook (5 min) ---
    h(doc, "PHASE 2 : Notebook d'analyse (5 min)")
    checkbox(doc, "Jupyter ouvert sur http://localhost:8889/tree", "Notebook: 26 cellules exécutées de bout en bout")
    checkbox(doc, "Sections à montrer :", "Cell 1–5 : données brutes + volumétrie ✓ Cell 10–15 : modélisation ML (comparaison 5 modèles) ✓ Cell 20–26 : DÉMONSTRATION CLÉE — moteur de règles expert génère alertes/recommandations indépendamment du ML")
    checkbox(doc, "Message clé", "Les recommandations ne viennent PAS du modèle ML, mais d'un moteur à règles explicable (AgentIndustrial_v1.core.rules)")

    # --- PHASE 3: Application (7 min) ---
    h(doc, "PHASE 3 : Application Streamlit (7 min)")
    checkbox(doc, "URL", "https://rondol-process-supervision-demo.streamlit.app  OU  streamlit run app/Supervision.py (localhost)")
    checkbox(doc, "Authentification (1 min)", "☐ Login avec email: demo@rondol.local / moteur: demo123 (local auto-provision) — Supabase auth via PBKDF2-HMAC-SHA256 (jamais en clair)")
    checkbox(doc, "Supervision (2 min)", "☐ KPIs (fill factor, residence time, volume) ☐ Score de risque (68/100 en démo) ☐ Alertes agent + recos ☐ Stabilité thermique")
    checkbox(doc, "Profile (1 min)", "☐ Screw design (81 éléments, zone Z1–Z8) ☐ +/− counters ☐ Résumé des KPIs")
    checkbox(doc, "Settings (1 min)", "☐ Feeders calibration ☐ AI thresholds ☐ Enregistrer validé")
    checkbox(doc, "Process Engine (2 min)", "☐ Moteur 7 couches : screw_logic → engine → torque → enrichissement ☐ Affichage nominal (non calibré industriel)")

    # --- PHASE 4: Questions attendues (3 min) ---
    h(doc, "QUESTIONS ATTENDUES + RÉPONSES CLÉS")

    qas = [
        ("Q: Comment avez-vous généré le dataset d'entraînement ?",
         "R: Données réelles captées (12 voies thermiques) lors de 8 essais avril 2026, "
         "puis segmentées en 798 fenêtres (fenêtres glissantes de 60 s). Voir notebook cell 5."),

        ("Q: Pourquoi RandomForest et pas Deep Learning ?",
         "R: RandomForest : robustesse sur petit dataset (8 essais), interprétabilité, pas de surapprentissage observé. "
         "Deep Learning nécessiterait 100s d'essais. F1-macro 0,918 ± 0,054 vs XGBoost 0,903 → RandomForest sélectionné."),

        ("Q: Comment avez-vous validé le modèle ?",
         "R: Leave-One-Group-Out (par essai, pas par fenêtre) pour éviter fuite data. "
         "Externe : AUC 0,753 sur jeu consolidé (3479 fenêtres). Voir reports/eval_consolidated_w60.json."),

        ("Q: Les alertes, c'est le modèle ML ou des règles ?",
         "R: RÈGLES PURES, pas ML. AgentIndustrial_v1.core.rules.py évalue seuils manuel : "
         "SME > 0.30 kWh/kg → alerte. ΔT_zone > 10°C → alerte. Moteur à règles = explainable + auditable."),

        ("Q: Pourquoi les valeurs ne sont pas calibrées industriellement ?",
         "R: 8 essais = trop peu pour calibration robuste. Affichage explicite 'valeurs nominales'. "
         "Positionnement : aide à la décision (tendance relative) PAS mesure absolue. Cohérent avec contexte R&D."),

        ("Q: Avez-vous fait du feature engineering ?",
         "R: Oui. Fenêtres glissantes + lags (T[t-1], T[t-2]). Variables brutes : 12 capteurs → 87 features ML. "
         "Voir src/build_dataset.py + notebook cell 15 (feature importance)."),
    ]

    for q, a in qas:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        run(p, q, sz=10, color=BLUE, bold=True)
        p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(0.8); p2.paragraph_format.space_after = Pt(6)
        run(p2, a, sz=9.5, color=DARK)

    # --- PHASE 5: Reproduction du dataset (CLI) ---
    h(doc, "POUR REPRODUIRE LE DATASET (commandes)")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    run(p, "Depuis la racine du projet (PowerShell ou bash) :", bold=True, sz=10)
    code_box(doc, [
        "# 1. Construire dataset (données brutes → features)",
        "python -m src.build_dataset",
        "",
        "# 2. Entraîner modèles (RF, XGBoost, SVM, etc.)",
        "python -m src.train_models --window 60",
        "",
        "# 3. Évaluation robustesse (100 graines aléatoires, LOGO validation)",
        "python -m src.robustness_check --window 60 --n-seeds 100",
        "",
        "# 4. Évaluation externe (sur jeu consolidé, hors entraînement)",
        "python scripts/evaluate_on_consolidated.py",
        "",
        "# Résultats : reports/ml_metrics_w60.json, reports/eval_consolidated_w60.json, models/RandomForest_w60_augmented.joblib"
    ])

    # --- Checklist finale ---
    h(doc, "CHECKLIST FINALE (à cocher pendant la démo)", color=RED)
    final = [
        ("Documents présentés", ""),
        ("Notebook démonstration exécuté", ""),
        ("Application live testée (login + pages)", ""),
        ("Q&A modèle ML couverts", ""),
        ("Encadrant a vu le code dataset (repo)", ""),
        ("Encadrant satisfait", ""),
    ]
    for item, note in final:
        checkbox(doc, item, note)

    doc.save(str(OUT))
    print("[OK]", OUT)
    print("\nOuvrire et imprimer (ou consulter en parallèle pendant la démo)")


if __name__ == "__main__":
    raise SystemExit(main())
