"""Build the INDUSTRIAL SHOWCASE v3 (DOCX + PDF) — premium symposium edition.

Visual direction: executive symposium showcase, full-bleed dark cockpits,
premium typography, very few text blocks, large KPI badges, timeline roadmap.

Outputs (never overwrites v1 or v2):
  reports/poster_abstract/figures/generated_v3/v3_*.png
  reports/poster_abstract/Mbeumi_2026_IndustrialShowcase_v3.docx
  reports/poster_abstract/Mbeumi_2026_IndustrialShowcase_v3.pdf
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Circle, Rectangle, Wedge
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
FIG_V1 = ROOT / "reports" / "poster_abstract" / "figures" / "generated"
FIG_V3 = ROOT / "reports" / "poster_abstract" / "figures" / "generated_v3"
FIG_V3.mkdir(parents=True, exist_ok=True)
OUT_DIR = ROOT / "reports" / "poster_abstract"
DOCX_PATH = OUT_DIR / "Mbeumi_2026_IndustrialShowcase_v3.docx"
PDF_PATH = OUT_DIR / "Mbeumi_2026_IndustrialShowcase_v3.pdf"

# ---------------------------------------------------------------------------
# Premium palette — dark navy base, electric green Rondol, accent orange
# ---------------------------------------------------------------------------
INK = "#0B1220"          # near-black navy
INK_2 = "#111B2E"        # panel background
INK_3 = "#1A2740"        # raised panel
LINE = "#2A3A5C"         # subtle separator on dark
LINE_2 = "#1F2C45"
PAPER = "#FFFFFF"
PAPER_TINT = "#F4F6FA"
SLATE = "#94A3B8"
SLATE_2 = "#64748B"
TXT = "#E5ECF6"
TXT_DIM = "#A8B5CC"
GREEN = "#1FB07A"        # electric green
GREEN_DEEP = "#0E7A52"
GREEN_SOFT = "#D8F4E7"
BLUE = "#3A8DDE"
BLUE_DEEP = "#1F5F9E"
ORANGE = "#F08A3E"
ORANGE_DEEP = "#C4651C"
RED = "#E04A4A"
AMBER = "#F5C342"
PURPLE = "#9A7BD8"

# Word colors
def C(hexcode: str) -> RGBColor:
    h = hexcode.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 11,
    "axes.edgecolor": LINE,
    "axes.labelcolor": TXT,
    "xtick.color": TXT,
    "ytick.color": TXT,
    "text.color": TXT,
    "figure.facecolor": PAPER,
    "axes.facecolor": PAPER,
    "savefig.dpi": 240,
    "savefig.pad_inches": 0.0,
})


# ============================================================================
# Helpers
# ============================================================================

def _box(ax, x, y, w, h, *, fc, ec=None, lw=0, radius=2.0, alpha=1.0):
    p = FancyBboxPatch(
        (x, y), w, h,
        boxstyle=f"round,pad=0.0,rounding_size={radius}",
        linewidth=lw, edgecolor=ec or fc, facecolor=fc, alpha=alpha,
    )
    ax.add_patch(p)
    return p


def _rect(ax, x, y, w, h, *, fc, ec=None, lw=0, alpha=1.0):
    r = Rectangle((x, y), w, h, linewidth=lw,
                  edgecolor=ec or fc, facecolor=fc, alpha=alpha)
    ax.add_patch(r)
    return r


def _hline(ax, x0, x1, y, *, c=LINE, lw=1.2, alpha=1.0):
    ax.plot([x0, x1], [y, y], color=c, lw=lw, alpha=alpha, solid_capstyle="butt")


def _vline(ax, x, y0, y1, *, c=LINE, lw=1.2, alpha=1.0):
    ax.plot([x, x], [y0, y1], color=c, lw=lw, alpha=alpha, solid_capstyle="butt")


def _arrow(ax, x0, y0, x1, y1, *, color=SLATE, lw=2.0, style="-|>", scale=18):
    a = FancyArrowPatch((x0, y0), (x1, y1), arrowstyle=style,
                        mutation_scale=scale, linewidth=lw, color=color)
    ax.add_patch(a)


# ============================================================================
# FIG v3-01 — COVER (full A4)
# ============================================================================

def v3_cover():
    # A4 portrait, full bleed
    fig = plt.figure(figsize=(8.27, 11.69), facecolor=INK)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 141); ax.axis("off")

    # Background gradient suggestion via two stacked rectangles
    _rect(ax, 0, 0, 100, 141, fc=INK)

    # Decorative side rail
    _rect(ax, 0, 0, 1.2, 141, fc=GREEN)

    # Top strap
    _rect(ax, 5, 132, 90, 0.4, fc=LINE)
    ax.text(5, 134.5, "RONDOL  ·  INSTITUT JEAN LAMOUR",
            color=TXT_DIM, fontsize=9, fontweight="bold",
            va="center")
    ax.text(95, 134.5, "INDUSTRIAL SHOWCASE  ·  v3",
            color=GREEN, fontsize=9, fontweight="bold",
            va="center", ha="right")

    # Eyebrow chip
    _box(ax, 5, 121, 38, 4.5, fc=INK_3, radius=1.2)
    ax.text(7, 123.25, "AI  ·  EXTRUSION  ·  LITHIUM-BEARING ELECTRODES",
            color=GREEN, fontsize=8.5, fontweight="bold",
            va="center")

    # MAIN TITLE — large, premium
    ax.text(5, 110,
            "Industrial AI Copilot",
            color=PAPER, fontsize=40, fontweight="bold", va="center")
    ax.text(5, 102,
            "for Twin-Screw Extrusion",
            color=PAPER, fontsize=40, fontweight="bold", va="center")
    ax.text(5, 95.5,
            "of Solid-State Battery Electrodes",
            color=GREEN, fontsize=24, fontweight="bold", va="center")

    # Subtitle / strategic line
    _rect(ax, 5, 92.6, 12, 0.5, fc=GREEN)
    ax.text(5, 89.5,
            "From recipe to ranked recommendation, on a real Rondol Ø 10.5 mm  ·  L/D 40 extruder.",
            color=TXT_DIM, fontsize=11, va="center", style="italic")

    # Mini workflow ribbon — 5 steps
    steps = [
        ("RECIPE",         GREEN),
        ("PROCESS",        BLUE),
        ("SCREW",          BLUE),
        ("AI ENGINE",      ORANGE),
        ("ACTION",         GREEN),
    ]
    y_rib = 76
    rw = 16.0; gap = 1.6
    total = 5 * rw + 4 * gap
    x_start = (100 - total) / 2
    for i, (lbl, c) in enumerate(steps):
        x = x_start + i * (rw + gap)
        _box(ax, x, y_rib - 3.5, rw, 7, fc=INK_3, ec=c, lw=2, radius=1.4)
        _rect(ax, x + 1, y_rib - 3.0, 1, 6, fc=c)
        ax.text(x + rw / 2 + 0.5, y_rib, lbl, color=PAPER,
                fontsize=9.5, fontweight="bold", ha="center", va="center")
        if i < 4:
            ax.text(x + rw + gap / 2, y_rib, "→", color=SLATE,
                    fontsize=14, ha="center", va="center")

    # KPI strip — 4 huge tiles
    kpis = [
        ("0.917",  "TEST F1-MACRO",        "Random Forest · 60 s window", GREEN),
        ("0.976",  "ROC-AUC  5-FOLD CV",   "± 0.021 across runs",         BLUE),
        ("+32",    "POINTS  ·  C3 → C5",   "compatibility score gain",    ORANGE),
        ("TRL 4–5","INDUSTRIAL READINESS", "demonstrator on real runs",   GREEN_DEEP),
    ]
    y_kpi = 53
    kw = 21; gap = 1.5
    total = 4 * kw + 3 * gap
    x0 = (100 - total) / 2
    for i, (val, lbl, sub, c) in enumerate(kpis):
        x = x0 + i * (kw + gap)
        _box(ax, x, y_kpi - 12, kw, 16, fc=INK_2, ec=LINE, lw=1.0, radius=1.6)
        _rect(ax, x, y_kpi + 4 - 0.6, kw, 0.6, fc=c)
        ax.text(x + kw / 2, y_kpi + 0.6, val,
                color=c, fontsize=24, fontweight="bold",
                ha="center", va="center")
        ax.text(x + kw / 2, y_kpi - 5.5, lbl,
                color=TXT, fontsize=8.5, fontweight="bold",
                ha="center", va="center")
        ax.text(x + kw / 2, y_kpi - 9.5, sub,
                color=SLATE, fontsize=7.8,
                ha="center", va="center", style="italic")

    # Pull-quote
    _rect(ax, 5, 31.5, 1, 4, fc=GREEN)
    ax.text(8, 33.5,
            "“The agent diagnoses a documented failure mode, ranks the corrective "
            "action, and proves the loop end-to-end on real cases.”",
            color=PAPER, fontsize=12, va="center", style="italic")

    # Author / affiliation card
    _box(ax, 5, 14, 90, 11, fc=INK_2, ec=LINE, lw=1, radius=1.6)
    _rect(ax, 5, 23.5, 90, 0.4, fc=GREEN)
    ax.text(7, 20.4, "Wilfried Galtier  Mbeumi",
            color=PAPER, fontsize=13, fontweight="bold", va="center")
    ax.text(7, 17.4, "Rondol Industrie  ·  Institut Jean Lamour (IJL), Université de Lorraine",
            color=TXT_DIM, fontsize=9.5, va="center")
    ax.text(93, 20.4, "Maël Gallas",
            color=PAPER, fontsize=11, fontweight="bold", va="center", ha="right")
    ax.text(93, 17.4, "Rondol Industrie  ·  industrial supervisor",
            color=TXT_DIM, fontsize=9, va="center", ha="right")

    # Bottom strap
    _rect(ax, 0, 0, 100, 7, fc=INK_3)
    ax.text(5, 3.5, "SYMPOSIUM  ·  15 MAY 2026",
            color=TXT, fontsize=9, fontweight="bold", va="center")
    ax.text(95, 3.5, "Mastère Data & IA  ·  RNCP 37137",
            color=TXT_DIM, fontsize=8.5, va="center", ha="right")

    fig.savefig(FIG_V3 / "v3_cover.png", facecolor=INK, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-02 — EXECUTIVE OVERVIEW (workflow + KPI banner)
# ============================================================================

def v3_executive_overview():
    fig = plt.figure(figsize=(16, 9.0), facecolor=PAPER)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 56); ax.axis("off")

    # Header band
    _rect(ax, 0, 50, 100, 6, fc=INK)
    _rect(ax, 0, 49.7, 100, 0.4, fc=GREEN)
    ax.text(3, 53, "EXECUTIVE OVERVIEW",
            color=PAPER, fontsize=14, fontweight="bold",
            va="center")
    ax.text(97, 53, "AI copilot for hot-melt extrusion of battery electrodes",
            color=TXT_DIM, fontsize=10, ha="right", va="center", style="italic")

    # Subtitle
    ax.text(3, 46,
            "An industrial decision-support agent that closes the formulation "
            "→ process → screw → risk → recommendation loop.",
            color="#1F2937", fontsize=11.5, va="center")

    # 5-step workflow with rich icons
    steps = [
        ("01", "RECIPE",       "LFP / PVDF / SuperP\nLATP / LiTFSI",       GREEN_DEEP, GREEN),
        ("02", "PROCESS",      "T(Z1..Z8) · rpm\nthroughput · L/D 40",    BLUE_DEEP,  BLUE),
        ("03", "SCREW PROFILE","conveying · kneading\ncompression · tip",  BLUE_DEEP,  BLUE),
        ("04", "AI ENGINE",    "Random Forest w60\nF1 0.917 · AUC 0.976",  ORANGE_DEEP, ORANGE),
        ("05", "ACTION",       "ranked recommendation\nformulation + screw\n+ process",
                                                                            GREEN_DEEP, GREEN),
    ]
    y0 = 22
    bw = 16.4; gap = (100 - 6 - 5 * bw) / 4
    for i, (num, lbl, body, cdark, cbright) in enumerate(steps):
        x = 3 + i * (bw + gap)
        _box(ax, x, y0, bw, 19, fc=PAPER_TINT, ec=cbright, lw=2, radius=1.6)
        # numeric badge
        _box(ax, x + 1.2, y0 + 14.5, 4.5, 4, fc=cbright, radius=1.0)
        ax.text(x + 1.2 + 2.25, y0 + 16.5, num, color=PAPER,
                fontsize=10.5, fontweight="bold", ha="center", va="center")
        # label
        ax.text(x + bw / 2 + 1.8, y0 + 16.5, lbl, color=cdark,
                fontsize=10.5, fontweight="bold", ha="center", va="center")
        # body
        ax.text(x + bw / 2, y0 + 7.5, body, color="#374151", fontsize=9.5,
                ha="center", va="center", linespacing=1.4)
        # accent strip bottom
        _rect(ax, x, y0, bw, 1.0, fc=cbright)

    # Closed loop arrow
    _arrow(ax, 85, 19.5, 22, 19.5, color=ORANGE, lw=2.2, scale=20)
    ax.text(53, 17.2, "closed loop  ·  the recommendation re-enters the process",
            color=ORANGE_DEEP, fontsize=10, ha="center", va="center",
            style="italic", fontweight="bold")

    # KPI banner
    _rect(ax, 0, 0, 100, 11, fc=INK_2)
    _rect(ax, 0, 10.6, 100, 0.4, fc=GREEN)
    kpis = [
        ("0.917",  "Test F1-macro (RF · 60 s)", GREEN),
        ("0.976",  "ROC-AUC · 5-fold CV",       BLUE),
        ("+32",    "Score gain C3 → C5",        ORANGE),
        ("+0.52",  "p_stable gain C3 → C5",     GREEN),
        ("8",      "industrial runs · 627 windows", PURPLE),
        ("TRL 4–5","industrial readiness",      GREEN),
    ]
    kw = (100 - 6 * 1.5 - 3) / 6
    for i, (val, lbl, c) in enumerate(kpis):
        x = 1.5 + i * (kw + 1.5)
        ax.text(x + kw / 2, 7.0, val,
                color=c, fontsize=18, fontweight="bold",
                ha="center", va="center")
        ax.text(x + kw / 2, 3.0, lbl,
                color=TXT_DIM, fontsize=8.5,
                ha="center", va="center")

    fig.savefig(FIG_V3 / "v3_executive_overview.png", facecolor=PAPER, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-03 — ARCHITECTURE (premium 4-layer with icons)
# ============================================================================

def v3_architecture():
    fig = plt.figure(figsize=(16, 9.0), facecolor=PAPER)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 56); ax.axis("off")

    # Header band
    _rect(ax, 0, 50, 100, 6, fc=INK)
    _rect(ax, 0, 49.7, 100, 0.4, fc=GREEN)
    ax.text(3, 53, "AGENT ARCHITECTURE",
            color=PAPER, fontsize=14, fontweight="bold", va="center")
    ax.text(97, 53, "four traceable layers  ·  no black box",
            color=TXT_DIM, fontsize=10, ha="right", va="center", style="italic")

    layers = [
        ("PRESENTATION  ·  Streamlit HMI",
         "Home  ·  Profile  ·  Settings  ·  Run analysis  ·  History",
         "OPERATOR", BLUE, BLUE_DEEP, "▣"),
        ("DECISION  ·  AI agent",
         "compatibility score (5 rules)  ·  ML stability classifier  ·  alerts  ·  ranked recommendation",
         "AGENT LOOP", ORANGE, ORANGE_DEEP, "◈"),
        ("PROCESS LOGIC  ·  Physics",
         "screw_logic.py  ·  fill factor  ·  residence time  ·  SME  ·  per-zone risk",
         "DETERMINISTIC", BLUE, BLUE_DEEP, "⚙"),
        ("DATA",
         "8 industrial runs  ·  627 sliding windows  ·  87 features  ·  9 trained models",
         "INDUSTRIAL", GREEN, GREEN_DEEP, "▤"),
    ]
    yh = 9.5; gap = 1.5
    y_top = 42
    for i, (title, body, badge, c, cdeep, icon) in enumerate(layers):
        y = y_top - (i + 1) * (yh + gap) + gap
        _box(ax, 3, y, 94, yh, fc=PAPER_TINT, ec=c, lw=2, radius=1.4)
        _rect(ax, 3, y, 1.6, yh, fc=c)
        # icon disc
        _box(ax, 7, y + yh / 2 - 2.5, 5, 5, fc=c, radius=2.5)
        ax.text(9.5, y + yh / 2, icon, color=PAPER,
                fontsize=14, fontweight="bold", ha="center", va="center")
        # title
        ax.text(14, y + yh - 2.3, title, color=cdeep,
                fontsize=12, fontweight="bold", va="center")
        # body
        ax.text(14, y + 2.3, body, color="#374151", fontsize=10, va="center")
        # right-side badge
        _box(ax, 81, y + yh / 2 - 1.6, 14, 3.2, fc=c, radius=1.0)
        ax.text(88, y + yh / 2, badge, color=PAPER,
                fontsize=8.5, fontweight="bold", ha="center", va="center")

    # Bottom strip — what makes this auditable
    _rect(ax, 0, 0, 100, 11, fc=INK_2)
    _rect(ax, 0, 10.6, 100, 0.4, fc=GREEN)
    ax.text(50, 7.0,
            "Every alert is justified by an explicit signal  —  rule  +  physics  +  probability.",
            color=PAPER, fontsize=12.5, fontweight="bold", ha="center", va="center",
            style="italic")
    ax.text(50, 3.0,
            "Traceable  ·  reproducible  ·  WCAG 2.1  ·  RGPD-friendly  ·  RNCP 37137",
            color=TXT_DIM, fontsize=9.5, ha="center", va="center")

    fig.savefig(FIG_V3 / "v3_architecture.png", facecolor=PAPER, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-04 — FIVE CASES PREMIUM (taller cards with mini-charts)
# ============================================================================

def v3_five_cases():
    fig = plt.figure(figsize=(16, 9.0), facecolor=PAPER)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 56); ax.axis("off")

    # Header
    _rect(ax, 0, 50, 100, 6, fc=INK)
    _rect(ax, 0, 49.7, 100, 0.4, fc=GREEN)
    ax.text(3, 53, "FIVE  ·  CASE  ·  DEMONSTRATION",
            color=PAPER, fontsize=14, fontweight="bold", va="center")
    ax.text(97, 53, "baseline  →  optimised  →  at risk  →  recommendation  →  applied",
            color=TXT_DIM, fontsize=9.5, ha="right", va="center", style="italic")

    cases = [
        ("C1", "BASELINE",       "LFP 65 · LATP 17",     "65",  "0.84", "STABLE",  GREEN,  None),
        ("C2", "OPTIMISED SCREW","Z5 45° → 30°",         "82",  "0.91", "STABLE",  GREEN,  None),
        ("C3", "AT RISK",        "LATP 17 → 35 %",       "46",  "0.35", "ALERT Z5",RED,    "alert"),
        ("C4", "AI RECO",        "4 ranked actions",     "—",   "—",    "READY",   ORANGE, "reco"),
        ("C5", "APPLIED",        "reco → process",       "78",  "0.87", "STABLE",  GREEN,  None),
    ]
    n = 5
    cw = 17.5; gap = (100 - 6 - n * cw) / (n - 1)
    y0 = 4

    for i, (cid, label, sub, score, p, status, c, marker) in enumerate(cases):
        x = 3 + i * (cw + gap)
        # main card
        _box(ax, x, y0, cw, 40, fc=PAPER_TINT, ec=c, lw=2.2, radius=1.6)
        # color header strip
        _rect(ax, x, y0 + 40 - 5.5, cw, 5.5, fc=c)
        ax.text(x + 1.5, y0 + 40 - 2.7, cid, color=PAPER,
                fontsize=16, fontweight="bold", va="center")
        ax.text(x + cw - 1.5, y0 + 40 - 2.7, label, color=PAPER,
                fontsize=8.8, fontweight="bold", ha="right", va="center")
        # subline
        ax.text(x + cw / 2, y0 + 40 - 8.5, sub,
                color="#374151", fontsize=10, ha="center", va="center",
                style="italic")

        # Score row
        _box(ax, x + 1.4, y0 + 22, cw - 2.8, 7.5, fc=PAPER, ec=LINE, lw=1, radius=1.2)
        ax.text(x + 3, y0 + 27, "SCORE", color=SLATE_2,
                fontsize=7.5, fontweight="bold", va="center")
        ax.text(x + cw - 3, y0 + 27, score, color=c,
                fontsize=18, fontweight="bold", ha="right", va="center")
        ax.text(x + 3, y0 + 24, "p_stable", color=SLATE_2,
                fontsize=7.5, fontweight="bold", va="center")
        ax.text(x + cw - 3, y0 + 24, p, color="#1F2937",
                fontsize=11, fontweight="bold", ha="right", va="center")

        # Status pill
        _box(ax, x + 2.5, y0 + 14, cw - 5, 5.5, fc=c, radius=1.2)
        ax.text(x + cw / 2, y0 + 16.8, status, color=PAPER,
                fontsize=10.5, fontweight="bold", ha="center", va="center")

        # Mini formulation bar (LFP / PVDF / SP / LATP / LiTFSI)
        if i in (0, 1):  # baseline / optimised — same formulation
            comp = [65, 8, 5, 17, 5]
        elif i in (2, 3):  # at risk / reco context
            comp = [47, 8, 5, 35, 5]
        else:
            comp = [65, 8, 5, 17, 5]
        bar_y = y0 + 6
        bar_x = x + 2.0
        bar_w = cw - 4
        bar_h = 4
        _box(ax, bar_x, bar_y, bar_w, bar_h, fc=PAPER, ec=LINE, lw=1, radius=0.6)
        cur = bar_x
        comp_colors = [GREEN_DEEP, BLUE_DEEP, SLATE_2, ORANGE if i in (2, 3) else PURPLE, AMBER]
        for w_pct, col in zip(comp, comp_colors):
            seg_w = bar_w * w_pct / 100
            _rect(ax, cur, bar_y, seg_w, bar_h, fc=col)
            cur += seg_w
        ax.text(x + 2.0, bar_y - 1.4, "FORMULATION", color=SLATE_2,
                fontsize=6.8, fontweight="bold", va="center")
        ax.text(x + cw - 2.0, bar_y - 1.4,
                f"LATP {comp[3]} %",
                color=ORANGE if i in (2, 3) else SLATE_2,
                fontsize=7.5, fontweight="bold", ha="right", va="center")

        # Connector arrow to next card
        if i < n - 1:
            ay = y0 + 22
            _arrow(ax, x + cw + 0.3, ay, x + cw + gap - 0.3, ay,
                   color=SLATE_2, lw=1.5, scale=12)

    fig.savefig(FIG_V3 / "v3_five_cases.png", facecolor=PAPER, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-05 — COCKPIT C3 (full-bleed dark, immersive)
# ============================================================================

def v3_cockpit_risk():
    fig = plt.figure(figsize=(16, 10.5), facecolor=INK)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 65); ax.axis("off")

    # All dark
    _rect(ax, 0, 0, 100, 65, fc=INK)

    # Topbar
    _rect(ax, 0, 60, 100, 5, fc=INK_3)
    _rect(ax, 0, 60, 100, 0.3, fc=RED)
    ax.text(2, 62.5, "● LIVE", color=RED, fontsize=10, fontweight="bold", va="center")
    ax.text(10, 62.5, "AI EXTRUSION COPILOT", color=PAPER, fontsize=11.5,
            fontweight="bold", va="center")
    ax.text(98, 62.5, "case C3  ·  Ø 10.5 mm  ·  L/D 40  ·  semi-dry  ·  180 rpm",
            color=SLATE, fontsize=9.5, ha="right", va="center")

    # Status banner — RED
    _rect(ax, 0, 53, 100, 6, fc=RED)
    ax.text(2, 56, "▲  AT RISK", color=PAPER, fontsize=15,
            fontweight="bold", va="center")
    ax.text(50, 56, "ceramic overload detected  ·  LATP 35 wt%",
            color=PAPER, fontsize=12, va="center")
    ax.text(98, 56, "STABILITY  CLASSIFIER:  UNSTABLE",
            color=PAPER, fontsize=10, ha="right", va="center",
            fontweight="bold")

    # KPI tiles (4) — large
    kpis = [
        ("46 / 100",   "COMPATIBILITY SCORE", RED, "▼"),
        ("0.35",       "p_stable  (RF · w60)", RED, "▼"),
        ("0.97",       "FILL FACTOR  Z5",      AMBER, "▲"),
        ("84 %",       "ESTIMATED TORQUE",     AMBER, "▲"),
    ]
    kw = 22; gap = 2
    total = 4 * kw + 3 * gap
    x_kpi = (100 - total) / 2
    for i, (val, lbl, c, arrow) in enumerate(kpis):
        x = x_kpi + i * (kw + gap)
        _box(ax, x, 39, kw, 12, fc=INK_2, ec=c, lw=2, radius=1.4)
        _rect(ax, x, 50.7, kw, 0.4, fc=c)
        ax.text(x + 2, 48, arrow, color=c, fontsize=16,
                fontweight="bold", va="center")
        ax.text(x + kw / 2, 46, val, color=c, fontsize=22,
                fontweight="bold", ha="center", va="center")
        ax.text(x + kw / 2, 41.2, lbl, color=TXT_DIM, fontsize=8.5,
                ha="center", va="center")

    # Per-zone risk chart (left)
    zones = [f"Z{i}" for i in range(1, 9)]
    risk = [0.10, 0.15, 0.22, 0.35, 0.92, 0.45, 0.30, 0.18]

    chart_x = 2; chart_y = 7; chart_w = 56; chart_h = 28
    _box(ax, chart_x, chart_y, chart_w, chart_h, fc=INK_2, ec=LINE, lw=1, radius=1.2)
    ax.text(chart_x + 2, chart_y + chart_h - 2.3, "PER-ZONE  RISK", color=TXT,
            fontsize=10, fontweight="bold", va="center")
    ax.text(chart_x + chart_w - 2, chart_y + chart_h - 2.3,
            "threshold 0.70", color=SLATE_2, fontsize=8.5,
            ha="right", va="center", style="italic")

    # Bars inside the inset region
    bar_y0 = chart_y + 4
    bar_yh = 18
    bar_x0 = chart_x + 4
    bar_xw = chart_w - 8
    bar_w = bar_xw / 8 * 0.65
    bar_gap = bar_xw / 8 - bar_w
    threshold_y = bar_y0 + bar_yh * 0.70

    # threshold line
    _hline(ax, bar_x0, bar_x0 + bar_xw, threshold_y,
           c=RED, lw=1.2, alpha=0.6)
    # bars
    for i, (z, r) in enumerate(zip(zones, risk)):
        bx = bar_x0 + i * (bar_w + bar_gap) + bar_gap / 2
        h = bar_yh * r
        col = RED if r >= 0.7 else (AMBER if r >= 0.4 else GREEN)
        _rect(ax, bx, bar_y0, bar_w, h, fc=col)
        ax.text(bx + bar_w / 2, bar_y0 - 1.4, z, color=TXT_DIM,
                fontsize=9, ha="center", va="center", fontweight="bold")
        ax.text(bx + bar_w / 2, bar_y0 + h + 1.0, f"{r:.2f}",
                color=col, fontsize=8.5, ha="center", va="center",
                fontweight="bold")

    # Agent log (right)
    log_x = 60; log_y = 7; log_w = 38; log_h = 28
    _box(ax, log_x, log_y, log_w, log_h, fc=INK_2, ec=LINE, lw=1, radius=1.2)
    ax.text(log_x + 2, log_y + log_h - 2.3, "AGENT  LOG", color=TXT,
            fontsize=10, fontweight="bold", va="center")
    log_lines = [
        ("12:04:18", "rule:  LATP 35 % > 30 %  ·  ceramic overload", RED),
        ("12:04:18", "ML:  p_stable = 0.35  (RF · w60)  ·  UNSTABLE", RED),
        ("12:04:19", "physics:  fill Z5 = 0.97  ·  torque 84 %",     AMBER),
        ("12:04:19", "alert raised  ·  zone Z5",                       RED),
        ("12:04:20", "recommendation engine:  READY",                  GREEN),
    ]
    for i, (ts, msg, c) in enumerate(log_lines):
        yy = log_y + log_h - 6 - i * 3.6
        ax.text(log_x + 2,   yy, ts,  color=SLATE_2, fontsize=8.5,
                va="center", family="monospace")
        ax.text(log_x + 8.5, yy, "■", color=c, fontsize=10, va="center")
        ax.text(log_x + 10,  yy, msg, color=TXT, fontsize=9, va="center")

    # Bottom call-out
    _rect(ax, 0, 0, 100, 5, fc=INK_3)
    _rect(ax, 0, 4.7, 100, 0.3, fc=RED)
    ax.text(50, 2.5,
            "diagnosis triggers the recommendation engine  →  see case C4",
            color=TXT_DIM, fontsize=10, ha="center", va="center",
            style="italic")

    fig.savefig(FIG_V3 / "v3_cockpit_risk.png", facecolor=INK, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-06 — COCKPIT C4 (recommendation, full-bleed dark)
# ============================================================================

def v3_cockpit_recommendation():
    fig = plt.figure(figsize=(16, 10.5), facecolor=INK)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 65); ax.axis("off")

    _rect(ax, 0, 0, 100, 65, fc=INK)

    # Topbar
    _rect(ax, 0, 60, 100, 5, fc=INK_3)
    _rect(ax, 0, 60, 100, 0.3, fc=ORANGE)
    ax.text(2, 62.5, "▣ RECO", color=ORANGE, fontsize=10, fontweight="bold", va="center")
    ax.text(10, 62.5, "AI EXTRUSION COPILOT", color=PAPER, fontsize=11.5,
            fontweight="bold", va="center")
    ax.text(98, 62.5, "case C4  ·  diagnostic + ranked actions",
            color=SLATE, fontsize=9.5, ha="right", va="center")

    # Diagnostic banner
    _rect(ax, 0, 52.5, 100, 6.5, fc=INK_2)
    _rect(ax, 0, 52.2, 100, 0.3, fc=ORANGE)
    ax.text(2, 56, "DIAGNOSTIC", color=ORANGE,
            fontsize=11.5, fontweight="bold", va="center")
    ax.text(2, 53.6,
            "ceramic overload (LATP 35 wt%)  ·  overflow risk Z5  ·  torque drift",
            color=PAPER, fontsize=10.5, va="center")

    # 4 ranked action cards — larger
    actions = [
        (1, "FORMULATION",    "Reduce LATP",
         "17–20 wt%",       "highest impact",       GREEN),
        (2, "SCREW PROFILE",  "Soften kneading Z4",
         "45° → 30°",        "reduces local SME",    BLUE),
        (3, "PROCESS",        "Specific mechanical energy",
         "−15 %",            "less viscous heating", BLUE_DEEP),
        (4, "PROCESS",        "Zone 5 temperature",
         "+5 °C",            "improves flow",        ORANGE),
    ]
    aw = 22.5; gap = 2
    total = 4 * aw + 3 * gap
    x0 = (100 - total) / 2
    for i, (rank, target, action, val, sub, c) in enumerate(actions):
        x = x0 + i * (aw + gap)
        _box(ax, x, 18, aw, 30, fc=INK_2, ec=c, lw=2, radius=1.6)
        # accent strip
        _rect(ax, x, 47.5, aw, 0.5, fc=c)
        # rank number — huge
        _box(ax, x + 1.5, 41, 6, 5, fc=c, radius=1.0)
        ax.text(x + 4.5, 43.5, f"#{rank}", color=PAPER,
                fontsize=14, fontweight="bold", ha="center", va="center")
        # target label
        ax.text(x + 9, 43.5, target, color=c,
                fontsize=10, fontweight="bold", va="center")
        # action title
        ax.text(x + aw / 2, 36, action, color=PAPER,
                fontsize=11.5, fontweight="bold", ha="center", va="center")
        # value (big)
        _box(ax, x + 2.5, 25, aw - 5, 7, fc=INK_3, ec=LINE, lw=1, radius=1.0)
        ax.text(x + aw / 2, 28.5, val, color=c,
                fontsize=17, fontweight="bold", ha="center", va="center")
        # sub
        ax.text(x + aw / 2, 21, sub, color=TXT_DIM,
                fontsize=9, ha="center", va="center", style="italic")

    # Projected outcome strip
    _box(ax, 2, 5, 96, 11, fc=INK_2, ec=LINE, lw=1.2, radius=1.6)
    _rect(ax, 2, 15.5, 96, 0.5, fc=GREEN)
    ax.text(4, 13.0, "PROJECTED OUTCOME",
            color=GREEN, fontsize=11, fontweight="bold",
            va="center")
    ax.text(4, 9.0,
            "compatibility score  46  →  70–80      ·      p_stable  ≥ 0.85      ·      Z5 alert cleared",
            color=PAPER, fontsize=11.5, va="center")
    # gain badge
    _box(ax, 78, 7, 18, 6, fc=GREEN, radius=1.4)
    ax.text(87, 10, "+30 PTS",
            color=PAPER, fontsize=13, fontweight="bold",
            ha="center", va="center")

    # Bottom strip
    _rect(ax, 0, 0, 100, 4, fc=INK_3)
    _rect(ax, 0, 3.7, 100, 0.3, fc=ORANGE)
    ax.text(50, 2,
            "operator review  →  apply  →  see case C5",
            color=TXT_DIM, fontsize=9.5, ha="center", va="center",
            style="italic")

    fig.savefig(FIG_V3 / "v3_cockpit_recommendation.png", facecolor=INK, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-07 — OUTCOME (before/after premium)
# ============================================================================

def v3_outcome():
    fig = plt.figure(figsize=(16, 10.5), facecolor=PAPER)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 65); ax.axis("off")

    # Header band — dark
    _rect(ax, 0, 58, 100, 7, fc=INK)
    _rect(ax, 0, 57.7, 100, 0.4, fc=GREEN)
    ax.text(3, 61.5, "OUTCOME",
            color=PAPER, fontsize=15, fontweight="bold", va="center")
    ax.text(97, 61.5, "before  →  after  ·  C3 (at risk)  vs  C5 (recommendation applied)",
            color=TXT_DIM, fontsize=10, ha="right", va="center", style="italic")

    # Lead statement
    ax.text(50, 53.5,
            "The agent changes the operator decision  —  and the change improves the process.",
            color="#1F2937", fontsize=13.5, fontweight="bold",
            ha="center", va="center", style="italic")

    # 4 large tiles
    tiles = [
        ("46",    "78",    "+32",     "COMPATIBILITY SCORE",  RED, GREEN),
        ("0.35",  "0.87",  "+0.52",   "p_stable  (RF · w60)", RED, GREEN),
        ("0.97",  "0.72",  "−0.25",   "FILL FACTOR  Z5",      AMBER, GREEN),
        ("84 %",  "62 %",  "−22 pts", "ESTIMATED TORQUE",     AMBER, GREEN),
    ]
    tw = 22; th = 32
    gap = (100 - 6 - 4 * tw) / 3
    y0 = 14
    for i, (b, a, d, lbl, cbefore, cafter) in enumerate(tiles):
        x = 3 + i * (tw + gap)
        _box(ax, x, y0, tw, th, fc=PAPER_TINT, ec=LINE, lw=1.2, radius=2.0)
        # title
        _rect(ax, x, y0 + th - 5, tw, 5, fc=INK)
        ax.text(x + tw / 2, y0 + th - 2.5, lbl,
                color=PAPER, fontsize=9, fontweight="bold",
                ha="center", va="center")
        # before
        ax.text(x + tw / 4, y0 + th - 12, "BEFORE",
                color=SLATE_2, fontsize=7.5, fontweight="bold",
                ha="center", va="center")
        ax.text(x + tw / 4, y0 + th - 17, b,
                color=cbefore, fontsize=24, fontweight="bold",
                ha="center", va="center")
        ax.text(x + tw / 4, y0 + th - 20.5, "C3",
                color=SLATE_2, fontsize=8.5, ha="center", va="center",
                style="italic")
        # arrow
        _arrow(ax, x + tw / 3 + 1.0, y0 + th - 17,
               x + 2 * tw / 3 - 1.0, y0 + th - 17,
               color=SLATE, lw=2, scale=16)
        # after
        ax.text(x + 3 * tw / 4, y0 + th - 12, "AFTER",
                color=SLATE_2, fontsize=7.5, fontweight="bold",
                ha="center", va="center")
        ax.text(x + 3 * tw / 4, y0 + th - 17, a,
                color=cafter, fontsize=24, fontweight="bold",
                ha="center", va="center")
        ax.text(x + 3 * tw / 4, y0 + th - 20.5, "C5",
                color=SLATE_2, fontsize=8.5, ha="center", va="center",
                style="italic")
        # delta badge
        _box(ax, x + 2, y0 + 2.5, tw - 4, 5.5, fc=GREEN, radius=1.4)
        ax.text(x + tw / 2, y0 + 5.3, f"Δ  {d}",
                color=PAPER, fontsize=12, fontweight="bold",
                ha="center", va="center")

    # Bottom takeaway strip
    _rect(ax, 0, 0, 100, 8, fc=INK_2)
    _rect(ax, 0, 7.7, 100, 0.3, fc=GREEN)
    ax.text(50, 4,
            "Score +32  ·  p_stable +0.52  ·  Z5 alert cleared",
            color=PAPER, fontsize=14, fontweight="bold",
            ha="center", va="center")

    fig.savefig(FIG_V3 / "v3_outcome.png", facecolor=PAPER, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-08 — SCIENCE BACKBONE (ML perf + feature importance compact)
# ============================================================================

def v3_science_backbone():
    fig = plt.figure(figsize=(16, 9.0), facecolor=PAPER)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 56); ax.axis("off")

    # Header
    _rect(ax, 0, 50, 100, 6, fc=INK)
    _rect(ax, 0, 49.7, 100, 0.4, fc=GREEN)
    ax.text(3, 53, "SCIENTIFIC  BACKBONE",
            color=PAPER, fontsize=14, fontweight="bold", va="center")
    ax.text(97, 53, "from raw runs to a production classifier",
            color=TXT_DIM, fontsize=10, ha="right", va="center", style="italic")

    # 4 KPI band
    kpis = [
        ("8",   "INDUSTRIAL RUNS  (≥ 15 min)", PURPLE),
        ("627", "60 s SLIDING WINDOWS",        BLUE),
        ("87",  "FEATURES PER WINDOW",         BLUE),
        ("3",   "CLASSIFIERS  ·  RF / XGB / SVM", ORANGE),
    ]
    kw = 22; gap = (100 - 6 - 4 * kw) / 3
    y_kpi = 38
    for i, (val, lbl, c) in enumerate(kpis):
        x = 3 + i * (kw + gap)
        _box(ax, x, y_kpi, kw, 9, fc=PAPER_TINT, ec=c, lw=2, radius=1.4)
        _rect(ax, x, y_kpi + 8.5, kw, 0.5, fc=c)
        ax.text(x + 5, y_kpi + 4.5, val, color=c,
                fontsize=20, fontweight="bold", va="center")
        ax.text(x + kw - 2, y_kpi + 4.5, lbl, color="#1F2937",
                fontsize=9.5, fontweight="bold",
                ha="right", va="center")

    # 3 model panels (RF/XGB/SVM)
    models = [
        ("RANDOM FOREST",  "PRODUCTION", "0.950 / 0.917 / 0.976", GREEN),
        ("SVM  (RBF)",     "EQUIVALENT", "0.953 / 0.916 / 0.958", BLUE),
        ("XGBOOST",        "REFERENCE",  "0.882 / 0.827 / 0.948", ORANGE),
    ]
    mw = 30.5; gap = 2.0
    total = 3 * mw + 2 * gap
    x0 = (100 - total) / 2
    y_m = 18
    for i, (name, badge, metrics, c) in enumerate(models):
        x = x0 + i * (mw + gap)
        _box(ax, x, y_m, mw, 16, fc=PAPER_TINT, ec=c, lw=2, radius=1.6)
        _rect(ax, x, y_m + 16, mw, 0.5, fc=c)
        ax.text(x + 2, y_m + 13, name, color="#1F2937",
                fontsize=12, fontweight="bold", va="center")
        _box(ax, x + mw - 14, y_m + 11.5, 12, 3.2, fc=c, radius=0.8)
        ax.text(x + mw - 8, y_m + 13, badge, color=PAPER,
                fontsize=8, fontweight="bold", ha="center", va="center")
        ax.text(x + 2, y_m + 7.5,
                "accuracy   ·   F1-macro   ·   ROC-AUC (CV)",
                color=SLATE_2, fontsize=8.5, va="center")
        ax.text(x + 2, y_m + 3.5, metrics, color=c,
                fontsize=14, fontweight="bold", va="center")

    # Bottom takeaway
    _rect(ax, 0, 0, 100, 11, fc=INK_2)
    _rect(ax, 0, 10.6, 100, 0.4, fc=GREEN)
    ax.text(3, 7.2, "INTERPRETABILITY  ·  TOP DRIVERS",
            color=GREEN, fontsize=10, fontweight="bold",
            va="center")
    ax.text(3, 3.5,
            "Variability downstream of the screw  ·  CastFilm P1/P2/Body and DIE "
            "head dominate the top-10 (9 entries) — barrel setpoints rank only "
            "from #14.",
            color=PAPER, fontsize=10.5, va="center")

    fig.savefig(FIG_V3 / "v3_science_backbone.png", facecolor=PAPER, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


# ============================================================================
# FIG v3-09 — ROADMAP (timeline visual)
# ============================================================================

def v3_roadmap():
    fig = plt.figure(figsize=(16, 9.0), facecolor=PAPER)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100); ax.set_ylim(0, 56); ax.axis("off")

    # Header
    _rect(ax, 0, 50, 100, 6, fc=INK)
    _rect(ax, 0, 49.7, 100, 0.4, fc=GREEN)
    ax.text(3, 53, "ROADMAP",
            color=PAPER, fontsize=14, fontweight="bold", va="center")
    ax.text(97, 53, "from demonstrator (TRL 4–5) to industrial product",
            color=TXT_DIM, fontsize=10, ha="right", va="center", style="italic")

    # Timeline axis
    y_tl = 27
    _hline(ax, 6, 94, y_tl, c=LINE, lw=2.5)

    steps = [
        ("Q2-26",   "DATASET\nEXTENSION",
         "Literature corpus of ~50 lithium-bearing recipes  ·  regression score 0–100",
         GREEN, "01"),
        ("Q3-26",   "LOCAL EXPLAINABILITY",
         "SHAP per-window  ·  per-zone risk cartography  ·  audit trail",
         BLUE, "02"),
        ("Q4-26",   "IN-LINE CLOSURE",
         "Sensor feedback  ·  closed-loop recommendation  ·  in-line trigger",
         ORANGE, "03"),
        ("Q1-27",   "SCALE-UP",
         "Validation on Rondol Ø 21 mm platform  ·  productization",
         GREEN_DEEP, "04"),
    ]
    sw = 22; gap = (100 - 12 - 4 * sw) / 3
    x_start = 6
    for i, (when, title, body, c, num) in enumerate(steps):
        x = x_start + i * (sw + gap)
        center = x + sw / 2
        # node
        _box(ax, center - 2.5, y_tl - 2.5, 5, 5, fc=c, radius=2.5)
        ax.text(center, y_tl, num, color=PAPER,
                fontsize=11, fontweight="bold", ha="center", va="center")
        # date pill (above timeline)
        _box(ax, center - 5.5, y_tl + 4, 11, 4, fc=INK_2, radius=1.0)
        ax.text(center, y_tl + 6, when, color=c,
                fontsize=10, fontweight="bold", ha="center", va="center")
        # card (below timeline)
        _box(ax, x, 6, sw, 16, fc=PAPER_TINT, ec=c, lw=2, radius=1.6)
        _rect(ax, x, 21.5, sw, 0.5, fc=c)
        ax.text(x + sw / 2, 19, title, color=c,
                fontsize=10.5, fontweight="bold", ha="center", va="center")
        ax.text(x + sw / 2, 12, body, color="#374151",
                fontsize=9, ha="center", va="center",
                wrap=True, linespacing=1.4)
        # connector tick from node down to card
        _vline(ax, center, y_tl - 2.5, 22, c=c, lw=1.4, alpha=0.5)

    # Bottom takeaway band
    _rect(ax, 0, 0, 100, 4, fc=INK_3)
    _rect(ax, 0, 3.7, 100, 0.3, fc=GREEN)
    ax.text(50, 2,
            "extrusion  ·  AI  ·  batteries  —  documented gap and strategic SSB opportunity",
            color=TXT_DIM, fontsize=10, fontweight="bold",
            ha="center", va="center", style="italic")

    fig.savefig(FIG_V3 / "v3_roadmap.png", facecolor=PAPER, dpi=240,
                bbox_inches=None, pad_inches=0)
    plt.close(fig)


def build_all_figures():
    v3_cover()
    v3_executive_overview()
    v3_architecture()
    v3_five_cases()
    v3_cockpit_risk()
    v3_cockpit_recommendation()
    v3_outcome()
    v3_science_backbone()
    v3_roadmap()
    print("[OK] 9 v3 figures generated in", FIG_V3.relative_to(ROOT))


# ============================================================================
# DOCX — minimal text, full-bleed figures
# ============================================================================

def _set_run_font(run, *, size=10, bold=False, italic=False,
                  color=C("#1F2937"), name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), name)


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color="2A3A5C", size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_cell_no_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_full_bleed_figure(doc, image_path, *, width_cm=19.0):
    """Add a centered figure, almost edge-to-edge (margins are 1cm)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def add_tiny_caption(doc, text, *, dark=False, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _set_run_font(r, size=8, italic=True,
                  color=C("#94A3B8") if dark else C("#64748B"))


def add_section_label(doc, text, *, color="#1FB07A", size=8.5,
                      space_before=4, space_after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text.upper())
    _set_run_font(r, size=size, bold=True, color=C(color))


def setup_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)
        section.footer_distance = Cm(0.4)
        section.header_distance = Cm(0.3)


def install_footer(doc, *, dark=False):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Rondol Industrie  ·  Institut Jean Lamour  ·  Industrial Showcase v3  ·  2026"
    )
    _set_run_font(run, size=7, italic=True,
                  color=C("#94A3B8") if dark else C("#64748B"))


# ----------------------------------------------------------------------------
# Pages
# ----------------------------------------------------------------------------

def page1_cover(doc):
    # Cover image fills the whole page — A4 portrait, sized to fit content area
    # (29.7 - 1.6 cm margins = 28.1 cm tall; aspect 0.707 → max width 19.86 cm,
    # but we use 19.0 cm to leave breathing room and avoid overflow).
    add_full_bleed_figure(doc, FIG_V3 / "v3_cover.png", width_cm=19.0)


def page2_executive(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_executive_overview.png", width_cm=19.4)
    add_tiny_caption(doc,
        "Workflow  ·  the five stages of the AI copilot and the closed loop.",
        space_before=4, space_after=2)


def page3_architecture(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_architecture.png", width_cm=19.4)
    add_tiny_caption(doc,
        "Architecture  ·  four traceable layers — each decision is justified "
        "by a rule, a physical KPI, and a probability.",
        space_before=4, space_after=2)


def page4_five_cases(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_five_cases.png", width_cm=19.4)
    add_tiny_caption(doc,
        "Five reproducible cases  ·  JSON-defined states under "
        "reports/poster_abstract/cases/states/.")


def page5_cockpit_risk(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_cockpit_risk.png", width_cm=19.4)


def page6_cockpit_recommendation(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_cockpit_recommendation.png",
                          width_cm=19.4)


def page7_outcome(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_outcome.png", width_cm=19.4)


def page8_science(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_science_backbone.png", width_cm=19.4)

    # Two scientific figures from v1 as a side-by-side strip
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in (table.cell(0, 0), table.cell(0, 1),
                 table.cell(1, 0), table.cell(1, 1)):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_no_borders(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(str(FIG_V1 / "fig03_ml_performance_w60.png"),
                            width=Cm(9.6))
    p = table.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confusion matrices  ·  60 s window, n = 340.")
    _set_run_font(r, size=8, italic=True, color=C("#64748B"))

    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(str(FIG_V1 / "fig04_feature_importance_RF_w60.png"),
                            width=Cm(9.6))
    p = table.cell(1, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Top-10 RF features  ·  downstream CastFilm + DIE dominate.")
    _set_run_font(r, size=8, italic=True, color=C("#64748B"))


def page9_roadmap(doc):
    add_page_break(doc)
    add_full_bleed_figure(doc, FIG_V3 / "v3_roadmap.png", width_cm=19.4)

    # Minimal trailing references — micro line at the bottom
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("KEY REFERENCES   ")
    _set_run_font(r, size=7.5, bold=True, color=C("#1FB07A"))
    r = p.add_run(
        "Repka, Int. J. Pharm. (2018)  ·  ECHA Annex XV PFAS (2023)  ·  "
        "Drakopoulos, Cell Rep. Phys. Sci. (2021)  ·  Haarmann, Energy Technol. (2021)  ·  "
        "Seeba, Batteries (2024)  ·  Kim, Nat. Commun. (2023)  ·  "
        "Maia, AMI Plastics & Completion AI (2025)  ·  Wang, Nano-Micro Lett. (2025)  ·  "
        "Daoudi, J. Power Sources (2024)  ·  Fraunhofer IWS, DRYtraec® (2024)."
    )
    _set_run_font(r, size=7, color=C("#64748B"))


def build_document():
    doc = Document()
    setup_margins(doc)
    install_footer(doc)

    page1_cover(doc)
    page2_executive(doc)
    page3_architecture(doc)
    page4_five_cases(doc)
    page5_cockpit_risk(doc)
    page6_cockpit_recommendation(doc)
    page7_outcome(doc)
    page8_science(doc)
    page9_roadmap(doc)

    doc.save(DOCX_PATH)
    print(f"[OK] DOCX -> {DOCX_PATH.relative_to(ROOT)}")


def convert_to_pdf():
    try:
        from docx2pdf import convert
    except ImportError:
        print("[WARN] docx2pdf missing — pip install docx2pdf")
        return
    convert(str(DOCX_PATH), str(PDF_PATH))
    print(f"[OK] PDF  -> {PDF_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    build_all_figures()
    build_document()
    convert_to_pdf()
