"""
Génère le poster une-page aligné mail Maël :
  reports/poster_abstract/Mbeumi_2026_ManagerAligned_OnePagePoster_vFinal.docx
  reports/poster_abstract/Mbeumi_2026_ManagerAligned_OnePagePoster_vFinal.pdf

Contraintes :
- Une seule page A4 portrait, marges 1.0 cm
- Logo Rondol intégré (assets/rondol_logo.png)
- 5 sections : Introduction / Materials & Methods / Results / Discussion / Conclusion
- Money shot C3 -> C5 et table de chiffres ML réels
- Aucun ancien fichier modifié
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets"
FIG = ROOT / "reports" / "poster_abstract" / "figures" / "generated"
OUT_DIR = ROOT / "reports" / "poster_abstract"
OUT_DOCX = OUT_DIR / "Mbeumi_2026_ManagerAligned_OnePagePoster_vFinal.docx"
OUT_PDF = OUT_DIR / "Mbeumi_2026_ManagerAligned_OnePagePoster_vFinal.pdf"

C_GREEN = RGBColor(0x0F, 0x6A, 0x3A)
C_BLUE = RGBColor(0x14, 0x49, 0x8B)
C_RED = RGBColor(0xB7, 0x2A, 0x2A)
C_DARK = RGBColor(0x1A, 0x1A, 0x1A)
C_GREY = RGBColor(0x55, 0x55, 0x55)


def shade(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_borders(cell, color_hex: str = "BFBFBF", size: int = 4) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def no_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_margins(cell, top=40, bottom=40, left=80, right=80) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_run(paragraph, text: str, *, bold=False, italic=False, size=8.5, color=C_DARK, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def set_par_spacing(par, before=0, after=2, line=1.05):
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def section_title(par, number: str, label: str, color):
    set_par_spacing(par, before=0, after=1, line=1.0)
    add_run(par, f"{number}  ", bold=True, size=9.5, color=color)
    add_run(par, label.upper(), bold=True, size=9.5, color=color)


def body_par(cell, *, before=0, after=1):
    par = cell.add_paragraph()
    set_par_spacing(par, before=before, after=after, line=1.05)
    return par


def bullet(cell, text_runs):
    par = cell.add_paragraph()
    set_par_spacing(par, before=0, after=0, line=1.05)
    add_run(par, "• ", bold=True, size=8.0, color=C_GREEN)
    for txt, kwargs in text_runs:
        kwargs.setdefault("size", 8.0)
        add_run(par, txt, **kwargs)
    return par


def build() -> None:
    doc = Document()

    # --- Page setup : A4 portrait, marges 1 cm ---
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0.6)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(8.5)

    # ============================================================
    # HEADER : logo (gauche) + titre + bandeau identité (droite)
    # ============================================================
    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_width(header.rows[0].cells[0], 3.6)
    set_cell_width(header.rows[0].cells[1], 15.2)

    logo_cell = header.rows[0].cells[0]
    no_borders(logo_cell)
    cell_margins(logo_cell, 20, 20, 20, 60)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_par_spacing(p_logo, 0, 0, 1.0)
    logo_path = ASSETS / "rondol_logo.png"
    if logo_path.exists():
        p_logo.add_run().add_picture(str(logo_path), width=Cm(3.3))

    title_cell = header.rows[0].cells[1]
    no_borders(title_cell)
    cell_margins(title_cell, 10, 10, 20, 20)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_t = title_cell.paragraphs[0]
    set_par_spacing(p_t, 0, 0, 1.0)
    add_run(
        p_t,
        "AI-Assisted Twin-Screw Hot-Melt Extrusion for Lithium / Solid-State Battery Components",
        bold=True, size=12.5, color=C_DARK,
    )

    p_sub = title_cell.add_paragraph()
    set_par_spacing(p_sub, 0, 0, 1.0)
    add_run(p_sub, "Wilfried Galtier Mbeumi", bold=True, size=7.8, color=C_DARK)
    add_run(p_sub, "  ·  Rondol Industrie & Institut Jean Lamour (IJL) — Campus ARTEM, Nancy, France",
            size=8.5, color=C_GREY)

    p_sup = title_cell.add_paragraph()
    set_par_spacing(p_sup, 0, 0, 1.0)
    add_run(p_sup, "Industrial supervisor: ", size=7.5, color=C_GREY)
    add_run(p_sup, "Maël Gallas", bold=True, size=7.5, color=C_DARK)
    add_run(p_sup, "  ·  Symposium IA + Extrusion + Batteries Li/SSB — May 2026  ·  Mastère Data & IA (RNCP 37137)",
            size=7.5, color=C_GREY)

    # ============================================================
    # ROW WORKFLOW : pipeline image compacte
    # ============================================================
    p_pipe = doc.add_paragraph()
    p_pipe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(p_pipe, 0, 0, 1.0)
    pipe = FIG / "fig01_pipeline_overview.png"
    if pipe.exists():
        p_pipe.add_run().add_picture(str(pipe), width=Cm(15.5))

    p_pipe_cap = doc.add_paragraph()
    p_pipe_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(p_pipe_cap, 0, 1, 1.0)
    add_run(p_pipe_cap, "Fig. 1 — ", bold=True, size=7.5, color=C_BLUE)
    add_run(p_pipe_cap,
            "Boucle industrielle : formulation Li → paramètres procédé → profil de vis → features 60 s "
            "→ prédiction IA → score risque / stabilité → recommandation → essai amélioré.",
            italic=True, size=7.5, color=C_GREY)

    # ============================================================
    # BLOCK 2 COLONNES : sections 1+2 (gauche) / sections 3 (droite)
    # ============================================================
    body = doc.add_table(rows=1, cols=2)
    body.autofit = False
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_width(body.rows[0].cells[0], 9.2)
    set_cell_width(body.rows[0].cells[1], 9.6)
    left = body.rows[0].cells[0]
    right = body.rows[0].cells[1]
    for c in (left, right):
        no_borders(c)
        cell_margins(c, 10, 10, 20, 20)

    # ---------------- LEFT COL ----------------
    # 1. INTRODUCTION
    section_title(left.paragraphs[0], "1.", "Introduction / Background", C_BLUE)
    bullet(left, [
        ("Hot Melt Extrusion (HME)", dict(bold=True, color=C_DARK)),
        (" : procédé continu, sans solvant, mature en pharma, en transfert vers les "
         "électrodes Li-ion et électrolytes solides (SSB).", dict(color=C_DARK)),
    ])
    bullet(left, [
        ("Verrou industriel : ", dict(bold=True, color=C_DARK)),
        ("formulations céramiques (LFP, LATP) abrasives, fenêtre procédé étroite, "
         "essais coûteux, rares jeux de données publiés.", dict(color=C_DARK)),
    ])
    bullet(left, [
        ("Gap scientifique : ", dict(bold=True, color=C_RED)),
        ("intersection ", dict(color=C_DARK)),
        ("extrusion + IA + batteries", dict(italic=True, color=C_DARK)),
        (" peu décrite (Drakopoulos 2021, Haarmann 2021, Kim 2023, Seeba 2024, Maia 2025).",
         dict(color=C_DARK)),
    ])
    bullet(left, [
        ("Objectif : ", dict(bold=True, color=C_GREEN)),
        ("démontrer un outil IA décisionnel formulation → procédé → vis → risque "
         "→ recommandation, sur recette lithiée réelle.", dict(color=C_DARK)),
    ])

    # 2. MATERIALS & METHODS
    p_m = left.add_paragraph()
    set_par_spacing(p_m, before=4, after=1, line=1.0)
    section_title(p_m, "2.", "Materials & Methods", C_BLUE)

    bullet(left, [
        ("Plateforme : ", dict(bold=True, color=C_DARK)),
        ("Rondol bi-vis Ø 10,5 mm, L/D 40:1, 8 zones thermiques, horizontale.",
         dict(color=C_DARK)),
    ])
    bullet(left, [
        ("Recette de référence (C1, semi-dry) : ", dict(bold=True, color=C_DARK)),
        ("LFP 65 / PVDF 8 / Super P 5 / ", dict(color=C_DARK)),
        ("LATP 17", dict(bold=True, color=C_GREEN)),
        (" / LiTFSI 5 wt%.", dict(color=C_DARK)),
    ])
    bullet(left, [
        ("Dataset : ", dict(bold=True, color=C_DARK)),
        ("11 runs (7–13 avril 2026), 8 conservés (≥1 5 min), fenêtres glissantes 60 s pas 30 s "
         "→ ", dict(color=C_DARK)),
        ("627 fenêtres", dict(bold=True, color=C_DARK)),
        (". Split ", dict(color=C_DARK)),
        ("GroupShuffleSplit", dict(italic=True, color=C_DARK)),
        (" par run_id (5 train / 3 test) → zéro fuite inter-run.", dict(color=C_DARK)),
    ])
    bullet(left, [
        ("Architecture IA : ", dict(bold=True, color=C_DARK)),
        ("frontend ", dict(color=C_DARK)),
        ("Streamlit", dict(bold=True, color=C_DARK)),
        (" (Supervision / Profile / Settings / Run Analysis) ; KPI physiques (fill, RT, SME) ; "
         "stack ML ", dict(color=C_DARK)),
        ("Random Forest / XGBoost / SVM", dict(bold=True, color=C_DARK)),
        ("; logique de décision = score formulation (5 critères pondérés) ", dict(color=C_DARK)),
        ("+", dict(bold=True, color=C_DARK)),
        (" classification ML stabilité → panneau risque unifié → recommandation actionnable.",
         dict(color=C_DARK)),
    ])

    # ---------------- RIGHT COL : RESULTS ----------------
    section_title(right.paragraphs[0], "3.", "Results", C_BLUE)

    # Mini-table KPI ML (compacte)
    kpi_tab = right.add_table(rows=2, cols=4)
    kpi_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Model (w=60 s)", "Accuracy", "F1-macro", "ROC-AUC CV"]
    values = [
        ("Random Forest", "0.950", "0.917", "0.976 ± 0.021"),
        ("XGBoost",        "0.882", "0.827", "0.983 ± 0.020"),
        ("SVM (RBF)",      "0.953", "0.916", "0.957 ± 0.037"),
    ]
    for i, h in enumerate(headers):
        c = kpi_tab.rows[0].cells[i]
        shade(c, "0F6A3A")
        set_borders(c, "FFFFFF", 4)
        cell_margins(c, 10, 10, 40, 40)
        p = c.paragraphs[0]
        set_par_spacing(p, 0, 0, 1.0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=8.0, color=RGBColor(0xFF, 0xFF, 0xFF))

    # On a déjà 1 ligne d'en-tête + 1 ligne ; en ajoute 2
    for vals in values[1:]:
        kpi_tab.add_row()
    for r, vals in enumerate(values):
        row = kpi_tab.rows[r + 1]
        for i, v in enumerate(vals):
            c = row.cells[i]
            set_borders(c, "DADADA", 4)
            cell_margins(c, 10, 10, 40, 40)
            p = c.paragraphs[0]
            set_par_spacing(p, 0, 0, 1.0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bold = (r == 0)  # ligne RF en gras
            color = C_GREEN if r == 0 else C_DARK
            add_run(p, v, bold=bold, size=8.0, color=color)

    cap_ml = right.add_paragraph()
    set_par_spacing(cap_ml, 1, 2, 1.0)
    cap_ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap_ml, "Table 1 — ", bold=True, size=7.5, color=C_BLUE)
    add_run(cap_ml,
            "performance test set (n=340), RF retenu pour la production (équilibre perf / interprétabilité).",
            italic=True, size=7.5, color=C_GREY)

    # Top features (texte condensé)
    p_feat = right.add_paragraph()
    set_par_spacing(p_feat, 2, 1, 1.10)
    add_run(p_feat, "Top features (Gini, RF w=60 s) : ", bold=True, size=7.8, color=C_DARK)
    add_run(p_feat,
            "CastFilmP2_iqr 0.072 · DIE_std 0.067 · CastFilmBody_std 0.066 · CastFilmP2_std 0.052 · DIE_iqr 0.052 — "
            "la stabilité dérive ",
            size=7.8, color=C_DARK)
    add_run(p_feat, "aval (cast-film + DIE)", bold=True, size=7.8, color=C_DARK)
    add_run(p_feat, ", pas des consignes de zone.", size=7.8, color=C_DARK)

    # MONEY SHOT C3 -> C5
    money_title = right.add_paragraph()
    set_par_spacing(money_title, 4, 1, 1.0)
    add_run(money_title, "Money shot — cas C3 (surcharge LATP) → C5 (post-reco IA)",
            bold=True, size=9.0, color=C_RED)

    p_money = right.add_paragraph()
    p_money.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_par_spacing(p_money, 0, 1, 1.0)
    money = FIG / "fig05_before_after_recommendation.png"
    if money.exists():
        p_money.add_run().add_picture(str(money), width=Cm(8.4))

    # 4 lignes Δ compactes
    deltas = [
        ("Score compatibilité /100", "46", "78", "+32", C_GREEN),
        ("Prob. stable (RF)",            "0.35", "0.87", "+0.52", C_GREEN),
        ("Fill factor Z5",               "0.97", "0.72", "−0.25", C_GREEN),
        ("Couple théorique %",       "84",   "62",   "−22",   C_GREEN),
    ]
    dt = right.add_table(rows=len(deltas), cols=4)
    dt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (kpi, before, after, delta, col) in enumerate(deltas):
        cells = dt.rows[i].cells
        for c in cells:
            set_borders(c, "EAEAEA", 4)
            cell_margins(c, 6, 6, 30, 30)
        p0 = cells[0].paragraphs[0]
        set_par_spacing(p0, 0, 0, 1.0)
        add_run(p0, kpi, size=8.0, color=C_DARK)
        p1 = cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_par_spacing(p1, 0, 0, 1.0)
        add_run(p1, before, bold=True, size=8.0, color=C_RED)
        p2 = cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_par_spacing(p2, 0, 0, 1.0)
        add_run(p2, after, bold=True, size=8.0, color=C_GREEN)
        p3 = cells[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_par_spacing(p3, 0, 0, 1.0)
        add_run(p3, delta, bold=True, size=8.0, color=col)

    # Reco IA (en bloc encadré)
    reco = right.add_paragraph()
    set_par_spacing(reco, 3, 1, 1.10)
    add_run(reco, "Reco IA déclenchée sur C3 : ", bold=True, size=7.8, color=C_DARK)
    add_run(reco,
            "(1) ramener LATP à 17 wt% ; (2) kneading Z4 → 30° ; (3) baisser SME estimée de 15 % ; "
            "(4) +5 °C en Z5. ⇒ essai C5 stable, alerte Z5 levée.",
            size=7.8, color=C_DARK)

    # Cas C1->C5 enchaînement texte
    cas_par = right.add_paragraph()
    set_par_spacing(cas_par, 2, 1, 1.10)
    add_run(cas_par, "Enchainement essais : ", bold=True, size=7.8, color=C_DARK)
    add_run(cas_par,
            "C1 baseline (65/100 → feu vert) · C2 optimisé (≥80/100) · ",
            size=7.8, color=C_DARK)
    add_run(cas_par, "C3 surcharge LATP 35 % (<50/100, alerte rouge Z5)", bold=True, size=7.8, color=C_RED)
    add_run(cas_par,
            " · C4 reco IA affichée · ",
            size=7.8, color=C_DARK)
    add_run(cas_par, "C5 post-reco (78/100, classe stable, alerte levée)",
            bold=True, size=7.8, color=C_GREEN)
    add_run(cas_par, ".", size=7.8, color=C_DARK)

    # 4. Discussion -> en bas de la colonne gauche
    p_disc = left.add_paragraph()
    set_par_spacing(p_disc, before=4, after=1, line=1.0)
    section_title(p_disc, "4.", "Discussion", C_BLUE)
    fd = left
    fc = right
    bullet(fd, [
        ("Résultats encourageants : ", dict(bold=True, color=C_DARK)),
        ("classification stabilité robuste (F1-macro 0.917, ROC-AUC 0.976), reco IA cohérente avec la "
         "physique de l’extrusion (overflow Z5 ↔ LATP).", dict(color=C_DARK)),
    ])
    bullet(fd, [
        ("Limites dataset : ", dict(bold=True, color=C_DARK)),
        ("8 runs / 627 fenêtres ; généralisation à d’autres formulations Li/SSB (NMC, sulfures) "
         "non encore validée.", dict(color=C_DARK)),
    ])
    bullet(fd, [
        ("Limites scientifiques : ", dict(bold=True, color=C_DARK)),
        ("score formulation rule-based (5 critères) → à remplacer par régression sur ~50 recettes "
         "littérature ; proxy thermique sans validation électrochimique ; pas de bouclage in-line.",
         dict(color=C_DARK)),
    ])
    bullet(fd, [
        ("Perspectives : ", dict(bold=True, color=C_GREEN)),
        ("enrichissement dataset (Maia 2025), capteurs in-line, validation NMC + sulfures, scale-up Rondol 21 mm. ",
         dict(color=C_DARK)),
    ])

    p_concl = fc.add_paragraph()
    set_par_spacing(p_concl, before=4, after=1, line=1.0)
    section_title(p_concl, "5.", "Conclusion", C_GREEN)
    bullet(fc, [
        ("Faisabilité TRL 4–5 : ", dict(bold=True, color=C_DARK)),
        ("outil IA décisionnel pour HME batteries Li/SSB démontré sur cas réel lithié.",
         dict(color=C_DARK)),
    ])
    bullet(fc, [
        ("Valeur industrielle : ", dict(bold=True, color=C_DARK)),
        ("optimisation procédé, aide au scale-up, réduction des essais itératifs, aide décisionnelle "
         "opérateur en temps réel.", dict(color=C_DARK)),
    ])
    bullet(fc, [
        ("Valeur scientifique : ", dict(bold=True, color=C_DARK)),
        ("contribution à un gap publié, ouvre une voie stratégique pour Rondol dans la chaîne de valeur SSB.",
         dict(color=C_DARK)),
    ])
    bullet(fc, [
        ("Prochains jalons : ", dict(bold=True, color=C_GREEN)),
        ("dataset étendu, capteurs in-line, validation électrochimique, plateforme 21 mm.",
         dict(color=C_DARK)),
    ])

    # ============================================================
    # FOOTER REFS / LEGAL
    # ============================================================
    refs = doc.add_paragraph()
    set_par_spacing(refs, 4, 0, 1.05)
    refs.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(refs, "References (sel.) : ", bold=True, size=7.5, color=C_GREY)
    add_run(refs,
            "Drakopoulos 2021 · Haarmann 2021 · Kim 2023 · Seeba 2024 · Daoudi 2024 · Maia 2025 · Wang 2025 · "
            "Fraunhofer IWS — DRYtraec®. ", italic=True, size=7.5, color=C_GREY)
    add_run(refs,
            "Acknowledgements : ", bold=True, size=7.5, color=C_GREY)
    add_run(refs,
            "Rondol Industrie ; Institut Jean Lamour (IJL) ; Campus ARTEM ; Mastère Data & IA (RNCP 37137). "
            "PFAS / PVDF — réglementation ECHA février 2023.",
            italic=True, size=7.5, color=C_GREY)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"[OK] DOCX  : {OUT_DOCX}")

    # PDF via docx2pdf (MS Word requis sous Windows)
    try:
        from docx2pdf import convert
        convert(str(OUT_DOCX), str(OUT_PDF))
        print(f"[OK] PDF   : {OUT_PDF}")
    except Exception as exc:
        print(f"[WARN] PDF non généré ({exc!r})")


if __name__ == "__main__":
    build()
